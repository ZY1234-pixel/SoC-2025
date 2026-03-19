"""HTML 表格到 python-docx Table 对象的转换器。

Migrated from original PaddleOCR table_process.py (~326 lines).
Converts an HTML ``<table>`` string into a native python-docx Table.

Dependencies: re, docx, bs4 (BeautifulSoup), html.parser (HTMLParser).
"""
from __future__ import annotations

import io
import re
import logging
from functools import lru_cache
from html.parser import HTMLParser
from typing import List, Optional, Tuple

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

logger = logging.getLogger(__name__)

_DENSE_PLAIN_CELL_MIN_LEN = 80
_MASK_VALUE_MAX = 190
_MASK_SAT_MAX = 55
_RULE_ROW_OCCUPANCY_MAX = 0.72

# ---------------------------------------------------------------------------
# 字体样式 / 名称映射
# ---------------------------------------------------------------------------

font_styles = {
    "b": "bold",
    "strong": "bold",
    "em": "italic",
    "i": "italic",
    "u": "underline",
    "s": "strike",
    "sup": "superscript",
    "sub": "subscript",
}

font_names = {
    "serif": "Times New Roman",
    "sans-serif": "Arial",
    "monospace": "Courier New",
}


# ---------------------------------------------------------------------------
# HTML 表格解析工具函数
# ---------------------------------------------------------------------------

def remove_whitespace(text: str) -> str:
    """将空白字符合并为单个空格并去除首尾空白。"""
    return re.sub(r"\s+", " ", text).strip()


def delete_paragraph(paragraph):
    """从 XML 树的父节点中移除一个段落元素。"""
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def get_table_rows(table_soup) -> list:
    """从 BeautifulSoup 表格元素中返回所有 ``<tr>`` 行。

    同时处理 ``<thead>/<tbody>`` 包裹的行和直接嵌套的 ``<tr>`` 子元素。
    """
    rows = []
    # 先尝试 thead/tbody/tfoot 内的行
    for section in table_soup.find_all(["thead", "tbody", "tfoot"], recursive=False):
        rows.extend(section.find_all("tr", recursive=False))
    # 再处理直接嵌套的 <tr> 子元素（若不在任何 section 内）
    for tr in table_soup.find_all("tr", recursive=False):
        if tr not in rows:
            rows.append(tr)
    return rows


def get_table_columns(row_soup) -> list:
    """返回单行中的所有 ``<td>`` 和 ``<th>`` 单元格。"""
    return row_soup.find_all(["td", "th"], recursive=False)


def get_table_dimensions(table_soup) -> Tuple[int, int]:
    """返回 HTML 表格的 ``(行数, 列数)``。

    考虑 row/colspan 属性来计算最大逻辑列数。
    """
    rows = get_table_rows(table_soup)
    num_rows = len(rows)
    max_cols = 0
    for row in rows:
        cols = get_table_columns(row)
        col_count = 0
        for cell in cols:
            colspan = int(cell.get("colspan", 1))
            col_count += colspan
        max_cols = max(max_cols, col_count)
    return num_rows, max_cols


def get_cell_html(cell_soup) -> str:
    """返回 ``<td>``/``<th>`` 元素的内部 HTML 字符串。"""
    return cell_soup.decode_contents()


# ---------------------------------------------------------------------------
# HtmlToDocx 转换器类
# ---------------------------------------------------------------------------

class HtmlToDocx:
    """将 HTML ``<table>`` 字符串转换为 python-docx ``Table``。

    用法::

        converter = HtmlToDocx()
        converter.table_style = "Table Grid"
        converter.handle_table(html_string, doc_or_cell)

    解析器使用 BeautifulSoup 解析 HTML，遍历行和单元格，
    创建具有正确尺寸的 docx 表格（包括 row/colspan 合并），
    并填充单元格文本。基本的行内样式（加粗、斜体、颜色、
    字号、对齐）会被保留。
    """

    def __init__(self):
        self.table_style: Optional[str] = None
        self.default_font_name: str = "Times New Roman"
        self.default_east_asia: str = "\u5b8b\u4f53"  # 宋体
        self.default_font_size: float = 10.0  # pt

    # -----------------------------------------------------------------
    # 公开入口
    # -----------------------------------------------------------------

    def handle_table(
        self,
        html_string: str,
        container,
        table_image_data: Optional[bytes] = None,
        table_width_pt: Optional[float] = None,
    ) -> Optional[object]:
        """解析 *html_string* 并向 *container* 添加表格。

        Args:
            html_string: An HTML string containing a ``<table>`` element.
            container: A python-docx container (Document or Cell).

        Returns:
            The python-docx Table object, or None on failure.
        """
        try:
            soup = BeautifulSoup(html_string, "html.parser")
            table_soup = soup.find("table")
            if table_soup is None:
                logger.warning("No <table> element found in HTML.")
                return None

            num_rows, num_cols = get_table_dimensions(table_soup)
            if num_rows == 0 or num_cols == 0:
                logger.warning("Empty table (%d rows, %d cols).", num_rows, num_cols)
                return None

            table = container.add_table(rows=num_rows, cols=num_cols)
            if self.table_style:
                table.style = self.table_style
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True

            # 跟踪已合并区域: grid[row][col] -> True if already consumed
            occupied = [[False] * num_cols for _ in range(num_rows)]

            table_image = None
            x_bounds = None
            y_bounds = None
            if table_image_data:
                try:
                    table_image = Image.open(io.BytesIO(table_image_data)).convert("RGB")
                    x_bounds = self._infer_axis_bounds(table_image, num_cols, axis="x")
                    y_bounds = self._infer_axis_bounds(table_image, num_rows, axis="y")
                except Exception:
                    logger.debug("Table image analysis unavailable.", exc_info=True)
                    table_image = None

            rows = get_table_rows(table_soup)
            for row_idx, row_soup in enumerate(rows):
                if row_idx >= num_rows:
                    break
                cells = get_table_columns(row_soup)
                col_cursor = 0
                for cell_soup in cells:
                    # 跳过已占用的单元格 (from previous rowspans)
                    while col_cursor < num_cols and occupied[row_idx][col_cursor]:
                        col_cursor += 1
                    if col_cursor >= num_cols:
                        break

                    colspan = int(cell_soup.get("colspan", 1))
                    rowspan = int(cell_soup.get("rowspan", 1))
                    colspan = min(colspan, num_cols - col_cursor)
                    rowspan = min(rowspan, num_rows - row_idx)

                    # 标记已占用的网格
                    for ri in range(row_idx, row_idx + rowspan):
                        for ci in range(col_cursor, col_cursor + colspan):
                            occupied[ri][ci] = True

                    # 按需合并单元格
                    top_left = table.cell(row_idx, col_cursor)
                    bottom_right = table.cell(
                        row_idx + rowspan - 1,
                        col_cursor + colspan - 1,
                    )
                    if rowspan > 1 or colspan > 1:
                        try:
                            top_left.merge(bottom_right)
                        except Exception:
                            logger.debug(
                                "Cell merge failed at (%d,%d).", row_idx, col_cursor
                            )

                    cell_image = None
                    cell_width_pt = None
                    if table_image is not None and x_bounds and y_bounds:
                        cell_box = self._grid_cell_box(
                            x_bounds=x_bounds,
                            y_bounds=y_bounds,
                            row_idx=row_idx,
                            col_idx=col_cursor,
                            rowspan=rowspan,
                            colspan=colspan,
                        )
                        cell_image = table_image.crop(cell_box)
                        if table_width_pt:
                            cell_width_pt = float(table_width_pt) * (colspan / max(num_cols, 1))

                    # 填充单元格内容
                    self.add_html_to_cell(
                        top_left,
                        cell_soup,
                        cell_image=cell_image,
                        cell_width_pt=cell_width_pt,
                    )

                    col_cursor += colspan

            return table

        except Exception:
            logger.warning("Failed to convert HTML table.", exc_info=True)
            return None

    # -----------------------------------------------------------------
    # 单元格内容
    # -----------------------------------------------------------------

    def add_html_to_cell(self, cell, cell_soup, cell_image=None, cell_width_pt: Optional[float] = None):
        """用 *cell_soup* 的内容填充 docx 表格 *cell*。

        处理文本内容、基本行内格式标签 (``<b>``、``<i>``、``<br>``)，
        并提取简单的样式属性 (font-size、color、text-align)。
        """
        # 清除已有的默认段落
        if cell.paragraphs:
            for i, p in enumerate(cell.paragraphs):
                if i == 0:
                    # 复用第一个段落
                    continue
                delete_paragraph(p)
            para = cell.paragraphs[0]
            para.clear()
        else:
            para = cell.add_paragraph()

        # 从 style 属性检测对齐方式
        style_str = cell_soup.get("style", "")
        alignment = self._parse_alignment(style_str)
        if alignment is not None:
            para.alignment = alignment

        # 检测背景色
        bg_color = self._parse_bg_color(style_str)
        if bg_color:
            self._set_cell_shading(cell, bg_color)

        dense_specs = self._infer_dense_plain_text_layout(
            cell_soup,
            cell_image=cell_image,
            cell_width_pt=cell_width_pt,
        )
        if dense_specs:
            self._write_inferred_dense_paragraphs(cell, para, dense_specs)
            return

        dense_lines = self._extract_dense_plain_text_lines(cell_soup)
        if dense_lines:
            self._write_dense_plain_text_cell(para, dense_lines)
            return

        # 遍历子节点并添加 run
        self._process_children(para, cell_soup, cell)

    @staticmethod
    def _is_dense_plain_text_cell(cell_soup) -> bool:
        child_tags = [
            child.name
            for child in cell_soup.children
            if getattr(child, "name", None)
        ]
        if any(tag not in {"br"} for tag in child_tags):
            return False

        text = cell_soup.get_text(" ", strip=True)
        if len(text) < _DENSE_PLAIN_CELL_MIN_LEN:
            return False

        alpha_like = sum(
            1
            for ch in text
            if ch.isalpha() or ("\u4e00" <= ch <= "\u9fff")
        )
        return alpha_like / max(len(text), 1) >= 0.35

    def _infer_dense_plain_text_layout(
        self,
        cell_soup,
        cell_image=None,
        cell_width_pt: Optional[float] = None,
    ) -> Optional[List[dict]]:
        """将长纯文本单元格视为“小页面”，从单元格图像中推断行与段落。"""
        if cell_image is None or not self._is_dense_plain_text_cell(cell_soup):
            return None

        visual_lines = self._detect_visual_lines(cell_image)
        if len(visual_lines) < 2:
            return None

        normalized = remove_whitespace(cell_soup.get_text(" ", strip=True).replace("\xa0", " "))
        line_texts = self._wrap_text_to_visual_lines(
            normalized,
            [line["width_ratio"] for line in visual_lines],
        )
        if len(line_texts) != len(visual_lines):
            return None

        cell_max_units = max((self._text_units(text) for text in line_texts if text), default=0)
        para_ranges = self._group_visual_lines_into_paragraphs(visual_lines, line_texts)
        specs: List[dict] = []
        for start, end in para_ranges:
            geom = visual_lines[start:end]
            texts = [text for text in line_texts[start:end] if text]
            if not geom or not texts:
                continue
            if start == 0 and end == len(visual_lines):
                paragraph_role = "body"
            elif start == 0:
                paragraph_role = "head"
            elif end == len(visual_lines):
                paragraph_role = "tail"
            else:
                paragraph_role = "body"
            alignment = self._infer_visual_alignment(
                geom,
                texts=texts,
                cell_max_units=cell_max_units,
                paragraph_role=paragraph_role,
            )
            indent_pt = self._infer_visual_indent_pt(geom, cell_width_pt)
            specs.append({
                "alignment": alignment,
                "indent_pt": indent_pt,
                "lines": texts,
                "paragraph_role": paragraph_role,
                "preserve_breaks": self._should_preserve_dense_paragraph_breaks(
                    paragraph_role=paragraph_role,
                    alignment=alignment,
                    lines=texts,
                    indent_pt=indent_pt,
                ),
            })
        return specs or None

    def _extract_dense_plain_text_lines(self, cell_soup) -> Optional[List[str]]:
        """识别“长纯文本单元格”并按句子切成多行。"""
        if not self._is_dense_plain_text_cell(cell_soup):
            return None

        text = cell_soup.get_text(" ", strip=True)
        lines = self._split_dense_plain_text(text)
        return lines if len(lines) >= 2 else None

    @staticmethod
    def _split_dense_plain_text(text: str) -> List[str]:
        """将长纯文本单元格按句末标点切分为视觉行。"""
        normalized = remove_whitespace(text.replace("\xa0", " ")).strip()
        if not normalized:
            return []

        lines: List[str] = []
        buf: List[str] = []
        length = len(normalized)
        for idx, ch in enumerate(normalized):
            buf.append(ch)
            next_ch = normalized[idx + 1] if idx + 1 < length else ""
            if ch in ".!?。！？；;":
                if next_ch and next_ch.isdigit():
                    continue
                seg = "".join(buf).strip()
                if seg:
                    lines.append(seg)
                buf = []

        tail = "".join(buf).strip()
        if tail:
            lines.append(tail)

        if len(lines) >= 2:
            return lines
        return [normalized]

    def _write_dense_plain_text_cell(self, para, lines: List[str]) -> None:
        """将拆分后的多行文本写入单元格，行间使用软换行。"""
        for idx, line in enumerate(lines):
            run = para.add_run(line)
            self._apply_run_defaults(run)
            if idx < len(lines) - 1:
                run.add_break()

    def _write_inferred_dense_paragraphs(self, cell, first_para, specs: List[dict]) -> None:
        for idx, spec in enumerate(specs):
            para = first_para if idx == 0 else cell.add_paragraph()
            if idx > 0:
                para.clear()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(1.0 if idx == len(specs) - 1 else 0.5)
            para.paragraph_format.line_spacing = 1.0
            para.alignment = spec["alignment"]
            indent_pt = float(spec.get("indent_pt") or 0.0)
            if indent_pt > 0:
                para.paragraph_format.first_line_indent = Pt(indent_pt)

            preserve_breaks = bool(spec.get("preserve_breaks"))
            if preserve_breaks:
                for li, line in enumerate(spec["lines"]):
                    run = para.add_run(line)
                    self._apply_run_defaults(run)
                    if li < len(spec["lines"]) - 1:
                        run.add_break()
                continue

            run = para.add_run(self._join_dense_paragraph_lines(spec["lines"]))
            self._apply_run_defaults(run)

    @staticmethod
    def _is_textish_pixel(rgb: Tuple[int, int, int]) -> bool:
        r, g, b = rgb
        mx = max(r, g, b)
        mn = min(r, g, b)
        sat = mx - mn
        val = (r + g + b) / 3.0
        return val < _MASK_VALUE_MAX and sat < _MASK_SAT_MAX

    def _infer_axis_bounds(self, image: Image.Image, count: int, axis: str) -> List[int]:
        size = image.size[0 if axis == "x" else 1]
        if count <= 1:
            return [0, size]

        projections = self._axis_projection(image, axis)
        step = size / count
        bounds = [0]
        for idx in range(1, count):
            guess = int(round(step * idx))
            search = max(12, int(step * 0.18))
            lo = max(bounds[-1] + 1, guess - search)
            hi = min(size - 2, guess + search)
            best = min(range(lo, hi + 1), key=lambda pos: projections[pos])
            bounds.append(best)
        bounds.append(size)
        return bounds

    def _axis_projection(self, image: Image.Image, axis: str) -> List[int]:
        width, height = image.size
        pixels = image.load()
        result: List[int] = []
        if axis == "x":
            for x in range(width):
                count = 0
                for y in range(height):
                    if self._is_textish_pixel(pixels[x, y]):
                        count += 1
                result.append(count)
        else:
            for y in range(height):
                count = 0
                for x in range(width):
                    if self._is_textish_pixel(pixels[x, y]):
                        count += 1
                result.append(count)
        return result

    @staticmethod
    def _grid_cell_box(
        x_bounds: List[int],
        y_bounds: List[int],
        row_idx: int,
        col_idx: int,
        rowspan: int,
        colspan: int,
    ) -> Tuple[int, int, int, int]:
        return (
            x_bounds[col_idx],
            y_bounds[row_idx],
            x_bounds[min(col_idx + colspan, len(x_bounds) - 1)],
            y_bounds[min(row_idx + rowspan, len(y_bounds) - 1)],
        )

    def _detect_visual_lines(self, cell_image: Image.Image) -> List[dict]:
        width, height = cell_image.size
        inset_x = max(8, int(width * 0.03))
        inset_y = max(8, int(height * 0.015))
        inner = cell_image.crop((
            inset_x,
            inset_y,
            max(inset_x + 1, width - inset_x),
            max(inset_y + 1, height - inset_y),
        ))
        iw, ih = inner.size
        pixels = inner.load()

        pixel_mask: List[List[int]] = []
        row_occupancies: List[float] = []
        for y in range(ih):
            row: List[int] = []
            dark = 0
            for x in range(iw):
                bit = 1 if self._is_textish_pixel(pixels[x, y]) else 0
                row.append(bit)
                if bit:
                    dark += 1
            pixel_mask.append(row)
            row_occupancies.append(dark / max(iw, 1))

        row_counts: List[int] = [sum(row) for row in pixel_mask]
        high_occ_start: Optional[int] = None
        for idx, occupancy in enumerate(row_occupancies + [0.0]):
            if occupancy >= _RULE_ROW_OCCUPANCY_MAX and high_occ_start is None:
                high_occ_start = idx
                continue
            if occupancy >= _RULE_ROW_OCCUPANCY_MAX:
                continue
            if high_occ_start is None:
                continue
            # 只清理很薄的高占用条带，避免把真实满行正文误判为边框。
            if idx - high_occ_start <= 3:
                for row_idx in range(high_occ_start, idx):
                    pixel_mask[row_idx] = [0] * iw
                    row_counts[row_idx] = 0
            high_occ_start = None

        threshold = max(2, int(iw * 0.016))
        raw_lines: List[Tuple[int, int]] = []
        in_run = False
        start = 0
        for idx, value in enumerate(row_counts):
            if value > threshold and not in_run:
                in_run = True
                start = idx
            elif value <= threshold and in_run:
                if idx - start >= 8:
                    raw_lines.append((start, idx - 1))
                in_run = False
        if in_run and ih - start >= 8:
            raw_lines.append((start, ih - 1))

        lines: List[dict] = []
        for y1, y2 in raw_lines:
            col_counts: List[int] = []
            row_occupancies = [
                sum(pixel_mask[y]) / max(iw, 1)
                for y in range(y1, y2 + 1)
            ]
            for x in range(iw):
                col_counts.append(sum(pixel_mask[y][x] for y in range(y1, y2 + 1)))
            x_thresh = max(2, int((y2 - y1 + 1) * 0.20))
            xs = [x for x, value in enumerate(col_counts) if value >= x_thresh]
            if not xs:
                continue
            left = min(xs)
            right = max(xs)
            lines.append({
                "y1": y1,
                "y2": y2,
                "left_ratio": left / max(iw, 1),
                "right_ratio": max(iw - right - 1, 0) / max(iw, 1),
                "width_ratio": (right - left + 1) / max(iw, 1),
                "avg_row_occupancy": sum(row_occupancies) / max(len(row_occupancies), 1),
            })
        if lines:
            occupancies = sorted(float(line["avg_row_occupancy"]) for line in lines)
            median_occupancy = occupancies[len(occupancies) // 2]
            reference_widths = sorted(
                float(line["width_ratio"])
                for line in lines
                if float(line["avg_row_occupancy"]) >= median_occupancy * 0.8
            )
            reference_width = reference_widths[len(reference_widths) // 2] if reference_widths else 0.95
            for line in lines:
                occupancy = float(line["avg_row_occupancy"])
                if (
                    float(line["width_ratio"]) >= 0.97
                    and median_occupancy > 0
                    and occupancy < median_occupancy * 0.65
                ):
                    line["width_ratio"] = max(
                        0.22,
                        min(reference_width, reference_width * ((occupancy / median_occupancy) ** 0.5)),
                    )
        return lines

    def _group_visual_lines_into_paragraphs(
        self,
        lines: List[dict],
        line_texts: Optional[List[str]] = None,
    ) -> List[Tuple[int, int]]:
        if not lines:
            return []

        ranges: List[Tuple[int, int]] = []
        start = 0
        gaps = [
            max(0.0, float(lines[idx + 1]["y1"]) - float(lines[idx]["y2"]))
            for idx in range(len(lines) - 1)
        ]
        positive_gaps = [gap for gap in gaps if gap > 0]
        median_gap = sorted(positive_gaps)[len(positive_gaps) // 2] if positive_gaps else 14.0
        large_gap = max(median_gap * 2.2, 36.0)
        text_units = [
            self._text_units(line_texts[idx]) if line_texts and idx < len(line_texts) else 0
            for idx in range(len(lines))
        ]
        max_units = max(text_units, default=0)
        effective_widths = [
            self._effective_line_width_ratio(line, text_units[idx], max_units)
            for idx, line in enumerate(lines)
        ]

        if (
            len(lines) >= 3
            and not (
                float(lines[0]["left_ratio"]) > 0.18
                and float(lines[0]["right_ratio"]) > 0.18
                and abs(float(lines[0]["left_ratio"]) - float(lines[0]["right_ratio"])) <= 0.14
            )
            and max(text_units[2:min(4, len(lines))], default=0) > 0
            and max(text_units[0], text_units[1]) <= max(text_units[2:min(4, len(lines))]) * 0.68
            and max(effective_widths[2:min(4, len(lines))], default=0.0) >= 0.76
        ):
            ranges.append((0, 2))
            start = 2
        elif (
            len(lines) >= 2
            and effective_widths[0] <= 0.74
            and effective_widths[1] >= min(0.88, effective_widths[0] + 0.12)
        ):
            ranges.append((0, 1))
            start = 1

        split_points: List[int] = []
        for idx in range(max(start, 1), len(lines)):
            gap = float(lines[idx]["y1"]) - float(lines[idx - 1]["y2"])
            if gap >= large_gap:
                split_points.append(idx)

        tail_split = self._infer_tail_cluster_start(
            lines=lines,
            text_units=text_units,
            effective_widths=effective_widths,
            start=start,
            split_points=split_points,
            median_gap=median_gap,
        )
        if tail_split is not None:
            split_points.append(tail_split)

        cursor = start
        for cut in sorted(set(point for point in split_points if point > cursor)):
            ranges.append((cursor, cut))
            cursor = cut
        ranges.append((cursor, len(lines)))
        return [(lo, hi) for lo, hi in ranges if hi > lo]

    @staticmethod
    def _tokenize_for_wrap(text: str) -> Tuple[List[str], bool]:
        if not text:
            return [], False
        if " " not in text:
            return list(text), True
        return text.split(), False

    @staticmethod
    def _join_dense_paragraph_lines(lines: List[str]) -> str:
        cleaned = [line.strip() for line in lines if line and line.strip()]
        if not cleaned:
            return ""
        merged = "".join(cleaned)
        cjk_count = sum(1 for ch in merged if "\u4e00" <= ch <= "\u9fff")
        latin_like = sum(1 for ch in merged if ch.isalpha() or ch.isdigit())
        if cjk_count > latin_like:
            return "".join(cleaned)
        return remove_whitespace(" ".join(cleaned))

    @staticmethod
    def _should_preserve_dense_paragraph_breaks(
        paragraph_role: str,
        alignment,
        lines: List[str],
        indent_pt: float,
    ) -> bool:
        if len(lines) <= 1:
            return False
        if paragraph_role == "tail":
            return True
        if paragraph_role == "body":
            return False
        # 段首短引言通常也应作为普通段落重排，只保留明显居中标题的换行。
        if paragraph_role == "head":
            return alignment == WD_ALIGN_PARAGRAPH.CENTER and indent_pt <= 0 and len(lines) >= 3
        return False

    @staticmethod
    def _text_units(text: str) -> int:
        normalized = remove_whitespace(text.replace("\xa0", " ")).strip()
        if not normalized:
            return 0
        tokens, char_mode = HtmlToDocx._tokenize_for_wrap(normalized)
        if not tokens:
            return 0
        return sum(len(token) for token in tokens) + (0 if char_mode else max(0, len(tokens) - 1))

    @staticmethod
    def _effective_line_width_ratio(line: dict, text_units: int, cell_max_units: int) -> float:
        visual_ratio = float(line["width_ratio"])
        if cell_max_units <= 0 or text_units <= 0:
            return visual_ratio
        text_ratio = text_units / max(cell_max_units, 1)
        return min(visual_ratio, max(0.18, text_ratio * 1.08))

    def _infer_tail_cluster_start(
        self,
        lines: List[dict],
        text_units: List[int],
        effective_widths: List[float],
        start: int,
        split_points: List[int],
        median_gap: float,
    ) -> Optional[int]:
        base_start = max([start] + split_points) if split_points else start
        if len(lines) - base_start < 5:
            return None

        for tail_size in (2, 3):
            tail_start = len(lines) - tail_size
            if tail_start <= base_start:
                continue
            body_units = [unit for unit in text_units[base_start:tail_start] if unit > 0]
            tail_units = [unit for unit in text_units[tail_start:] if unit > 0]
            if len(body_units) < 3 or len(tail_units) != tail_size:
                continue

            body_avg = sum(body_units) / len(body_units)
            tail_avg = sum(tail_units) / len(tail_units)
            body_max = max(body_units)
            tail_max = max(tail_units)
            body_width = sum(effective_widths[base_start:tail_start]) / max(tail_start - base_start, 1)
            tail_width = sum(effective_widths[tail_start:]) / tail_size
            prev_gap = max(0.0, float(lines[tail_start]["y1"]) - float(lines[tail_start - 1]["y2"]))
            tail_similar = min(tail_units) / max(tail_max, 1) >= 0.55

            if (
                tail_avg <= body_avg * 0.72
                and tail_max <= body_max * 0.82
                and tail_similar
                and (
                    tail_width <= body_width * 0.84
                    or prev_gap >= max(median_gap * 1.35, 18.0)
                )
            ):
                return tail_start
        return None

    def _wrap_text_to_visual_lines(self, text: str, width_ratios: List[float]) -> List[str]:
        tokens, char_mode = self._tokenize_for_wrap(text)
        if not tokens:
            return []
        if len(tokens) <= len(width_ratios):
            lines = [self._join_tokens([tok], char_mode) for tok in tokens]
            while len(lines) < len(width_ratios):
                lines.append("")
            return lines

        token_lengths = [len(tok) for tok in tokens]
        total_units = sum(token_lengths) + (0 if char_mode else max(0, len(tokens) - 1))
        ratios = [max(0.2, float(r)) for r in width_ratios]
        total_ratio = sum(ratios)
        target_units = [max(4.0, total_units * ratio / total_ratio) for ratio in ratios]
        prefix_units = [0]
        for token_len in token_lengths:
            prefix_units.append(prefix_units[-1] + token_len)

        def span_units(start_idx: int, end_idx: int) -> int:
            raw_units = prefix_units[end_idx] - prefix_units[start_idx]
            if char_mode:
                return raw_units
            return raw_units + max(0, end_idx - start_idx - 1)

        @lru_cache(maxsize=None)
        def solve(line_idx: int, start_idx: int) -> Tuple[float, Tuple[int, ...]]:
            remaining_lines = len(target_units) - line_idx
            remaining_tokens = len(tokens) - start_idx
            if remaining_lines == 0:
                return (0.0, ()) if start_idx == len(tokens) else (float("inf"), ())
            if remaining_tokens < remaining_lines:
                return float("inf"), ()
            if line_idx == len(target_units) - 1:
                diff = span_units(start_idx, len(tokens)) - target_units[line_idx]
                return diff * diff * 1.8, (len(tokens),)

            best_cost = float("inf")
            best_path: Tuple[int, ...] = ()
            min_end = start_idx + 1
            max_end = len(tokens) - (remaining_lines - 1)
            for end_idx in range(min_end, max_end + 1):
                diff = span_units(start_idx, end_idx) - target_units[line_idx]
                local_cost = diff * diff * (1.7 if diff > 0 else 1.0)
                future_cost, future_path = solve(line_idx + 1, end_idx)
                total_cost = local_cost + future_cost
                if total_cost < best_cost:
                    best_cost = total_cost
                    best_path = (end_idx,) + future_path
            return best_cost, best_path

        _, path = solve(0, 0)
        if not path:
            return [remove_whitespace(text)] + [""] * (len(width_ratios) - 1)

        lines: List[str] = []
        cursor = 0
        for end_idx in path:
            lines.append(self._join_tokens(tokens[cursor:end_idx], char_mode))
            cursor = end_idx
        while len(lines) < len(width_ratios):
            lines.append("")
        return lines[:len(width_ratios)]

    @staticmethod
    def _join_tokens(tokens: List[str], char_mode: bool) -> str:
        if char_mode:
            return "".join(tokens).strip()
        return " ".join(tokens).strip()

    def _infer_visual_alignment(
        self,
        lines: List[dict],
        texts: Optional[List[str]] = None,
        cell_max_units: int = 0,
        paragraph_role: str = "body",
    ):
        lefts = [float(line["left_ratio"]) for line in lines]
        rights = [float(line["right_ratio"]) for line in lines]
        visual_widths = [float(line["width_ratio"]) for line in lines]
        text_units = [
            self._text_units(texts[idx])
            for idx in range(len(lines))
        ] if texts else [0] * len(lines)
        widths = [
            self._effective_line_width_ratio(line, text_units[idx], cell_max_units)
            for idx, line in enumerate(lines)
        ]
        short_two_line_block = (
            len(lines) == 2
            and cell_max_units > 0
            and max(text_units, default=0) <= cell_max_units * 0.72
            and min(text_units, default=0) / max(max(text_units, default=1), 1) >= 0.60
        )

        if len(lines) == 1:
            left = lefts[0]
            right = rights[0]
            if left > 0.18 and right > 0.18 and abs(left - right) <= 0.12:
                return WD_ALIGN_PARAGRAPH.CENTER
            if left > 0.34 and right < 0.16:
                return WD_ALIGN_PARAGRAPH.RIGHT
            return WD_ALIGN_PARAGRAPH.LEFT

        body = lines[:-1] if len(lines) >= 3 else lines
        left_hit = sum(1 for line in lines if float(line["left_ratio"]) <= 0.08) / len(lines)
        right_hit = sum(1 for line in body if float(line["right_ratio"]) <= 0.10) / max(len(body), 1)
        avg_left = sum(lefts) / len(lefts)
        avg_right = sum(rights) / len(rights)
        avg_width = sum(widths) / len(widths)

        if left_hit >= 0.70 and right_hit >= 0.70 and avg_width >= 0.84 and not short_two_line_block:
            return WD_ALIGN_PARAGRAPH.JUSTIFY
        if avg_left > 0.24 and avg_right < 0.16:
            return WD_ALIGN_PARAGRAPH.RIGHT
        if avg_left > 0.16 and avg_right > 0.16 and abs(avg_left - avg_right) <= 0.10:
            return WD_ALIGN_PARAGRAPH.CENTER
        if short_two_line_block and paragraph_role == "tail":
            return WD_ALIGN_PARAGRAPH.CENTER
        return WD_ALIGN_PARAGRAPH.LEFT

    @staticmethod
    def _infer_visual_indent_pt(lines: List[dict], cell_width_pt: Optional[float]) -> float:
        if not cell_width_pt or len(lines) < 2:
            return 0.0
        first_left = float(lines[0]["left_ratio"])
        rest_left = min(float(line["left_ratio"]) for line in lines[1:])
        delta = first_left - rest_left
        if delta < 0.08:
            return 0.0
        return max(0.0, float(cell_width_pt) * delta * 0.95)

    def _process_children(self, para, element, cell):
        """递归遍历 *element* 的子节点并向 *para* 添加 run。"""
        for child in element.children:
            if isinstance(child, str):
                # NavigableString -- 纯文本
                text = child
                if text.strip():
                    self.handle_data(para, text)
            elif child.name == "br":
                # 换行符
                run = para.add_run()
                run.add_break()
            elif child.name == "p":
                # 单元格内的子段落
                if para.text.strip():
                    # 需要一个新段落
                    para = cell.add_paragraph()
                sub_align = self._parse_alignment(child.get("style", ""))
                if sub_align is not None:
                    para.alignment = sub_align
                self._process_children(para, child, cell)
            elif child.name in font_styles:
                # 行内格式标签
                style_attr = font_styles[child.name]
                text = child.get_text()
                if text.strip():
                    run = para.add_run(text)
                    self._apply_run_defaults(run)
                    setattr(run.font, style_attr, True)
                # 同时递归处理嵌套标签
                for nested in child.find_all(True):
                    if nested.name in font_styles and nested.string:
                        nr = para.add_run(nested.string)
                        self._apply_run_defaults(nr)
                        setattr(nr.font, font_styles[nested.name], True)
            elif child.name == "span":
                text = child.get_text()
                if text.strip():
                    run = para.add_run(text)
                    self._apply_run_defaults(run)
                    self._apply_span_style(run, child.get("style", ""))
            elif child.name in ("td", "th", "tr", "table"):
                # 跳过可能嵌套的结构性表格元素
                pass
            else:
                # 未知标签：直接提取文本
                text = child.get_text()
                if text.strip():
                    self.handle_data(para, text)

    def handle_data(self, para, text: str):
        """向 *para* 添加一个带默认字体设置的纯文本 run。"""
        text = text.replace("\n", " ")
        if not text.strip():
            return
        run = para.add_run(text)
        self._apply_run_defaults(run)

    # -----------------------------------------------------------------
    # 样式解析辅助
    # -----------------------------------------------------------------

    def _apply_run_defaults(self, run):
        """将默认字体、CJK 字体和字号应用到 run。"""
        font = run.font
        font.name = self.default_font_name
        font.size = Pt(self.default_font_size)
        # 东亚字体
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), self.default_east_asia)

    def _apply_span_style(self, run, style_str: str):
        """解析行内 ``style`` 属性并应用到 *run*。"""
        if not style_str:
            return

        # font-size
        m = re.search(r"font-size:\s*([\d.]+)\s*(?:px|pt)", style_str)
        if m:
            size = float(m.group(1))
            # 假定 px 值需要粗略转换
            if "px" in style_str[m.start(): m.end()]:
                size = size * 0.75  # px to pt approximation
            run.font.size = Pt(max(size, 5.0))

        # color
        m = re.search(r"(?:^|;)\s*color:\s*#([0-9a-fA-F]{6})", style_str)
        if m:
            hex_color = m.group(1)
            run.font.color.rgb = RGBColor(
                int(hex_color[0:2], 16),
                int(hex_color[2:4], 16),
                int(hex_color[4:6], 16),
            )

        # font-weight
        if "font-weight" in style_str:
            if "bold" in style_str or re.search(r"font-weight:\s*[6-9]\d{2}", style_str):
                run.font.bold = True

        # font-style
        if "font-style:" in style_str and "italic" in style_str:
            run.font.italic = True

    @staticmethod
    def _parse_alignment(style_str: str):
        """从 CSS 样式字符串中提取 text-align 值。"""
        if not style_str:
            return None
        m = re.search(r"text-align:\s*(\w+)", style_str)
        if m:
            align_val = m.group(1).lower()
            mapping = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            return mapping.get(align_val)
        return None

    @staticmethod
    def _parse_bg_color(style_str: str) -> Optional[str]:
        """从 CSS 样式字符串中提取 background-color 十六进制值。"""
        if not style_str:
            return None
        m = re.search(r"background(?:-color)?:\s*#([0-9a-fA-F]{6})", style_str)
        if m:
            return m.group(1)
        return None

    @staticmethod
    def _set_cell_shading(cell, color_hex: str):
        """将底纹（背景色）应用到单元格。"""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        shading = tcPr.find(qn("w:shd"))
        if shading is None:
            shading = OxmlElement("w:shd")
            tcPr.append(shading)
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), color_hex)
