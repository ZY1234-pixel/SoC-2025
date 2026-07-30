import pytest
from docx import Document

from test import _native_docx_table_count, _validate_content_integrity
from docflow.model.stages import (
    AnalysisPage,
    DocumentAnalysis,
    RecognitionEvidence,
    RecognitionItem,
    RecognitionPage,
    Rect,
    SemanticElement,
)


def test_content_integrity_rejects_unresolved_recognition_evidence() -> None:
    evidence = RecognitionEvidence(
        (RecognitionPage(0, 100, 100, (RecognitionItem("raw", "text", Rect(0, 0, 10, 10), 0),)),)
    )
    analysis = DocumentAnalysis(
        (AnalysisPage(0, 100, 100, (SemanticElement("body", "paragraph_group", Rect(0, 0, 10, 10), 0, ("other",)),)),)
    )

    with pytest.raises(RuntimeError, match="missing=\['raw'\].*duplicated=\['other'\]"):
        _validate_content_integrity(evidence, analysis)


def test_native_table_validation_counts_tables_nested_in_layout_cells(tmp_path) -> None:
    document = Document()
    layout = document.add_table(rows=1, cols=1)
    native = layout.cell(0, 0).add_table(rows=1, cols=1)
    native.style = "Table Grid"
    output = tmp_path / "nested.docx"
    document.save(output)

    assert _native_docx_table_count(output) == 1
