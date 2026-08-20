from __future__ import annotations

import base64
import io

import pytest
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from PIL import Image
from bs4 import BeautifulSoup

from docflow.model.stages import (
    AnalysisPage,
    DocumentAnalysis,
    FlowKind,
    FlowSection,
    GridCell,
    PlannedElement,
    Rect,
    SemanticElement,
    TextRow,
    TextSpan,
    TextStructure,
    TextParagraphLayout,
    TypographicRole,
)
from docflow.planning import ReflowPlanner
from docflow.planning.text_metrics import (
    estimate_text_units,
    estimate_wrapped_lines,
    fit_font_size_to_lines,
    infer_occupancy_line_height,
)
from docflow.renderer.reflow_docx_renderer import ReflowDocxRenderer
from docflow.renderer.docx_utils.html_table import get_table_column_weights


def _png_base64() -> str:
    image = Image.new("RGB", (120, 40), "white")
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    return base64.b64encode(buffer.getvalue()).decode("ascii")


def test_document_without_furniture_has_no_empty_header_footer_parts() -> None:
    analysis = DocumentAnalysis(
        (
            AnalysisPage(
                0,
                1000,
                1400,
                (
                    SemanticElement(
                        "body",
                        "paragraph_group",
                        Rect(100, 100, 900, 1200),
                        1,
                        ("raw",),
                        text="正文",
                        role_id="body",
                    ),
                ),
            ),
        ),
        (TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0),),
    )

    document = ReflowDocxRenderer().build(ReflowPlanner().plan(analysis))

    section_properties = document.sections[0]._sectPr
    assert not section_properties.xpath("./w:headerReference | ./w:footerReference")
    margins = section_properties.pgMar
    assert margins.get(qn("w:header")) == "0"
    assert margins.get(qn("w:footer")) == "0"


def test_table_column_weights_follow_cell_content() -> None:
    table = BeautifulSoup("<table><tr><td>近12个月最高/最低（元）</td><td>15.06/9.98</td></tr></table>", "html.parser").table

    left, right = get_table_column_weights(table)

    assert left > right * 2


def test_table_column_weights_prefer_source_geometry() -> None:
    table = BeautifulSoup("<table><tr><td>very long content</td><td>x</td></tr></table>", "html.parser").table

    assert get_table_column_weights(table, (0.25, 0.75)) == (0.25, 0.75)


def test_text_width_units_match_latin_and_cjk_font_metrics() -> None:
    assert estimate_text_units("AAAA") == pytest.approx(1.68)
    assert estimate_text_units("中文") == 2.0


def test_source_line_font_limit_rounds_down_to_word_half_points() -> None:
    assert fit_font_size_to_lines(10.0, ("中文" * 10,), (194.0,), 0.99) == 9.5


def test_text_occupancy_uses_rendered_lines_to_fill_tight_content_height() -> None:
    lines = estimate_wrapped_lines("正文" * 20, 10, 100, 6, 80, 1.0)

    assert lines == 5
    assert infer_occupancy_line_height(10, 10.5, 60, lines) == 12


def test_dense_grid_text_does_not_reinflate_line_spacing_past_source_rows() -> None:
    paragraph = Document().add_paragraph()
    element = PlannedElement(
        "body",
        "paragraph_group",
        "body",
        "正文" * 45,
        payload={
            "lines": ("正文" * 5,) * 9,
            "visual_line_count": 9,
            "line_height_pt": 13.25,
            "source_scale": 0.55,
            "grid_fit_scale": 0.9,
        },
        content_bbox=Rect(0, 0, 350, 216),
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 12.5, 1.0)

    ReflowDocxRenderer()._write_text(paragraph, element, {"body": role}, 0.817 * 0.9, 195)

    assert paragraph.paragraph_format.line_spacing.pt < 10


def test_tabular_text_rows_render_with_right_aligned_fields() -> None:
    cell = Document().add_table(rows=1, cols=1).cell(0, 0)
    cell.text = ""
    element = PlannedElement(
        "contact",
        "paragraph_group",
        "body",
        "phone email license number",
        payload={
            "alignment": "left",
            "line_height_pt": 12,
        },
        text_structure=TextStructure(preserve_source_lines=True, tabular_rows=True),
        text_rows=(
            TextRow(
                "phone email",
                Rect(0, 0, 200, 20),
                (TextSpan("phone", Rect(0, 0, 50, 20)), TextSpan("email", Rect(150, 0, 200, 20))),
            ),
            TextRow(
                "license number",
                Rect(0, 20, 200, 40),
                (TextSpan("license", Rect(0, 20, 60, 40)), TextSpan("number", Rect(150, 20, 200, 40))),
            ),
        ),
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    renderer = ReflowDocxRenderer()
    renderer._clear_container(cell)
    renderer._render_element(cell, element, {"body": role}, 1.0, 200)

    assert [paragraph.text for paragraph in cell.paragraphs] == ["phone\temail", "license\tnumber"]
    assert all(paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT for paragraph in cell.paragraphs)
    assert all(paragraph._p.xpath("./w:pPr/w:tabs/w:tab[@w:val='right']") for paragraph in cell.paragraphs)


def test_two_line_text_renders_each_source_alignment() -> None:
    cell = Document().add_table(rows=1, cols=1).cell(0, 0)
    cell.text = ""
    element = PlannedElement(
        "caption",
        "paragraph_group",
        "body",
        "centered caption right credit",
        payload={
            "alignment": "left",
            "line_height_pt": 12,
            "source_scale": 1.0,
        },
        content_bbox=Rect(0, 0, 200, 60),
        text_rows=(
            TextRow("centered caption", Rect(0, 0, 200, 20)),
            TextRow("right credit", Rect(100, 20, 200, 40)),
        ),
        row_alignments=("center", "right"),
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    renderer = ReflowDocxRenderer()
    renderer._clear_container(cell)
    renderer._render_element(cell, element, {"body": role}, 1.0, 200)

    assert [paragraph.text for paragraph in cell.paragraphs] == ["centered caption", "right credit"]
    assert [paragraph.alignment for paragraph in cell.paragraphs] == [
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.RIGHT,
    ]
    assert all(paragraph.paragraph_format.line_spacing.pt == 12 for paragraph in cell.paragraphs)


def test_renderer_executes_resolved_text_layout_without_reinferring_payload() -> None:
    cell = Document().add_table(rows=1, cols=1).cell(0, 0)
    cell.text = ""
    element = PlannedElement(
        "body",
        "paragraph_group",
        "body",
        "planned text",
        payload={"lines": ("conflicting payload",), "line_height_pt": 40, "alignment": "left"},
        text_layout=(
            TextParagraphLayout(
                "planned text",
                "right",
                9.0,
                11.0,
                left_indent_pt=4.0,
            ),
        ),
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 20, 1.0)

    renderer = ReflowDocxRenderer()
    renderer._clear_container(cell)
    renderer._render_element(cell, element, {"body": role}, 0.5, 200)
    paragraph = cell.paragraphs[0]

    assert paragraph.text == "planned text"
    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert paragraph.runs[0].font.size.pt == 9
    assert paragraph.paragraph_format.line_spacing.pt == 11


def test_single_line_heading_fits_with_uniform_font_scaling() -> None:
    role = TypographicRole("heading", "黑体", "Times New Roman", 16, 1.0)
    element = SemanticElement(
        "heading",
        "heading",
        Rect(100, 100, 350, 150),
        1,
        ("source",),
        "投资评级：买入（维持）",
        "heading",
        payload={"lines": ("投资评级：买入（维持）",), "background_color": "#4671A6"},
    )
    plan = ReflowPlanner().plan(DocumentAnalysis((AnalysisPage(0, 1000, 1400, (element,)),), (role,)))

    document = ReflowDocxRenderer().build(plan)
    paragraph = document.paragraphs[0]
    run = paragraph.runs[0]

    assert run.font.size.pt <= round(role.font_size_pt * plan.pages[0].fit_scale * 2) / 2
    assert run._element.rPr.find(qn("w:w")) is None
    assert paragraph._p.pPr.find(qn("w:shd")) is not None


def test_sparse_grid_keeps_valid_empty_cells() -> None:
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)
    elements = (
        SemanticElement("top-left", "paragraph_group", Rect(0, 0, 400, 200), 1, ("a",), "A", "body"),
        SemanticElement("bottom-right", "paragraph_group", Rect(600, 300, 1000, 500), 2, ("b",), "B", "body"),
        SemanticElement("top-right", "paragraph_group", Rect(600, 0, 1000, 200), 3, ("c",), "C", "body"),
        SemanticElement("bottom-left", "paragraph_group", Rect(0, 600, 400, 800), 4, ("d",), "D", "body"),
    )
    plan = ReflowPlanner().plan(DocumentAnalysis((AnalysisPage(0, 1000, 1400, elements),), (role,)))

    document = ReflowDocxRenderer().build(plan)
    grid = document.tables[0]

    assert all(cell.paragraphs or cell.tables for row in grid.rows for cell in row.cells)


def test_grid_renderer_hides_cell_marks_without_splitting_row_spans() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    flow = FlowSection(
        "grid",
        FlowKind.GRID,
        ("span", "side"),
        (100, 100, 100),
        gutter_pt=10,
        grid_cells=(
            GridCell(0, 2, ("side",), row_span=2),
            GridCell(1, 0, ("span",), column_span=2),
        ),
        row_heights_pt=(30, 40),
    )
    elements = {
        "span": PlannedElement("span", "paragraph_group", "body", "image area"),
        "side": PlannedElement(
            "side",
            "paragraph_group",
            "body",
            "side flow",
            payload={"grid_fit_scale": 0.8},
        ),
    }
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._render_grid(container, flow, elements, {"body": role}, 1.0)

    assert container.tables[0]._tbl.xpath('.//w:gridSpan[@w:val="2"]')
    assert container.tables[0]._tbl.xpath(".//w:vMerge")
    cell_properties = container.tables[0]._tbl.xpath(".//w:tcPr")
    assert cell_properties
    assert all(properties.find(qn("w:hideMark")) is not None for properties in cell_properties)
    assert all(row.height is None for row in container.tables[0].rows)
    assert container.tables[0].cell(0, 2).paragraphs[0].runs[0].font.size.pt == 8.5
    assert container.tables[0].cell(1, 0)._tc.xpath("./w:tcPr/w:tcMar/w:end/@w:w") == ["100"]


def test_row_spanning_grid_figure_floats_without_expanding_its_start_row() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    flow = FlowSection(
        "grid",
        FlowKind.GRID,
        ("figure", "top", "bottom"),
        (120, 120),
        grid_cells=(
            GridCell(0, 0, ("figure",), row_span=2),
            GridCell(0, 1, ("top",)),
            GridCell(1, 1, ("bottom",)),
        ),
        row_heights_pt=(40, 40),
    )
    elements = {
        "figure": PlannedElement(
            "figure",
            "figure_group",
            "body",
            payload={
                "image_base64": _png_base64(),
                "source_bbox": (0, 0, 120, 80),
                "source_scale": 1.0,
                "width_fraction": 1.0,
                "alignment": "left",
            },
        ),
        "top": PlannedElement("top", "paragraph_group", "body", "top"),
        "bottom": PlannedElement("bottom", "paragraph_group", "body", "bottom"),
    }
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._render_grid(container, flow, elements, {"body": role}, 1.0)

    table = container.tables[0]
    assert table._tbl.xpath(".//wp:anchor")
    assert not table._tbl.xpath(".//wp:inline")
    assert not table._tbl.xpath(".//w:vMerge")
    assert table._tbl.xpath(".//wp:positionH[@relativeFrom='column']/wp:align[text()='left']")
    assert all(row.height_rule == WD_ROW_HEIGHT_RULE.EXACTLY for row in table.rows)
    assert [row.height.pt for row in table.rows] == pytest.approx((40, 40))


def test_mixed_figure_text_cell_uses_normal_merged_flow() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    flow = FlowSection(
        "grid",
        FlowKind.GRID,
        ("figure", "body"),
        (120,),
        grid_cells=(GridCell(0, 0, ("figure", "body"), row_span=2),),
        row_heights_pt=(40, 40),
    )
    elements = {
        "figure": PlannedElement(
            "figure",
            "figure_group",
            "body",
            payload={
                "image_base64": _png_base64(),
                "source_bbox": (0, 0, 120, 80),
                "source_scale": 1.0,
                "width_fraction": 1.0,
            },
        ),
        "body": PlannedElement("body", "paragraph_group", "body", "continued text"),
    }
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._render_grid(container, flow, elements, {"body": role}, 1.0)

    table = container.tables[0]
    assert table._tbl.xpath(".//w:vMerge")
    assert not table._tbl.xpath(".//wp:anchor")
    assert all(row.height is None for row in table.rows)
    assert "continued text" in table.cell(0, 0).text


def test_header_with_image_payload_renders_as_image() -> None:
    document = Document()
    header = document.sections[0].header
    element = PlannedElement(
        "logo",
        "header",
        "body",
        text="OCR fallback",
        payload={
            "image_base64": _png_base64(),
            "source_bbox": (100, 20, 300, 80),
            "source_scale": 0.5,
        },
    )

    ReflowDocxRenderer()._render_furniture(header, ("logo",), {"logo": element}, {}, 1.0, 500)

    assert header._element.xpath(".//w:drawing")
    assert header._element.xpath('.//wp:anchor/wp:positionH[@relativeFrom="page"]')
    assert header._element.xpath('.//wp:anchor/wp:positionV[@relativeFrom="page"]')
    assert header._element.xpath("string(.//wp:positionH/wp:posOffset)") == str(50 * 12700)
    assert header._element.xpath("string(.//wp:positionV/wp:posOffset)") == str(10 * 12700)
    assert "OCR fallback" not in "\n".join(paragraph.text for paragraph in header.paragraphs)


def test_multiple_header_images_preserve_independent_page_positions() -> None:
    document = Document()
    header = document.sections[0].header
    elements = {
        identifier: PlannedElement(
            identifier,
            "header",
            "body",
            payload={
                "image_base64": _png_base64(),
                "source_bbox": bbox,
                "source_scale": 0.5,
            },
        )
        for identifier, bbox in (("top", (20, 30, 120, 60)), ("bottom", (200, 100, 300, 140)))
    }

    ReflowDocxRenderer()._render_furniture(header, tuple(elements), elements, {}, 1.0, 500)

    anchors = header._element.xpath(".//wp:anchor")
    assert len(anchors) == 2
    assert [anchor.xpath("string(./wp:positionH/wp:posOffset)") for anchor in anchors] == [
        str(10 * 12700),
        str(100 * 12700),
    ]
    assert [anchor.xpath("string(./wp:positionV/wp:posOffset)") for anchor in anchors] == [
        str(15 * 12700),
        str(50 * 12700),
    ]


def test_wrapped_flow_uses_floating_editable_media_table() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    flow = FlowSection(
        "wrapped",
        FlowKind.WRAPPED,
        ("body", "image"),
        floating_element_id="image",
        floating_width_pt=40,
    )
    elements = {
        "body": PlannedElement("body", "paragraph_group", "body", "Editable body"),
        "image": PlannedElement(
            "image",
            "figure_group",
            "body",
            payload={"image_base64": _png_base64(), "caption": "Editable caption"},
        ),
    }
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._render_wrapped(container, flow, elements, {"body": role}, 0.5, 100)

    positioning = container.tables[0]._tbl.tblPr.find(qn("w:tblpPr"))
    assert positioning.get(qn("w:tblpXSpec")) == "right"
    assert container.tables[0].cell(0, 0).paragraphs[0].runs[0]._r.xpath(".//wp:extent")[0].get("cx") == str(
        int(20 * 12700)
    )
    assert container.tables[0].cell(0, 0).paragraphs[-1].text == "Editable caption"
    assert container.paragraphs[-1].text == "Editable body"


def test_grid_image_does_not_exceed_its_source_physical_width() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    element = PlannedElement(
        "image",
        "figure_group",
        "body",
        payload={
            "image_base64": _png_base64(),
            "source_bbox": (0, 0, 50, 20),
            "source_scale": 0.5,
            "width_fraction": 1.0,
        },
    )

    ReflowDocxRenderer()._write_image(container, element, 0.5, 100)

    assert container.paragraphs[-1].runs[0]._r.xpath(".//wp:extent")[0].get("cx") == str(int(25 * 12700))


def test_left_wrapped_flow_uses_absolute_position_and_editable_text() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    flow = FlowSection(
        "wrapped",
        FlowKind.WRAPPED,
        ("body", "note"),
        floating_element_id="note",
        floating_width_pt=40,
        floating_side="left",
        floating_offset_x_pt=20,
    )
    elements = {
        "body": PlannedElement("body", "paragraph_group", "body", "Editable body"),
        "note": PlannedElement("note", "paragraph_group", "body", "Editable note"),
    }
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._render_wrapped(container, flow, elements, {"body": role}, 1.0, 100)

    positioning = container.tables[0]._tbl.tblPr.find(qn("w:tblpPr"))
    assert positioning.get(qn("w:tblpX")) == "400"
    assert positioning.get(qn("w:tblpXSpec")) is None
    assert container.tables[0].cell(0, 0).paragraphs[-1].text == "Editable note"


def test_native_table_respects_element_width_and_sets_fixed_table_width() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    element = PlannedElement(
        "table",
        "table_group",
        "body",
        payload={"html": "<table><tr><td>A</td><td>B</td></tr></table>", "width_fraction": 0.75},
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._write_native_table(container, element, {"body": role}, 0.8, 100)

    grid = container.tables[0]._tbl.tblGrid
    assert sum(int(column.get(qn("w:w"))) for column in grid.gridCol_lst) == pytest.approx(75 * 20, abs=2)
    table_width = container.tables[0]._tbl.tblPr.find(qn("w:tblW"))
    assert table_width.get(qn("w:type")) == "dxa"
    assert int(table_width.get(qn("w:w"))) == pytest.approx(75 * 20, abs=2)
    assert container.tables[0]._tbl.tblPr.find(qn("w:tblCaption")).get(qn("w:val")) == "docflow-native-table"


def test_native_table_respects_source_left_anchor() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    element = PlannedElement(
        "table",
        "table_group",
        "body",
        payload={
            "html": "<table><tr><td>A</td></tr></table>",
            "width_fraction": 0.5,
            "left_indent_pt": 24,
        },
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._write_native_table(container, element, {"body": role}, 1.0, 100)

    table = container.tables[0]
    assert table.alignment == WD_TABLE_ALIGNMENT.LEFT
    assert table._tbl.tblPr.find(qn("w:tblInd")).get(qn("w:w")) == "480"


def test_native_table_preserves_sparse_source_columns_and_header_span() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    element = PlannedElement(
        "table",
        "table_group",
        "body",
        payload={
            "html": (
                "<table><tr><td>Letter</td><td colspan='3'>Use</td></tr>"
                "<tr><td>K</td><td>kite</td><td></td><td></td></tr>"
                "<tr><td>L</td><td></td><td></td><td>lion</td></tr></table>"
            )
        },
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._write_native_table(container, element, {"body": role}, 1.0, 100)

    table = container.tables[0]
    assert len(table.columns) == 4
    assert table.cell(1, 1).text == "kite"
    assert table.cell(2, 3).text == "lion"
    assert table._tbl.xpath('.//w:gridSpan[@w:val="3"]')


def test_native_table_renders_data_uri_cell_image() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    data_uri = f"data:image/png;base64,{_png_base64()}"
    element = PlannedElement(
        "table",
        "table_group",
        "body",
        payload={"html": f"<table><tr><td><img src='{data_uri}'>型号 A</td></tr></table>"},
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._write_native_table(container, element, {"body": role}, 1.0, 100)

    table = container.tables[0]
    assert table._tbl.xpath(".//a:blip")
    assert table.rows[0].height_rule == WD_ROW_HEIGHT_RULE.AT_LEAST
    assert "型号 A" in table.cell(0, 0).text


def test_native_table_does_not_append_break_after_image_only_cell() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    data_uri = f"data:image/png;base64,{_png_base64()}"
    element = PlannedElement(
        "table",
        "table_group",
        "body",
        payload={"html": f"<table><tr><td><img src='{data_uri}'></td></tr></table>"},
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._write_native_table(container, element, {"body": role}, 1.0, 100)

    assert not container.tables[0].cell(0, 0)._tc.xpath(".//w:br")


def test_native_table_applies_resolved_cell_semantics() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    element = PlannedElement(
        "table",
        "table_group",
        "body",
        payload={
            "html": "<table><tr><th>Heading</th><td>Value</td></tr></table>",
            "table_cell_styles": ((0, 0, "center", True), (0, 1, "right", False)),
        },
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._write_native_table(container, element, {"body": role}, 1.0, 100)

    table = container.tables[0]
    assert table.cell(0, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert table.cell(0, 0).paragraphs[0].runs[0].bold is True
    assert table.cell(0, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert all(cell.vertical_alignment == 1 for cell in table.rows[0].cells)


def test_native_table_preserves_cell_lines_and_semantic_header_styles() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    element = PlannedElement(
        "table",
        "table_group",
        "body",
        payload={
            "html": "<table><tr><th>Column</th><td class='row-header'>Line one\nLine two</td></tr></table>",
        },
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._write_native_table(container, element, {"body": role}, 1.0, 100)

    table = container.tables[0]
    assert table.cell(0, 1).text == "Line one\nLine two"
    assert table.cell(0, 0)._tc.xpath('./w:tcPr/w:shd[@w:fill="F2F5F8"]')
    assert table.cell(0, 1)._tc.xpath('./w:tcPr/w:shd[@w:fill="EEF7FF"]')
    assert table.cell(0, 1).paragraphs[0].runs[-1].bold is True


def test_horizontal_table_uses_th_rows_without_thead_as_header() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    element = PlannedElement(
        "table",
        "table_group",
        "body",
        payload={
            "html": "<table><tr><th>A</th><th>B</th></tr><tr><td>1</td><td>2</td></tr></table>",
            "table_rule_style": "horizontal",
        },
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._write_native_table(container, element, {"body": role}, 1.0, 100)

    table = container.tables[0]
    assert all(cell._tc.xpath("./w:tcPr/w:tcBorders/w:bottom") for cell in table.rows[0].cells)
    assert not any(cell._tc.xpath("./w:tcPr/w:tcBorders/w:bottom") for cell in table.rows[1].cells)


def test_numbered_equation_clears_word_default_cell_paragraph_spacing() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    element = PlannedElement(
        "equation",
        "equation_group",
        "body",
        payload={"image_base64": _png_base64(), "number": "(1)", "width_fraction": 0.8},
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._write_equation(container, element, {"body": role}, 1.0, 100)

    for cell in container.tables[0].rows[0].cells:
        assert cell.paragraphs[0].paragraph_format.space_before.pt == 0
        assert cell.paragraphs[0].paragraph_format.space_after.pt == 0


def test_native_table_uses_inferred_table_font() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    element = PlannedElement(
        "table",
        "table_group",
        "body",
        payload={"html": "<table><tr><td>内容</td></tr></table>", "font_family": "仿宋"},
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._write_native_table(container, element, {"body": role}, 1.0, 100)

    fonts = container.tables[0].cell(0, 0).paragraphs[0].runs[0]._r.rPr.find(qn("w:rFonts"))
    assert fonts.get(qn("w:eastAsia")) == "仿宋"


def test_native_table_uses_source_height_as_row_minimum() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    element = PlannedElement(
        "table",
        "table_group",
        "body",
        payload={"html": "<table><tr><td>A</td></tr><tr><td>B</td></tr></table>", "table_height_pt": 100},
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._write_native_table(container, element, {"body": role}, 0.8, 100)

    assert all(row.height.pt == pytest.approx(40) for row in container.tables[0].rows)
    assert all(row.height_rule == WD_ROW_HEIGHT_RULE.EXACTLY for row in container.tables[0].rows)


def test_native_table_uses_fused_row_height_ratios() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    element = PlannedElement(
        "table",
        "table_group",
        "body",
        payload={
            "html": "<table><tr><td>A</td><td>B</td></tr><tr><td>C</td><td>D</td></tr></table>",
            "table_height_pt": 100,
            "table_content_fit": True,
            "table_row_height_ratios": (0.75, 0.25),
            "table_column_width_ratios": (0.7, 0.3),
        },
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._write_native_table(container, element, {"body": role}, 1.0, 100)

    assert [row.height.pt for row in container.tables[0].rows] == pytest.approx([75, 25])
    grid = container.tables[0]._tbl.tblGrid.gridCol_lst
    assert [int(column.get(qn("w:w"))) for column in grid] == pytest.approx([1400, 600], abs=1)


def test_native_table_fixed_row_height_accounts_for_wrapped_content() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    element = PlannedElement(
        "table",
        "table_group",
        "body",
        payload={
            "html": "<table><tr><td>content that cannot fit on one line</td></tr></table>",
            "table_height_pt": 10,
            "table_min_font_size_pt": 6.5,
        },
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._write_native_table(container, element, {"body": role}, 1.0, 40)

    row = container.tables[0].rows[0]
    assert row.height_rule == WD_ROW_HEIGHT_RULE.EXACTLY
    assert row.height.pt > 20


def test_heading_preserves_visual_rows_but_joins_overlapping_lines() -> None:
    element = PlannedElement(
        "heading",
        "heading",
        "heading",
        text="2 Chapter title Continued",
        payload={
            "lines": ("2", "Chapter title", "Continued"),
            "line_tops_px": (10, 12, 40),
            "line_heights_px": (30, 24, 20),
        },
    )

    assert ReflowDocxRenderer._visual_text(element) == "2 Chapter title\nContinued"


def test_question_block_preserves_option_rows() -> None:
    element = PlannedElement(
        "question",
        "paragraph_group",
        "body",
        text="1 Question A. First B. Second C. Third",
        payload={"lines": ("1 Question", "A. First", "B. Second", "C. Third")},
        text_structure=TextStructure(preserve_source_lines=True, is_list=True),
    )

    assert ReflowDocxRenderer._visual_text(element) == "1 Question\nA. First\nB. Second\nC. Third"

    paragraph = Document().add_paragraph()
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)
    ReflowDocxRenderer()._write_text(paragraph, element, {"body": role}, 1.0, 100)
    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT


def test_vertical_text_remains_editable_in_a_vertical_cell() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    element = PlannedElement(
        "aside",
        "paragraph_group",
        "body",
        text="公司证券研究报告",
        payload={
            "lines": ("公司证券研究报告",),
            "source_bbox": (20, 100, 60, 700),
            "source_scale": 0.5,
            "background_color": "#034579",
        },
        text_structure=TextStructure(orientation="vertical"),
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._render_element(container, element, {"body": role}, 1.0, 100)

    assert container.tables[0].cell(0, 0).text == "公\n司\n证\n券\n研\n究\n报\n告"
    assert not container.tables[0]._tbl.xpath('.//w:textDirection')
    assert container._tc.xpath('./w:tcPr/w:shd[@w:fill="034579"]')
    assert container.tables[0].cell(0, 0).paragraphs[0].runs[0].font.size.pt == role.font_size_pt


def test_multiline_heading_font_fits_each_preserved_source_line() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    element = PlannedElement(
        "heading",
        "heading",
        "heading",
        text="A deliberately long heading Continued",
        payload={
            "lines": ("A deliberately long heading", "Continued"),
            "line_tops_px": (0, 20),
            "line_heights_px": (18, 18),
            "alignment": "left",
        },
    )
    role = TypographicRole("heading", "宋体", "Times New Roman", 20, 1.0)

    ReflowDocxRenderer()._render_element(container, element, {"heading": role}, 1.0, 100)

    run = container.paragraphs[-1].runs[0]
    assert run._element.rPr.find(qn("w:w")) is None
    assert run.font.size.pt <= 100 * 0.98 / estimate_text_units("A deliberately long heading") + 0.1


def test_centered_source_rows_fit_without_word_adding_wraps() -> None:
    paragraph = Document().add_paragraph()
    element = PlannedElement(
        "affiliations",
        "paragraph_group",
        "body",
        text="First unit Second deliberately longer unit",
        payload={"lines": ("First unit", "Second deliberately longer unit"), "alignment": "center"},
        text_structure=TextStructure(preserve_source_lines=True),
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 20, 1.0)

    ReflowDocxRenderer()._write_text(paragraph, element, {"body": role}, 1.0, 100)

    assert paragraph.runs[0].font.size.pt < role.font_size_pt


def test_continuous_paragraph_keeps_natural_character_width_without_hard_breaks() -> None:
    paragraph = Document().add_paragraph()
    line = "中文" * 10
    element = PlannedElement(
        "body",
        "paragraph_group",
        "body",
        text=line * 2,
        payload={"lines": (line, line), "alignment": "justify"},
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10, 1.0)

    ReflowDocxRenderer()._write_text(paragraph, element, {"body": role}, 1.0, 194)

    assert paragraph.runs[0].font.size.pt == role.font_size_pt
    assert paragraph.runs[0]._element.rPr.find(qn("w:w")) is None
    assert "\n" not in paragraph.text


def test_column_layout_collapses_word_trailing_paragraph() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)
    elements = {
        "left": PlannedElement("left", "paragraph_group", "body", text="left", payload={"column": 0}),
        "right": PlannedElement("right", "paragraph_group", "body", text="right", payload={"column": 1}),
    }
    flow = FlowSection("columns", FlowKind.SEQUENTIAL_COLUMNS, tuple(elements), column_widths_pt=(50, 50))

    ReflowDocxRenderer()._render_columns(container, flow, elements, {"body": role}, 1.0)

    spacing = container.paragraphs[-1]._p.pPr.find(qn("w:spacing"))
    assert spacing.get(qn("w:line")) == "1"


def test_reflow_docx_keeps_text_tables_and_equation_numbers_editable(tmp_path) -> None:
    roles = (
        TypographicRole("heading", "黑体", "Times New Roman", 18, 1.0, bold=True),
        TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0),
    )
    elements = (
        SemanticElement("heading", "heading", Rect(100, 100, 900, 160), 1, ("r1",), text="Editable heading", role_id="heading"),
        SemanticElement(
            "table",
            "table_group",
            Rect(100, 220, 900, 500),
            2,
            ("r2",),
            payload={"html": "<table><tr><th>Name</th><th>Value</th></tr><tr><td>A</td><td>7</td></tr></table>"},
        ),
        SemanticElement(
            "equation",
            "equation_group",
            Rect(200, 560, 800, 640),
            3,
            ("r3",),
            payload={"image_base64": _png_base64(), "number": "(7)"},
        ),
    )
    analysis = DocumentAnalysis((AnalysisPage(0, 1000, 1400, elements),), roles)
    plan = ReflowPlanner().plan(analysis)
    output = tmp_path / "result.docx"

    ReflowDocxRenderer().render(plan, str(output))
    document = Document(output)
    text = " ".join(node.text or "" for node in document.element.body.iter(qn("w:t")))

    assert "Editable heading" in text
    assert "Name" in text and "Value" in text
    assert "(7)" in text
    assert len(document.element.body.xpath(".//w:tbl")) >= 2
    native_table = document.element.body.xpath('.//w:tbl[w:tblPr/w:tblStyle[@w:val="TableGrid"]]')[0]
    assert native_table.xpath('.//w:pPr/w:spacing[@w:lineRule="exact"]')


def test_reflow_docx_creates_one_unlinked_section_per_source_page(tmp_path) -> None:
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)
    pages = tuple(
        AnalysisPage(
            index,
            1000,
            1400,
            (
                SemanticElement(f"header-{index}", "header", Rect(100, 20, 900, 60), 1, (f"h{index}",), text=f"H{index}", role_id="body"),
                SemanticElement(f"body-{index}", "paragraph_group", Rect(100, 100, 900, 1200), 2, (f"b{index}",), text=f"B{index}", role_id="body"),
            ),
        )
        for index in range(2)
    )
    output = tmp_path / "two-pages.docx"

    ReflowDocxRenderer().render(ReflowPlanner().plan(DocumentAnalysis(pages, (role,))), str(output))
    document = Document(output)

    assert len(document.sections) == 2
    first_header = " ".join(paragraph.text for paragraph in document.sections[0].header.paragraphs)
    second_header = " ".join(paragraph.text for paragraph in document.sections[1].header.paragraphs)
    assert "H0" in first_header
    assert "H1" in second_header
    assert not document.sections[1].header.is_linked_to_previous


def test_reflow_docx_uses_body_flow_with_a_collapsed_end_mark(tmp_path) -> None:
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)
    element = SemanticElement("body", "paragraph_group", Rect(100, 100, 900, 1200), 1, ("r1",), text="Body", role_id="body")
    plan = ReflowPlanner().plan(DocumentAnalysis((AnalysisPage(0, 1000, 1400, (element,)),), (role,)))
    output = tmp_path / "page-frame.docx"

    ReflowDocxRenderer().render(plan, str(output))
    document = Document(output)
    assert not document.tables
    assert document.paragraphs[0].text == "Body"
    assert document.paragraphs[-1]._p.xpath('./w:pPr/w:spacing[@w:line="1"]')
    assert document.paragraphs[-1]._p.xpath("./w:pPr/w:rPr/w:vanish")
    assert not document.element.body.xpath(".//w:sectPr/w:docGrid")


def test_scaled_page_adds_only_a_collapsed_end_mark() -> None:
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)
    element = SemanticElement(
        "body",
        "paragraph_group",
        Rect(100, 100, 900, 1300),
        1,
        ("r1",),
        text="x" * 10000,
        role_id="body",
    )
    plan = ReflowPlanner().plan(
        DocumentAnalysis(
            (AnalysisPage(0, 1000, 1400, (element,)),),
            (role,),
        )
    )
    page = plan.pages[0]

    document = ReflowDocxRenderer().build(plan)

    assert page.fit_scale < 1.0
    assert document.paragraphs[-1]._p.xpath('./w:pPr/w:spacing[@w:line="1"]')
    assert document.paragraphs[-1]._p.xpath("./w:pPr/w:rPr/w:vanish")


def test_layout_table_gutter_preserves_planned_content_width() -> None:
    table = Document().add_table(rows=1, cols=3)

    ReflowDocxRenderer._format_layout_table(table, (100, 100, 100), 20)

    assert [cell.width.pt for cell in table.rows[0].cells] == [110, 120, 110]
    assert not table._tbl.xpath(".//w:cantSplit")


def test_nested_table_trailing_paragraph_is_collapsed() -> None:
    cell = Document().add_table(rows=1, cols=1).cell(0, 0)
    cell.add_table(rows=1, cols=1)

    ReflowDocxRenderer._collapse_trailing_paragraph(cell)

    assert cell.paragraphs[-1]._p.xpath('./w:pPr/w:spacing[@w:line="1"]')
    assert cell.paragraphs[-1]._p.xpath("./w:pPr/w:rPr/w:vanish")


def test_reflow_docx_writes_planned_vertical_spacing(tmp_path) -> None:
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)
    elements = (
        SemanticElement("first", "heading", Rect(100, 100, 900, 200), 1, ("r1",), text="First", role_id="body"),
        SemanticElement("second", "paragraph_group", Rect(100, 300, 900, 400), 2, ("r2",), text="Second", role_id="body"),
    )
    plan = ReflowPlanner().plan(DocumentAnalysis((AnalysisPage(0, 1000, 1400, elements),), (role,)))
    output = tmp_path / "spacing.docx"

    ReflowDocxRenderer().render(plan, str(output))
    paragraph = next(item for item in Document(output).paragraphs if item.text == "Second")

    assert plan.pages[0].elements[1].payload["space_before_pt"] > 0
    assert paragraph.paragraph_format.space_before.pt == pytest.approx(
        plan.pages[0].elements[1].payload["space_before_pt"] * plan.pages[0].fit_scale,
        abs=0.1,
    )


def test_background_paragraph_uses_source_width() -> None:
    paragraph = Document().add_paragraph()
    element = PlannedElement(
        "label",
        "paragraph_group",
        "body",
        "Label",
        payload={"background_color": "#892549", "width_fraction": 0.25},
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._write_text(paragraph, element, {"body": role}, 0.8, 100)

    assert paragraph.paragraph_format.right_indent.pt == pytest.approx(75)


def test_reflow_docx_writes_source_bbox_paragraph_indents(tmp_path) -> None:
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)
    element = SemanticElement(
        "author",
        "paragraph_group",
        Rect(300, 100, 700, 150),
        1,
        ("r1",),
        text="Author Name",
        role_id="body",
        payload={"lines": ("Author Name",)},
    )
    body = SemanticElement("body", "paragraph_group", Rect(100, 200, 900, 300), 2, ("r2",), text="Body", role_id="body")
    plan = ReflowPlanner().plan(DocumentAnalysis((AnalysisPage(0, 1000, 1400, (element, body)),), (role,)))
    output = tmp_path / "indents.docx"

    ReflowDocxRenderer().render(plan, str(output))
    paragraph = next(item for item in Document(output).paragraphs if item.text == "Author Name")

    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert paragraph.paragraph_format.left_indent.pt == pytest.approx(paragraph.paragraph_format.right_indent.pt, abs=0.1)


def test_reflow_docx_uses_deterministic_line_height(tmp_path) -> None:
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)
    element = SemanticElement(
        "body",
        "paragraph_group",
        Rect(100, 100, 900, 200),
        1,
        ("r1",),
        text="First line Second line",
        role_id="body",
        payload={"lines": ("First line", "Second line"), "line_heights_px": (20, 20)},
    )
    plan = ReflowPlanner().plan(DocumentAnalysis((AnalysisPage(0, 1000, 1400, (element,)),), (role,)))
    output = tmp_path / "source-lines.docx"

    ReflowDocxRenderer().render(plan, str(output))
    paragraph = next(item for item in Document(output).paragraphs if "First line" in item.text)

    assert plan.pages[0].elements[0].payload["line_height_pt"] == pytest.approx(20 * 841.89 / 1400)
    assert paragraph.text == "First line Second line"
    assert paragraph.paragraph_format.line_spacing_rule == WD_LINE_SPACING.EXACTLY
    assert paragraph.paragraph_format.widow_control is False
    assert paragraph.paragraph_format.keep_together is False
    assert paragraph.paragraph_format.keep_with_next is False
