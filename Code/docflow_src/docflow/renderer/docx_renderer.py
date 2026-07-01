"""DOCX 文档渲染器。

实现分区域渲染编排、逐区块类型渲染、跨列布局表格等核心功能。
"""
from __future__ import annotations

import re
import io
import os
import logging
import math
from pathlib import Path
from collections import defaultdict
from copy import copy
from typing import TYPE_CHECKING, List, Optional, Tuple

from docx import Document as DocxDocument
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docflow.renderer.base import BaseRenderer
from docflow.renderer.docx_utils.paragraph_fmt import (
    reset_paragraph_format, set_paragraph_spacing, add_spacing_para,
)
from docflow.renderer.docx_utils.section_fmt import set_section_columns as _set_section_columns_fmt
from docflow.renderer.docx_utils.run_fmt import set_run_font
from docflow.renderer.docx_utils.table_fmt import (
    clear_table_borders, set_table_col_widths, set_cell_right_margin,
    fit_table_to_width,
)
from docflow.renderer.docx_utils.html_table import HtmlToDocx
from docflow.renderer.context import RenderContext
from docflow.schema.style import resolve_textline_style
from docflow.model.base import Block, BlockType, Alignment
from docflow.model.blocks.text_block import TextBlock, should_insert_space
from docflow.model.blocks.table_block import TableBlock
from docflow.model.blocks.image_block import ImageBlock
from docflow.model.blocks.equation_block import EquationBlock
from docflow.model.zone import Zone
from docflow.utils.constants import MIN_LINE_SPACING_PT

if TYPE_CHECKING:
    from docflow.model.page import Page, Document

logger = logging.getLogger(__name__)

# 对齐字符串 → python-docx 枚举（模块级常量，避免重复构造）
_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

# caption 类区块集合
_CAPTION_TYPES = frozenset({
    BlockType.TABLE_CAPTION,
    BlockType.FIGURE_CAPTION,
    BlockType.TABLE_FOOTNOTE,
    BlockType.FORMULA_CAPTION,
})
_LISTISH_RE = re.compile(
    r"^\s*(\d+[\.．、\)](?!\d)|[A-Za-z][\.\)]|[（(]?[一二三四五六七八九十百]+[)）\.、]?|[•\-·])\s*"
)
_QUESTION_OR_OPTION_RE = re.compile(
    r"^\s*(?:\d{1,2}\s*[A-Z][a-z]|\d{1,2}[\.\)．、]\s*\S|[A-H][\.\)]\s*\S|[•●]\s*\S)"
)
_NUMBERED_TITLE_RE = re.compile(
    r"^\s*(\d+(?:\.\d+)*[\.、]|\d+[)）]|\(?\d+\)|[（(]?[一二三四五六七八九十百]+[)）\.、])\s*\S+"
)
_NUMBERED_TITLE_LEVEL_RE = re.compile(
    r"^\s*(?:(\d+(?:\.\d+)*)(?:[\.、])?|[（(]?([一二三四五六七八九十百]+)[)）\.、])\s*\S"
)
_FORMULA_NUMBER_TEXT_RE = re.compile(r"^\s*\(?\s*\d{1,3}[a-zA-Z]?\s*\)?\s*$")
_FIELD_LIKE_LINE_RE = re.compile(
    r"^\s*(?:"
    r"(?:this\s+article\s+was\s+downloaded\s+by|on|publisher|registered\s+(?:number|office)|"
    r"received|accepted|published\s+online|doi|to\s+cite\s+this\s+article|to\s+link\s+to\s+this\s+article)"
    r"\s*:|"
    r"https?://|www\.|doi\s*:|©|copyright"
    r")",
    re.IGNORECASE,
)


class DocxRenderer(BaseRenderer):
    """将 Document 模型渲染为 DOCX 格式。

    实现三种基于区域的渲染策略：
      - single_col: 直接顺序段落输出
      - multi_col_wcols: Word 原生 w:cols 分栏
      - multi_col_table: 无边框布局表格（用于跨列内容）
    """

    _PAGE_FIT_SCALES = (
        1.00, 0.99, 0.98, 0.97, 0.96, 0.95, 0.94, 0.93, 0.92, 0.91,
        0.90, 0.89, 0.88, 0.87, 0.86, 0.85, 0.84, 0.83, 0.82, 0.81,
        0.80, 0.79, 0.78, 0.77, 0.76, 0.75, 0.74, 0.73, 0.72, 0.71,
        0.70,
    )
    _PAGE_FIT_MIN_SCALE = 0.70
    _PAGE_FIT_HEADROOM = 1.00
    _TEXT_WRAP_RISK_CJK = 1.04
    _TEXT_WRAP_RISK_LATIN = 1.08
    _TABLE_HEIGHT_RISK = 1.07
    _IMAGE_HEIGHT_RISK = 1.05

    def __init__(self, config=None) -> None:
        super().__init__(config=config)
        self._fit_scale: float = 1.0
        # Hierarchical correction budget (设计文档 §5.2):
        # Before applying global fit_scale, try local corrections first.
        self._corr_space_after_pt: float = 0.0  # 段后间距削减量
        self._corr_gap_pt: float = 0.0          # 区块间隙削减量
        self._corr_font_pt: float = 0.0         # 字号削减量（pt），溢出控制用
        self._font_floor: float = 8.5           # 最小可读字号（pt）

    def render(self, document: "Document", output_path: str, **options) -> None:
        expected_pages = int(options.get("expected_pages", len(document.pages)))
        enforce_single_page = bool(options.get("enforce_single_page", True))
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        tmp_output = output.with_name(f".{output.name}.tmp")
        if enforce_single_page and expected_pages == 1:
            doc = self._render_single_page_fit(document, expected_pages, **options)
            doc.save(tmp_output)
            tmp_output.replace(output)
            return
        doc = self._build_docx(document, **options)
        doc.save(tmp_output)
        tmp_output.replace(output)

    def render_bytes(self, document: "Document", **options) -> bytes:
        expected_pages = int(options.get("expected_pages", len(document.pages)))
        enforce_single_page = bool(options.get("enforce_single_page", True))
        if enforce_single_page and expected_pages == 1:
            doc = self._render_single_page_fit(document, expected_pages, **options)
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        doc = self._build_docx(document, **options)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------
    # 核心构建
    # ------------------------------------------------------------------

    def _build_docx(self, document: "Document", **options) -> DocxDocument:
        doc = DocxDocument()

        # 默认样式 -- 由 RecoveryConfig 驱动
        cfg = self.config
        style_normal = doc.styles['Normal']
        style_normal.font.name = cfg.default_font
        style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), cfg.default_cjk_font)
        style_normal.font.size = Pt(self._scale(cfg.default_font_size_pt))
        style_normal.paragraph_format.space_after = Pt(0)
        style_normal.paragraph_format.space_before = Pt(0)
        style_normal.paragraph_format.line_spacing = cfg.default_line_spacing

        for page in document.pages:
            self._render_page(doc, page)

        return doc

    def _render_page(self, doc: DocxDocument, page: "Page") -> None:
        sect = doc.sections[0] if len(doc.sections) == 1 and not doc.paragraphs else doc.add_section(WD_SECTION.NEW_PAGE)
        self._setup_section(sect, page)
        self._set_section_columns(sect, 1)
        base_sect = sect
        current_cols = 1

        mapper = page.coord_mapper
        usable_w_pt = page.usable_width_pt
        img_w = page.image_width

        # RenderPlan 提示（通过 test.py 注入到 page.attributes）
        page_render_mode = ""
        if page.attributes:
            page_render_mode = page.attributes.get("render_mode", "")

        render_zones = self._prepare_edge_decorative_text_zones(page.zones, page)
        render_zones = self._merge_adjacent_visual_text_zones(render_zones, page)
        render_zones = self._split_embedded_visual_text_bands(render_zones, page)

        for zi, zone in enumerate(render_zones):
            # 预计算每个区域的列像素边界
            col_px = self._build_render_col_px(zone)

            strategy = zone.rendering_strategy
            local_visual_band = zone.col_count > 1 and self._is_local_visual_zone(zone, page)
            decorative_sidecar = zone.region_kind == "decorative_sidecar"
            # RenderPlan 模式覆盖。reflow 页面仍允许局部图文 sidecar band
            # 使用布局表格，否则图片会独占一整行并切断正文流。
            if page_render_mode == "reflow" and not (local_visual_band or decorative_sidecar):
                strategy = "single_col"

            use_native_cols = self._should_use_native_columns(
                zone,
                col_px=col_px,
                page_width_px=img_w,
            )
            if local_visual_band:
                use_native_cols = False
            if decorative_sidecar:
                use_native_cols = False
            # RenderPlan 的 native_columns 只是页面级偏好，实际仍由 zone
            # 安全检查决定；复杂图文混排若强制 Word 原生分栏，容易把后续
            # 块流入不可见/错误栏位。
            # RenderPlan 的 grid 模式只对跨栏 zone 强制表格；普通多栏正文
            # 继续使用 Word 原生分栏，和旧版 newspaper/magazine 输出一致。
            if page_render_mode == "grid" and zone.has_spanned:
                use_native_cols = False

            desired_cols = zone.col_count if use_native_cols else 1
            if desired_cols != current_cols:
                new_sect = doc.add_section(WD_SECTION.CONTINUOUS)
                self._copy_section(base_sect, new_sect)
                if use_native_cols:
                    col_widths = self._column_widths_pt(
                        zone.col_count, col_px, img_w, usable_w_pt,
                    )
                    self._set_section_columns(new_sect, desired_cols, col_widths_pt=col_widths)
                else:
                    self._set_section_columns(new_sect, desired_cols)
                current_cols = desired_cols

            if strategy == 'single_col':
                cl = col_px.get(0, [0, img_w])[0]
                cr = col_px.get(0, [0, img_w])[1]
                ctx = RenderContext(
                    coord_mapper=mapper, page=page,
                    col_width_pt=usable_w_pt,
                    col_left_px=cl, col_right_px=cr,
                )
                setattr(ctx, "render_mode", page_render_mode)
                prev_y = 0
                for block in zone.blocks:
                    gap = max(0, block.bbox.y1 - prev_y)
                    sp = self._scale(max(min(mapper.h(gap), 18) - self._corr_gap_pt, 0)) if (prev_y > 0 and gap > 2) else 0
                    self._render_block(doc, block, ctx, space_before=sp)
                    prev_y = block.bbox.y2

            elif strategy == 'strip_row':
                self._render_strip_row_zone(doc, zone, page)

            elif use_native_cols:
                self._render_native_columns_zone(doc, zone, page, col_px)
            else:  # multi_col_table
                self._render_layout_table_zone(doc, zone, page, col_px)

        # 记录 render_fit 与 style_inferred 元数据（设计文档 §2.2）
        if page.attributes is None:
            page.attributes = {}
        if self._fit_scale < 1.0:
            est_h = self._estimate_page_content_height_pt(page)
            page.attributes["render_fit"] = {
                "page_scale": self._fit_scale,
                "reason": "content_exceeds_page",
                "estimated_content_pt": round(est_h, 1),
                "usable_height_pt": round(page.usable_height_pt, 1),
                "excess_pt": round(est_h - page.usable_height_pt, 1),
            }
        page.attributes["style_inferred"] = self._build_page_style_inferred(page)

        # 清理尾部空段落
        self._cleanup_trailing_paragraphs(doc)

    # ------------------------------------------------------------------
    # 节格式辅助
    # ------------------------------------------------------------------

    def _setup_section(self, sect, page: "Page") -> None:
        sect.page_width = Cm(page.page_width_pt / 72 * 2.54)
        sect.page_height = Cm(page.page_height_pt / 72 * 2.54)
        is_landscape = page.orientation == 'landscape'
        sect.orientation = WD_ORIENT.LANDSCAPE if is_landscape else WD_ORIENT.PORTRAIT
        sect.left_margin = Pt(page.margin_left_pt)
        sect.right_margin = Pt(page.margin_right_pt)
        sect.top_margin = Pt(page.margin_top_pt)
        sect.bottom_margin = Pt(page.margin_bottom_pt)

    # ------------------------------------------------------------------
    # 渲染前 zone 规整
    # ------------------------------------------------------------------

    def _prepare_edge_decorative_text_zones(self, zones: List[Zone], page: "Page") -> List[Zone]:
        """Keep marginal labels visible without inserting them into body flow.

        Textbook pages often contain vertical section labels or teacher-note
        stickers on the outer margin.  They are useful visual decoration, but
        rendering them as ordinary paragraphs shifts the real article body.
        Left-side labels are rendered in a narrow sidecar column; right-edge
        labels are still removed from the editable flow until we can anchor them
        reliably without disturbing surrounding content.
        """
        if not zones or getattr(page, "image_width", 0) <= 0:
            return list(zones)

        output: List[Zone] = []
        for zone in zones:
            if zone.rendering_strategy == "strip_row" or not zone.blocks:
                output.append(zone)
                continue
            left_decorative: List[Block] = []
            kept: List[Block] = []
            for block in zone.blocks:
                decorative_side = self._edge_decorative_side(block, page, zone.blocks)
                if decorative_side == "left":
                    left_decorative.append(block)
                    continue
                if decorative_side == "right":
                    if block.attributes is None:
                        block.attributes = {}
                    block.attributes["docx_suppressed_reason"] = "edge_decorative_text"
                    continue
                kept.append(block)
            if not kept:
                continue
            if not left_decorative and len(kept) == len(zone.blocks):
                output.append(zone)
                continue

            if left_decorative:
                render_blocks: List[Block] = []
                for block in sorted(left_decorative, key=lambda item: (item.bbox.y1, item.bbox.x1)):
                    clone = self._clone_block_for_render_column(block, col_index=0, col_count=2)
                    clone.attributes["docx_decorative_role"] = "left_sidecar"
                    render_blocks.append(clone)
                for block in kept:
                    clone = self._clone_block_for_render_column(block, col_index=1, col_count=2)
                    render_blocks.append(clone)
                output.append(Zone(
                    col_count=2,
                    blocks=render_blocks,
                    has_spanned=False,
                    flow_id=zone.flow_id,
                    flow_kind=zone.flow_kind,
                    region_id=zone.region_id,
                    region_kind="decorative_sidecar",
                ))
                continue

            output.append(Zone(
                col_count=max(1, max((int(getattr(block, "col_index", 0) or 0) for block in kept), default=0) + 1),
                blocks=kept,
                has_spanned=any(len(getattr(block, "spanned_cols", []) or []) > 1 for block in kept),
                flow_id=zone.flow_id,
                flow_kind=zone.flow_kind,
                region_id=zone.region_id,
                region_kind=zone.region_kind,
            ))
        return output

    @staticmethod
    def _clone_block_for_render_column(block: Block, *, col_index: int, col_count: int) -> Block:
        clone = copy(block)
        clone.col_index = col_index
        clone.col_count = col_count
        clone.spanned_cols = [col_index]
        clone.attributes = dict(getattr(block, "attributes", None) or {})
        return clone

    @classmethod
    def _edge_decorative_side(cls, block: Block, page: "Page", zone_blocks: List[Block]) -> Optional[str]:
        if not cls._is_edge_decorative_text(block, page, zone_blocks):
            return None
        page_w = max(float(getattr(page, "image_width", 0) or 0), 1.0)
        body_blocks = [
            other for other in zone_blocks
            if other is not block
            and isinstance(other, TextBlock)
            and other.block_type in {BlockType.TEXT, BlockType.ABSTRACT, BlockType.REFERENCE}
            and len((other.full_text() or "").strip()) >= 40
            and float(other.bbox.width) >= page_w * 0.30
        ]
        if not body_blocks:
            return None
        body_left = min(float(other.bbox.x1) for other in body_blocks)
        body_right = max(float(other.bbox.x2) for other in body_blocks)
        if float(block.bbox.x2) <= body_left - page_w * 0.025 or float(block.bbox.x1) <= page_w * 0.12:
            return "left"
        if float(block.bbox.x1) >= body_right + page_w * 0.035 or float(block.bbox.x2) >= page_w * 0.88:
            return "right"
        return None

    def _suppress_edge_decorative_text_zones(self, zones: List[Zone], page: "Page") -> List[Zone]:
        """Backward-compatible alias for tests and older call sites."""
        return self._prepare_edge_decorative_text_zones(zones, page)

    @staticmethod
    def _is_edge_decorative_text(block: Block, page: "Page", zone_blocks: List[Block]) -> bool:
        if not isinstance(block, TextBlock):
            return False
        if block.block_type not in {BlockType.TEXT, BlockType.TITLE}:
            return False
        text = (block.full_text() or "").strip()
        if not text:
            return False

        page_w = max(float(getattr(page, "image_width", 0) or 0), 1.0)
        page_h = max(float(getattr(page, "image_height", 0) or 0), 1.0)
        x1 = float(block.bbox.x1)
        x2 = float(block.bbox.x2)
        width = max(float(block.bbox.width), 1.0)
        height = max(float(block.bbox.height), 1.0)
        left_edge = x1 <= page_w * 0.12
        right_edge = x2 >= page_w * 0.88
        vertical_label = height >= width * 1.6 and len(text) <= 18
        short_margin_label = len(text) <= 18 and width <= page_w * 0.16

        body_blocks = [
            other for other in zone_blocks
            if other is not block
            and isinstance(other, TextBlock)
            and other.block_type in {BlockType.TEXT, BlockType.ABSTRACT, BlockType.REFERENCE}
            and len((other.full_text() or "").strip()) >= 40
            and float(other.bbox.width) >= page_w * 0.30
        ]
        if not body_blocks:
            return False

        body_left = min(float(other.bbox.x1) for other in body_blocks)
        body_right = max(float(other.bbox.x2) for other in body_blocks)
        left_of_body = x2 <= body_left - page_w * 0.025
        right_of_body = x1 >= body_right + page_w * 0.035
        if not ((left_edge or right_edge or left_of_body or right_of_body) and (vertical_label or short_margin_label)):
            return False

        if left_edge or left_of_body:
            if body_left - x2 < page_w * 0.025:
                return False
        if right_edge:
            if x1 - body_right < page_w * 0.05:
                return False

        # Large sidebar notes are content, not decoration.
        if height >= page_h * 0.10 and len(text) >= 24:
            return False
        return True

    def _merge_adjacent_visual_text_zones(self, zones: List[Zone], page: "Page") -> List[Zone]:
        """Merge local side-by-side image/text bands split by upstream zones.

        PP-DocLayoutV3 often emits the right text, caption line, and visual block
        as separate reading-order items.  If those items become separate zones,
        Word renders them sequentially and creates large blank areas.  This
        renderer-only pass groups adjacent, vertically-overlapping visual/text
        zones into a compact two-column table while leaving page reading order
        untouched for non-DOCX outputs.
        """
        if len(zones) < 2 or getattr(page, "image_width", 0) <= 0:
            return list(zones)

        merged: List[Zone] = []
        idx = 0
        while idx < len(zones):
            zone = zones[idx]
            if not self._zone_has_local_visual_text_evidence(zone, page):
                merged.append(zone)
                idx += 1
                continue

            cluster = [zone]
            visual_blocks = [
                block for block in zone.blocks
                if block.block_type in {BlockType.FIGURE, BlockType.TABLE}
            ]
            band_top = min(float(b.bbox.y1) for b in visual_blocks)
            band_bottom = max(float(b.bbox.y2) for b in visual_blocks)
            lookahead = idx + 1
            while lookahead < len(zones) and len(cluster) < 4:
                cand = zones[lookahead]
                if cand.rendering_strategy == "strip_row":
                    break
                if not self._zone_can_join_visual_text_band(
                    cand,
                    page=page,
                    band_top=band_top,
                    band_bottom=band_bottom,
                    visual_blocks=visual_blocks,
                ):
                    break
                cluster.append(cand)
                for block in cand.blocks:
                    band_top = min(band_top, float(block.bbox.y1))
                    band_bottom = max(band_bottom, float(block.bbox.y2))
                lookahead += 1

            if len(cluster) == 1 and not self._is_local_visual_zone(zone, page):
                merged.append(zone)
                idx += 1
                continue

            if len(cluster) == 1:
                merged.append(zone)
                idx += 1
                continue

            cluster_blocks: List[Block] = []
            for item in cluster:
                cluster_blocks.extend(item.blocks)
            self._assign_render_band_columns(cluster_blocks, page)
            merged.append(Zone(
                col_count=2,
                blocks=cluster_blocks,
                has_spanned=any(len(getattr(b, "spanned_cols", []) or []) > 1 for b in cluster_blocks),
                flow_id=zone.flow_id,
                flow_kind=zone.flow_kind,
                region_id=zone.region_id,
                region_kind=zone.region_kind,
            ))
            idx = lookahead

        return merged

    def _split_embedded_visual_text_bands(self, zones: List[Zone], page: "Page") -> List[Zone]:
        if not zones or getattr(page, "image_width", 0) <= 0:
            return list(zones)

        output: List[Zone] = []
        for zone in zones:
            if zone.rendering_strategy == "strip_row" or len(zone.blocks) < 4:
                output.append(zone)
                continue
            if zone.col_count >= 3:
                output.append(zone)
                continue
            if zone.col_count > 1 and not zone.has_spanned:
                output.append(zone)
                continue
            visual_bands = self._embedded_visual_bands(zone, page)
            if not visual_bands:
                output.append(zone)
                continue

            scan_blocks = self._zone_blocks_for_embedded_visual_split(zone, visual_bands)
            consumed: set[int] = set()
            cursor: List[Block] = []
            band_by_first_id = {id(band[0]): band for band in visual_bands}
            for block in scan_blocks:
                if id(block) in consumed:
                    continue
                band = band_by_first_id.get(id(block))
                if band is None:
                    cursor.append(block)
                    continue
                if cursor:
                    output.append(Zone(
                        col_count=zone.col_count,
                        blocks=cursor,
                        has_spanned=any(len(getattr(b, "spanned_cols", []) or []) > 1 for b in cursor),
                        flow_id=zone.flow_id,
                        flow_kind=zone.flow_kind,
                        region_id=zone.region_id,
                        region_kind=zone.region_kind,
                    ))
                    cursor = []
                for band_block in band:
                    consumed.add(id(band_block))
                leading_blocks, prefix_titles, band = self._split_local_section_titles_from_band(band, page)
                if leading_blocks:
                    output.append(Zone(
                        col_count=1,
                        blocks=leading_blocks,
                        has_spanned=False,
                        flow_id=zone.flow_id,
                        flow_kind=zone.flow_kind,
                        region_id=zone.region_id,
                        region_kind=zone.region_kind,
                    ))
                if prefix_titles:
                    output.append(Zone(
                        col_count=1,
                        blocks=prefix_titles,
                        has_spanned=False,
                        flow_id=zone.flow_id,
                        flow_kind=zone.flow_kind,
                        region_id=zone.region_id,
                        region_kind=zone.region_kind,
                    ))
                if len(band) < 3:
                    cursor.extend(band)
                    continue
                self._assign_render_band_columns(band, page)
                output.append(Zone(
                    col_count=2,
                    blocks=band,
                    has_spanned=False,
                    flow_id=zone.flow_id,
                    flow_kind=zone.flow_kind,
                    region_id=zone.region_id,
                    region_kind=zone.region_kind,
                ))
            if cursor:
                output.append(Zone(
                    col_count=zone.col_count,
                    blocks=cursor,
                    has_spanned=any(len(getattr(b, "spanned_cols", []) or []) > 1 for b in cursor),
                    flow_id=zone.flow_id,
                    flow_kind=zone.flow_kind,
                    region_id=zone.region_id,
                    region_kind=zone.region_kind,
                ))
        return output

    @staticmethod
    def _zone_blocks_for_embedded_visual_split(
        zone: Zone,
        visual_bands: List[List[Block]],
    ) -> List[Block]:
        if not visual_bands or zone.col_count != 1:
            return list(zone.blocks)

        band_ids = {id(block) for band in visual_bands for block in band}
        band_first_ids = {id(band[0]) for band in visual_bands if band}
        band_top_by_id = {
            id(block): min(float(item.bbox.y1) for item in band)
            for band in visual_bands
            for block in band
        }

        def _key(block: Block) -> tuple[float, float, int]:
            if id(block) in band_ids:
                return (band_top_by_id[id(block)], float(block.bbox.x1), 1)
            return (float(block.bbox.y1), float(block.bbox.x1), 0)

        scan_blocks = [
            block for block in zone.blocks
            if id(block) not in band_ids or id(block) in band_first_ids
        ]
        return sorted(scan_blocks, key=_key)

    @staticmethod
    def _split_local_section_titles_from_band(
        band: List[Block],
        page: "Page",
    ) -> Tuple[List[Block], List[Block], List[Block]]:
        """Keep local section headings above a sidecar visual band.

        The sidecar band represents paragraph text flowing beside a visual. A
        short local heading with following body/visual evidence is section
        structure, not part of the visual/caption group, so it should render as
        a single-column heading above the paragraph. Blocks geometrically above
        that heading remain in the preceding normal flow.
        """
        visuals = [block for block in band if block.block_type in {BlockType.FIGURE, BlockType.TABLE}]
        if not visuals:
            return [], [], band
        page_h = max(float(getattr(page, "image_height", 0) or 0), 1.0)
        visual_top = min(float(block.bbox.y1) for block in visuals)
        local_titles = [
            block for block in band
            if isinstance(block, TextBlock)
            and block.block_type == BlockType.TITLE
            and (block.full_text() or "").strip()
            and not _NUMBERED_TITLE_RE.match((block.full_text() or "").strip())
            and float(block.bbox.y1) <= visual_top + page_h * 0.08
            and not (
                any(
                    DocxRenderer._is_side_of_visual(block, visual)
                    and float(block.bbox.y2) >= float(visual.bbox.y1) + 6.0
                    for visual in visuals
                )
            )
        ]
        if not local_titles:
            return [], [], band

        def _has_following_body(title: Block) -> bool:
            return any(
                isinstance(block, TextBlock)
                and block.block_type in {BlockType.TEXT, BlockType.REFERENCE, BlockType.ABSTRACT}
                and float(block.bbox.y1) >= float(title.bbox.y2) - max(float(title.bbox.height), 24.0)
                for block in band
            )

        local_titles = [title for title in local_titles if _has_following_body(title)]
        if not local_titles:
            return [], [], band

        split_title = min(local_titles, key=lambda block: (float(block.bbox.y1), float(block.bbox.x1)))
        title_top = float(split_title.bbox.y1)
        title_bottom = float(split_title.bbox.y2)
        leading: List[Block] = []
        prefix_titles: List[Block] = [split_title]
        remaining: List[Block] = []
        for block in band:
            if block is split_title:
                block.col_count = 1
                block.col_index = 0
                block.spanned_cols = [0]
                continue
            if float(block.bbox.y2) <= title_top + max(float(split_title.bbox.height), 24.0):
                block.col_count = 1
                block.col_index = 0
                block.spanned_cols = [0]
                leading.append(block)
            elif (
                isinstance(block, TextBlock)
                and block.block_type == BlockType.TITLE
                and not _NUMBERED_TITLE_RE.match((block.full_text() or "").strip())
                and abs(float(block.bbox.y1) - title_top) <= max(float(split_title.bbox.height), 24.0)
            ):
                block.col_count = 1
                block.col_index = 0
                block.spanned_cols = [0]
                prefix_titles.append(block)
            else:
                remaining.append(block)
        return leading, prefix_titles, remaining

    @staticmethod
    def _embedded_visual_anchors(zone: Zone, page: "Page") -> List[Block]:
        page_w = max(float(getattr(page, "image_width", 0) or 0), 1.0)
        page_h = max(float(getattr(page, "image_height", 0) or 0), 1.0)
        return [
            block for block in zone.blocks
            if block.block_type in {BlockType.FIGURE, BlockType.TABLE}
            and float(block.bbox.width) <= page_w * 0.46
            and float(block.bbox.height) >= max(page_h * 0.10, 120.0)
        ]

    def _embedded_visual_bands(self, zone: Zone, page: "Page") -> List[List[Block]]:
        bands: List[List[Block]] = []
        used: set[int] = set()
        order = {id(block): idx for idx, block in enumerate(zone.blocks)}
        for visual in self._embedded_visual_anchors(zone, page):
            if id(visual) in used:
                continue
            band = self._embedded_visual_band_blocks(zone, visual, page)
            if len(band) < 3:
                continue
            if any(id(block) in used for block in band):
                continue
            for block in band:
                used.add(id(block))
            bands.append(band)
        bands.sort(key=lambda band: min(order.get(id(block), 10**9) for block in band))
        return bands

    @staticmethod
    def _embedded_visual_band_blocks(zone: Zone, visual: Block, page: "Page") -> List[Block]:
        page_w = max(float(getattr(page, "image_width", 0) or 0), 1.0)
        page_h = max(float(getattr(page, "image_height", 0) or 0), 1.0)
        y_top = float(visual.bbox.y1) - page_h * 0.04
        y_bottom = float(visual.bbox.y2) + page_h * 0.06
        allowed = {
            BlockType.TEXT,
            BlockType.TITLE,
            BlockType.FIGURE_CAPTION,
            BlockType.TABLE_CAPTION,
            BlockType.REFERENCE,
            BlockType.ABSTRACT,
        }
        side_text: List[Block] = []
        captions: List[Block] = []
        for block in zone.blocks:
            if block is visual or block.block_type not in allowed:
                continue
            overlap = min(float(block.bbox.y2), float(visual.bbox.y2)) - max(float(block.bbox.y1), float(visual.bbox.y1))
            close_vertical = float(block.bbox.y1) <= y_bottom and float(block.bbox.y2) >= y_top
            side_by_side = DocxRenderer._is_side_of_visual(block, visual)
            near_vertical = (
                abs(float(block.bbox.y2) - float(visual.bbox.y1)) <= page_h * 0.04
                or abs(float(block.bbox.y1) - float(visual.bbox.y2)) <= page_h * 0.04
            )
            if side_by_side and close_vertical and (overlap >= 18.0 or near_vertical or block.block_type == BlockType.TITLE):
                side_text.append(block)
                continue
            caption_like = (
                block.block_type in {BlockType.FIGURE_CAPTION, BlockType.TABLE_CAPTION}
                and abs(float(block.bbox.y1) - float(visual.bbox.y2)) <= page_h * 0.08
            )
            if caption_like:
                captions.append(block)
        if not side_text:
            return []
        band = [visual] + side_text + captions
        order = {id(block): idx for idx, block in enumerate(zone.blocks)}
        band.sort(key=lambda block: order.get(id(block), 10**9))
        return band

    @staticmethod
    def _is_side_of_visual(block: Block, visual: Block) -> bool:
        block_cx = (float(block.bbox.x1) + float(block.bbox.x2)) * 0.5
        visual_cx = (float(visual.bbox.x1) + float(visual.bbox.x2)) * 0.5
        x_overlap = max(
            0.0,
            min(float(block.bbox.x2), float(visual.bbox.x2)) - max(float(block.bbox.x1), float(visual.bbox.x1)),
        )
        max_allowed_x_overlap = min(float(block.bbox.width), float(visual.bbox.width)) * 0.12
        return (
            (
                block_cx < visual_cx
                and float(block.bbox.x1) < float(visual.bbox.x1)
                and x_overlap <= max_allowed_x_overlap
            )
            or (
                block_cx > visual_cx
                and float(block.bbox.x2) > float(visual.bbox.x2)
                and x_overlap <= max_allowed_x_overlap
            )
        )

    @classmethod
    def _zone_has_local_visual_text_evidence(cls, zone: Zone, page: "Page") -> bool:
        if cls._is_local_visual_zone(zone, page):
            return True
        if zone.rendering_strategy == "strip_row" or not zone.blocks:
            return False
        visuals = [
            block for block in zone.blocks
            if block.block_type in {BlockType.FIGURE, BlockType.TABLE}
        ]
        if len(visuals) != 1:
            return False
        visual = visuals[0]
        text_blocks = [
            block for block in zone.blocks
            if block is not visual
            and block.block_type in {
                BlockType.TEXT,
                BlockType.TITLE,
                BlockType.FIGURE_CAPTION,
                BlockType.TABLE_CAPTION,
                BlockType.REFERENCE,
                BlockType.ABSTRACT,
            }
        ]
        if not text_blocks:
            return False
        return cls._zone_can_join_visual_text_band(
            Zone(col_count=zone.col_count, blocks=text_blocks, has_spanned=zone.has_spanned),
            page=page,
            band_top=float(visual.bbox.y1),
            band_bottom=float(visual.bbox.y2),
            visual_blocks=[visual],
        )

    @staticmethod
    def _is_local_visual_zone(zone: Zone, page: "Page") -> bool:
        if zone.rendering_strategy == "strip_row" or not zone.blocks:
            return False
        page_w = max(float(getattr(page, "image_width", 0) or 0), 1.0)
        page_h = max(float(getattr(page, "image_height", 0) or 0), 1.0)
        visuals = [
            block for block in zone.blocks
            if block.block_type in {BlockType.FIGURE, BlockType.TABLE}
        ]
        if len(visuals) != 1:
            return False
        visual = visuals[0]
        return (
            float(visual.bbox.width) <= page_w * 0.50
            and float(visual.bbox.height) >= max(page_h * 0.14, 160.0)
        )

    @staticmethod
    def _zone_can_join_visual_text_band(
        zone: Zone,
        *,
        page: "Page",
        band_top: float,
        band_bottom: float,
        visual_blocks: List[Block],
    ) -> bool:
        if not zone.blocks:
            return False
        page_w = max(float(getattr(page, "image_width", 0) or 0), 1.0)
        page_h = max(float(getattr(page, "image_height", 0) or 0), 1.0)
        if page_w >= 900 and zone.col_count >= 4:
            return False
        allowed = {
            BlockType.TEXT,
            BlockType.TITLE,
            BlockType.FIGURE_CAPTION,
            BlockType.TABLE_CAPTION,
            BlockType.REFERENCE,
            BlockType.ABSTRACT,
        }
        if any(block.block_type not in allowed for block in zone.blocks):
            return False

        zone_top = min(float(block.bbox.y1) for block in zone.blocks)
        zone_bottom = max(float(block.bbox.y2) for block in zone.blocks)
        y_overlap = min(band_bottom, zone_bottom) - max(band_top, zone_top)
        near_band = y_overlap >= 24.0 or abs(zone_top - band_bottom) <= page_h * 0.08 or abs(band_top - zone_bottom) <= page_h * 0.06
        if not near_band:
            return False

        for visual in visual_blocks:
            left_gap = float(visual.bbox.x1) - max(float(block.bbox.x2) for block in zone.blocks)
            right_gap = min(float(block.bbox.x1) for block in zone.blocks) - float(visual.bbox.x2)
            side_by_side = left_gap >= page_w * 0.02 or right_gap >= page_w * 0.02
            caption_like = (
                zone_bottom <= float(visual.bbox.y1) + page_h * 0.04
                or zone_top <= float(visual.bbox.y1) + page_h * 0.04
            )
            wide_band_text = max(float(block.bbox.width) for block in zone.blocks) >= page_w * 0.52
            if side_by_side or (caption_like and wide_band_text):
                return True
        return False

    @staticmethod
    def _assign_render_band_columns(blocks: List[Block], page: "Page") -> None:
        page_w = max(float(getattr(page, "image_width", 0) or 0), 1.0)
        visuals = [
            block for block in blocks
            if block.block_type in {BlockType.FIGURE, BlockType.TABLE}
        ]
        visual = visuals[0] if visuals else None
        divider = page_w * 0.5
        if visual is not None:
            side_text = [
                block for block in blocks
                if block is not visual
                and DocxRenderer._is_side_of_visual(block, visual)
            ]
            if side_text:
                text_cx = sum((float(b.bbox.x1) + float(b.bbox.x2)) * 0.5 for b in side_text) / len(side_text)
                if text_cx > (float(visual.bbox.x1) + float(visual.bbox.x2)) * 0.5:
                    divider = (float(visual.bbox.x2) + min(float(b.bbox.x1) for b in side_text)) * 0.5
                else:
                    divider = (float(visual.bbox.x1) + max(float(b.bbox.x2) for b in side_text)) * 0.5

        for block in blocks:
            cx = (float(block.bbox.x1) + float(block.bbox.x2)) * 0.5
            block.col_count = 2
            block.col_index = 0 if cx < divider else 1
            if (
                visual is not None
                and block is not visual
                and float(block.bbox.x1) < divider < float(block.bbox.x2)
            ):
                visual_cx = (float(visual.bbox.x1) + float(visual.bbox.x2)) * 0.5
                block.col_index = 0 if visual_cx < divider else 1
            block.spanned_cols = [block.col_index]

    @staticmethod
    def _local_visual_band_col_widths_pt(
        zone: Zone,
        page: "Page",
        *,
        total_width_pt: float,
        visual_gap_pt: float,
    ) -> List[float]:
        if zone.col_count != 2 or total_width_pt <= 0:
            return DocxRenderer._column_widths_pt(
                max(int(zone.col_count or 1), 1),
                DocxRenderer._build_render_col_px(zone),
                float(getattr(page, "image_width", 0) or 0),
                total_width_pt,
            )

        page_w = max(float(getattr(page, "image_width", 0) or 0), 1.0)
        visuals = [block for block in zone.blocks if block.block_type in {BlockType.FIGURE, BlockType.TABLE}]
        visual = visuals[0] if visuals else None
        if visual is None:
            return [total_width_pt / 2.0, total_width_pt / 2.0]

        side_blocks = [
            block for block in zone.blocks
            if block is not visual and DocxRenderer._is_side_of_visual(block, visual)
        ]
        if not side_blocks:
            return [total_width_pt / 2.0, total_width_pt / 2.0]

        if int(getattr(visual, "col_index", 0) or 0) == 0:
            side_left = min(float(block.bbox.x1) for block in side_blocks)
            divider_px = (float(visual.bbox.x2) + side_left) * 0.5
            left_ratio = max(0.12, min(0.48, divider_px / page_w))
            left_width = total_width_pt * left_ratio
            right_width = max(24.0, total_width_pt - left_width - visual_gap_pt)
            return [left_width, right_width]

        side_right = max(float(block.bbox.x2) for block in side_blocks)
        divider_px = (side_right + float(visual.bbox.x1)) * 0.5
        left_ratio = max(0.12, min(0.88, divider_px / page_w))
        left_width = total_width_pt * left_ratio
        right_width = max(24.0, total_width_pt - left_width - visual_gap_pt)
        return [left_width, right_width]

    def _copy_section(self, src_sect, dst_sect) -> None:
        dst_sect.page_width = src_sect.page_width
        dst_sect.page_height = src_sect.page_height
        dst_sect.left_margin = src_sect.left_margin
        dst_sect.right_margin = src_sect.right_margin
        dst_sect.top_margin = src_sect.top_margin
        dst_sect.bottom_margin = src_sect.bottom_margin

    def _set_section_columns(self, sect, num_cols: int, col_widths_pt=None) -> None:
        """设置节分栏，委托 section_fmt 处理（支持等宽/自定义列宽）。

        自定义列宽时，列宽总和必须扣除栏间距占用的空间，否则
        总宽度 + 间距 会超过可用宽度，导致内容从右边溢出。
        """
        n = max(1, int(num_cols))
        gap_twips = int(getattr(self.config, "docx_column_gap_twips", 720))
        spacing_pt = max(0, gap_twips) / 20.0

        if col_widths_pt is not None:
            # 计算当前节可用宽度
            page_w = sect.page_width
            left_m = sect.left_margin
            right_m = sect.right_margin
            # Convert EMU to pt: 1pt = 914400 EMU / 72 = 12700 EMU
            emu_per_pt = 914400 / 72
            usable_w = float(page_w) / emu_per_pt - float(left_m) / emu_per_pt - float(right_m) / emu_per_pt
            # 扣除栏间距后的总可用列宽
            gap_space = spacing_pt * max(n - 1, 0)
            available_for_cols = usable_w - gap_space
            current_total = sum(col_widths_pt)
            if current_total > 0 and available_for_cols < current_total:
                scale = available_for_cols / current_total
                col_widths_pt = [w * scale for w in col_widths_pt]

        _set_section_columns_fmt(sect._sectPr, n, col_widths_pt=col_widths_pt, spacing_pt=spacing_pt)

    def _render_strip_row_zone(self, doc: DocxDocument, zone: Zone, page: "Page") -> None:
        """将页眉/页脚条带按单行布局渲染，避免混入正文分栏。"""
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

        blocks = [
            block for block in sorted(zone.blocks, key=lambda b: b.bbox.x1)
            if self._has_renderable_content(block)
        ]
        if not blocks:
            return

        mapper = page.coord_mapper
        usable_w_pt = page.usable_width_pt
        centers = [((block.bbox.x1 + block.bbox.x2) * 0.5) for block in blocks]
        boundaries_px = [0.0]
        for idx in range(len(centers) - 1):
            boundaries_px.append((centers[idx] + centers[idx + 1]) * 0.5)
        boundaries_px.append(float(page.image_width))
        widths_pt = [
            max(self._scale(36.0), usable_w_pt * max(boundaries_px[idx + 1] - boundaries_px[idx], 1.0) / max(float(page.image_width), 1.0))
            for idx in range(len(blocks))
        ]

        total_width = sum(widths_pt)
        if total_width > 0:
            widths_pt = [width * (usable_w_pt / total_width) for width in widths_pt]

        tbl = doc.add_table(rows=1, cols=len(blocks))
        tbl.autofit = False
        clear_table_borders(tbl)
        set_table_col_widths(tbl, widths_pt)
        fit_table_to_width(tbl, usable_w_pt)
        row = tbl.rows[0]
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            reset_paragraph_format(cell.paragraphs[0])
            set_paragraph_spacing(
                cell.paragraphs[0],
                line_spacing=self._scale(MIN_LINE_SPACING_PT),
                exact=True,
            )

        for idx, block in enumerate(blocks):
            cell = row.cells[idx]
            ctx = RenderContext(
                coord_mapper=mapper,
                page=page,
                col_width_pt=widths_pt[idx],
                col_left_px=boundaries_px[idx],
                col_right_px=boundaries_px[idx + 1],
                in_table_cell=True,
            )
            self._render_block(cell, block, ctx, space_before=0)
            self._prune_leading_empty_cell_paragraphs(cell)

    @staticmethod
    def _has_renderable_content(block: Block) -> bool:
        if isinstance(block, TextBlock):
            return bool((block.full_text() or "").strip())
        if isinstance(block, (ImageBlock, EquationBlock, TableBlock)):
            return bool(getattr(block, "image_data", None) or getattr(block, "html", None))
        return True

    @staticmethod
    def _column_widths_pt(
        num_cols: int,
        col_px: dict,
        page_width_px: float,
        total_width_pt: float,
    ) -> List[float]:
        if num_cols <= 0 or total_width_pt <= 0:
            return []
        if page_width_px <= 0 or not col_px:
            return [total_width_pt / num_cols] * num_cols

        raw_widths = []
        default_px = max(page_width_px / max(num_cols, 1), 1.0)
        for ci in range(num_cols):
            span = col_px.get(ci)
            if not span:
                raw_widths.append(default_px)
                continue
            raw_widths.append(max(float(span[1]) - float(span[0]), 1.0))

        total_raw = sum(raw_widths)
        if total_raw <= 0:
            return [total_width_pt / num_cols] * num_cols
        return [total_width_pt * width / total_raw for width in raw_widths]

    @staticmethod
    def _build_render_col_px(zone: Zone) -> dict:
        col_px = {}
        num_cols = max(int(getattr(zone, "col_count", 1) or 1), 1)
        for ci in range(num_cols):
            candidates = [
                block for block in zone.blocks
                if block.col_index == ci and len(getattr(block, "spanned_cols", [])) <= 1
            ]
            if not candidates:
                continue

            primary = [
                block for block in candidates
                if isinstance(block, TextBlock) and block.count_lines() >= 2
            ]
            selected = primary or candidates
            x1 = min(float(block.bbox.x1) for block in selected)
            if primary:
                x2_values = sorted(float(block.bbox.x2) for block in selected)
                x2 = x2_values[-2] if len(x2_values) >= 3 else x2_values[-1]
            else:
                x2 = max(float(block.bbox.x2) for block in selected)
            col_px[ci] = [x1, x2]
        return col_px

    @classmethod
    def _has_irregular_column_widths(cls, col_px: dict, page_width_px: float, num_cols: int) -> bool:
        if num_cols <= 1 or page_width_px <= 0 or not col_px:
            return False
        widths_pt = cls._column_widths_pt(num_cols, col_px, page_width_px, 1.0)
        if not widths_pt:
            return False
        positive = [width for width in widths_pt if width > 0]
        if len(positive) < 2:
            return False
        return max(positive) / max(min(positive), 1e-6) >= 1.45

    @staticmethod
    def _column_segments(num_cols: int, spanned_cols: List[int]) -> List[Tuple[str, List[int]]]:
        if not spanned_cols:
            return [("standalone", [ci]) for ci in range(num_cols)]
        spanned_set = set(spanned_cols)
        first = min(spanned_cols)
        last = max(spanned_cols)
        segments: List[Tuple[str, List[int]]] = []
        ci = 0
        while ci < num_cols:
            if ci == first:
                segments.append(("spanned", list(spanned_cols)))
                ci = last + 1
                continue
            if ci in spanned_set:
                ci += 1
                continue
            segments.append(("standalone", [ci]))
            ci += 1
        return segments

    @staticmethod
    def _layout_block_cols(block: Block, num_cols: int, page_width_px: float = 0.0) -> List[int]:
        """Return columns that should affect DOCX layout table segmentation."""
        raw = getattr(block, "spanned_cols", []) or [getattr(block, "col_index", 0)]
        cols = sorted({int(ci) for ci in raw if 0 <= int(ci) < num_cols})
        if not cols:
            cols = [min(max(int(getattr(block, "col_index", 0) or 0), 0), num_cols - 1)]
        if len(cols) <= 1:
            return cols

        strip_like = {BlockType.HEADER, BlockType.FOOTER, BlockType.PAGE_NUMBER}
        width = float(getattr(getattr(block, "bbox", None), "width", 0.0) or 0.0)
        if block.block_type in strip_like and (page_width_px <= 0 or width <= float(page_width_px) * 0.42):
            return [min(max(int(getattr(block, "col_index", cols[0]) or cols[0]), 0), num_cols - 1)]
        return cols

    def _should_use_native_columns(self, zone: Zone, col_px: Optional[dict] = None, page_width_px: float = 0.0) -> bool:
        """无跨列时优先使用 Word 原生分栏，提升可编辑性与还原度。"""
        if not bool(getattr(self.config, "docx_prefer_native_columns", True)):
            return False
        if zone.col_count <= 1:
            return False
        if any(len(getattr(b, "spanned_cols", [])) > 1 for b in zone.blocks):
            return False
        if zone.col_count >= 4:
            return True
        if zone.col_count == 2 and page_width_px >= 1200:
            return True
        if self._has_irregular_column_widths(col_px or {}, page_width_px, zone.col_count):
            return False
        return True

    def _config_style_defaults(self) -> dict:
        return {
            "font_size_pt": self.config.default_font_size_pt,
            "font_family": self.config.default_cjk_font,
            "font_family_western": self.config.default_font,
            "bold": False,
            "italic": False,
            "underline": False,
            "strikethrough": False,
            "superscript": False,
            "subscript": False,
            "alignment": "left",
            "line_spacing": self.config.default_line_spacing,
            "space_after_pt": 1.0,
        }

    @staticmethod
    def _build_page_style_inferred(page: "Page") -> dict:
        """构建页面级 style_inferred 摘要（设计文档 §2.2）。

        记录从源页面坐标系推断出的真实样式，与 render_fit 严格分离。
        """
        from collections import defaultdict

        # 收集各 block 类型的字号/行距中位数
        by_type: dict = defaultdict(list)
        by_type_ls: dict = defaultdict(list)
        for zone in getattr(page, "zones", []):
            for block in getattr(zone, "blocks", []):
                if isinstance(block, TextBlock) and block.style:
                    if block.style.font_size_pt is not None:
                        by_type[block.block_type].append(block.style.font_size_pt)
                    if block.style.line_spacing is not None:
                        by_type_ls[block.block_type].append(block.style.line_spacing)

        def _median(values):
            if not values:
                return None
            s = sorted(values)
            return round(s[len(s) // 2] * 2) / 2.0

        canonical_styles = {}
        for btype, sizes in by_type.items():
            canonical_styles[str(btype)] = {
                "font_size_pt": _median(sizes),
                "line_spacing": _median(by_type_ls.get(btype, [])),
            }

        return {
            "page_margin": {
                "top": round(page.margin_top_pt, 1),
                "bottom": round(page.margin_bottom_pt, 1),
                "left": round(page.margin_left_pt, 1),
                "right": round(page.margin_right_pt, 1),
            },
            "default_font_size_pt": page.style_defaults.get("font_size_pt") if page.style_defaults else None,
            "default_line_spacing": page.style_defaults.get("line_spacing") if page.style_defaults else None,
            "block_styles": canonical_styles,
        }

    @staticmethod
    def _parse_css_hex(color: Optional[str]) -> Optional[RGBColor]:
        if not color:
            return None
        try:
            hex_c = color.lstrip("#")
            if len(hex_c) != 6:
                return None
            return RGBColor(int(hex_c[0:2], 16), int(hex_c[2:4], 16), int(hex_c[4:6], 16))
        except Exception:
            return None

    def _scale(self, value: float) -> float:
        return value * self._fit_scale

    def _scale_font(self, value: float) -> float:
        """缩放字号：应用 fit_scale，带动态最小可读字号保护。

        fit_scale 控制全局内容高度以防止溢页。_font_floor 保护字号
        不致过小（默认 8.5pt）。两遍渲染策略：首遍用 8.5pt 保证可读性，
        若仍溢页则降至 7.0pt 作为兜底。
        """
        return max(self._font_floor, value * self._fit_scale)

    def _scale_title_font(self, block: TextBlock, page: "Page", value: float) -> float:
        """Scale title fonts while preserving source-like masthead emphasis."""
        layout_profile = str((getattr(page, "attributes", None) or {}).get("layout_profile", "") or "")
        if (
            layout_profile == "newspaper_mixed"
            and self._is_masthead_title(block, page, str(getattr(block.style, "alignment", "") or "center"))
        ):
            return max(self._font_floor, value * max(self._fit_scale, 0.88))
        return self._scale_font(value)

    def _render_single_page_fit(
        self, document: "Document", expected_pages: int, **build_options
    ) -> DocxDocument:
        """纯内置单页适配：按内容高度预算选择缩放系数。

        分级修正策略（设计文档 §5.2）：
        1. 先应用局部修正（段后间距、区块间隙、字号）
        2. 对剩余溢出才应用全局 fit_scale
        """
        build_options = dict(build_options)
        build_options.pop("expected_pages", None)
        build_options.pop("enforce_single_page", None)

        scale = 1.0
        self._fit_scale = 1.0
        self._font_floor = 8.5
        trial_doc = self._build_docx(document, **build_options)
        if not self._check_overflow(trial_doc, expected_pages):
            self._fit_scale = 1.0
            self._font_floor = 8.5
            return trial_doc

        scale = self._select_builtin_fit_scale(document, expected_pages)
        if scale < 1.0:
            # 计算溢出量，尝试用局部修正减少部分溢出
            for page in document.pages:
                est_h = self._estimate_page_content_height_pt(page)
                excess = est_h - page.usable_height_pt
                if excess > 0:
                    # 分级修正预算（设计文档 §5.2）：从最不敏感属性开始
                    # 大溢出时使用激进预算
                    if excess > 50:
                        space_after_budget = min(excess * 0.15, 30.0)   # 段后间距：最多 30pt
                        gap_budget = min(excess * 0.20, 45.0)            # 区块间隙：最多 45pt
                        font_budget = min(excess * 0.012, 2.5)           # 字号削减：最多 2.5pt
                    else:
                        space_after_budget = min(excess * 0.20, 16.0)    # 段后间距：最多 16pt
                        gap_budget = min(excess * 0.15, 10.0)             # 区块间隙：最多 10pt
                        font_budget = 0.0

                    n_paragraphs = self._count_paragraphs(page)
                    space_saving = min(space_after_budget, n_paragraphs * 3.0)
                    gap_saving = gap_budget
                    font_saving = font_budget * n_paragraphs
                    local_reduction = space_saving + gap_saving + font_saving

                    self._corr_space_after_pt = space_after_budget
                    self._corr_gap_pt = gap_budget
                    self._corr_font_pt = font_budget

                    # 重新计算剩余溢出所需的 fit_scale
                    remaining_excess = max(0, excess - local_reduction)
                    if remaining_excess > 0:
                        remaining_scale = page.usable_height_pt / max(est_h - remaining_excess, 1.0)
                    else:
                        remaining_scale = 1.0
                    scale = min(scale, remaining_scale)
                    break  # 单页文档，只需处理第一页

        searched_doc = self._render_largest_fitting_scale(
            document,
            expected_pages,
            build_options,
            min_scale=self._PAGE_FIT_MIN_SCALE,
            font_floors=(8.5, 7.0, 6.5),
        )
        if searched_doc is not None:
            self._fit_scale = 1.0
            self._corr_space_after_pt = 0.0
            self._corr_gap_pt = 0.0
            self._corr_font_pt = 0.0
            self._font_floor = 8.5
            return searched_doc

        logger.info("single-page builtin fit selected scale=%.3f expected_pages=%d corr_space=%.1f corr_gap=%.1f corr_font=%.2f",
                    scale, expected_pages,
                    self._corr_space_after_pt, self._corr_gap_pt, self._corr_font_pt)
        try:
            # 两遍渲染：首遍用 8.5pt 字号下限保证可读性
            self._fit_scale = scale
            self._font_floor = 8.5
            doc = self._build_docx(document, **build_options)
            if not self._check_overflow(doc, expected_pages):
                return doc
            # 首遍仍溢页：降低字号下限作为兜底
            self._font_floor = 7.0
            doc = self._build_docx(document, **build_options)
            if not self._check_overflow(doc, expected_pages):
                return doc
            # 仍然溢页：进一步降低到 6.5pt
            logger.info("second-pass overflow, retrying with font_floor=6.5")
            self._font_floor = 6.5
            doc = self._build_docx(document, **build_options)
            return doc
        finally:
            self._fit_scale = 1.0
            self._corr_space_after_pt = 0.0
            self._corr_gap_pt = 0.0
            self._corr_font_pt = 0.0
            self._font_floor = 8.5

    def _render_largest_fitting_scale(
        self,
        document: "Document",
        expected_pages: int,
        build_options: dict,
        *,
        min_scale: float,
        font_floors: tuple[float, ...] = (8.5,),
    ) -> Optional[DocxDocument]:
        """Find the largest readable scale that LibreOffice confirms fits in-page."""
        scales = sorted(
            {
                scale for scale in self._PAGE_FIT_SCALES
                if self._PAGE_FIT_MIN_SCALE <= scale < 1.0 and scale >= min_scale
            }
        )
        if not scales:
            return None

        for floor in font_floors:
            best_doc: Optional[DocxDocument] = None
            best_scale: Optional[float] = None
            lo = 0
            hi = len(scales) - 1
            while lo <= hi:
                mid = (lo + hi) // 2
                candidate_scale = scales[mid]
                self._fit_scale = candidate_scale
                self._font_floor = floor
                doc = self._build_docx(document, **build_options)
                if self._check_overflow(doc, expected_pages):
                    hi = mid - 1
                    continue
                best_doc = doc
                best_scale = candidate_scale
                lo = mid + 1

            if best_doc is not None:
                logger.info(
                    "single-page actual fit selected scale=%.3f font_floor=%.1f expected_pages=%d",
                    best_scale or 1.0,
                    floor,
                    expected_pages,
                )
                return best_doc
        return None

    @classmethod
    def _check_overflow(cls, doc: DocxDocument, expected_pages: int) -> bool:
        """通过 LibreOffice 转换 PDF 并检查实际页数。

        若 LibreOffice 不可用，回退到基于内容密度的启发式估算。
        """
        import tempfile
        import subprocess
        import shutil

        # 尝试使用 LibreOffice 进行准确的页数检查
        lo_path = shutil.which("libreoffice") or shutil.which("soffice")
        if lo_path:
            try:
                with tempfile.TemporaryDirectory() as tmpdir:
                    tmp_doc = os.path.join(tmpdir, "check.docx")
                    doc.save(tmp_doc)
                    result = subprocess.run(
                        [lo_path, "--headless", "--convert-to", "pdf", "--outdir", tmpdir, tmp_doc],
                        capture_output=True, timeout=30,
                    )
                    if result.returncode == 0:
                        pdf_path = os.path.join(tmpdir, "check.pdf")
                        if os.path.exists(pdf_path):
                            # 用 pdfinfo 读取页数
                            info = subprocess.run(
                                ["pdfinfo", pdf_path],
                                capture_output=True, text=True, timeout=10,
                            )
                            for line in info.stdout.splitlines():
                                if line.startswith("Pages:"):
                                    actual_pages = int(line.split(":")[1].strip())
                                    return actual_pages > expected_pages
            except Exception:
                pass  # 回退到启发式方法

        # 回退：基于内容密度的启发式估算
        return cls._check_overflow_heuristic(doc, expected_pages)

    @staticmethod
    def _check_overflow_heuristic(doc: DocxDocument, expected_pages: int) -> bool:
        """通过 DOCX 内容密度估算是否可能溢出（回退方法）。"""
        from lxml import etree

        body = doc.element.body
        ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        text_len = 0
        table_count = 0
        para_count = 0
        for child in body:
            if child.tag == etree.QName(ns_w, 'p'):
                para_count += 1
            elif child.tag == etree.QName(ns_w, 'tbl'):
                table_count += 1
                for p in child.iter(etree.QName(ns_w, 'p').text):
                    para_count += 1

        for t in body.iter(etree.QName(ns_w, 't').text):
            text_len += len(t.text or '')

        if table_count >= 2 and 20 <= para_count <= 45 and 2000 <= text_len <= 6000:
            return True
        if table_count == 1 and 25 <= para_count <= 45 and 2000 <= text_len <= 3500:
            return True
        if text_len >= 8000:
            return True
        return False

    def _count_paragraphs(self, page: "Page") -> int:
        """估算页面段落数，用于计算间距修正的实际效果。"""
        count = 0
        for zone in page.zones:
            for block in zone.blocks:
                from docflow.model.blocks.text_block import TextBlock
                if isinstance(block, TextBlock):
                    if block.paragraphs:
                        count += max(1, len(block.paragraphs))
                    else:
                        count += 1
        return max(count, 1)

    def _select_builtin_fit_scale(self, document: "Document", expected_pages: int) -> float:
        """基于内容高度预算选择缩放档位（纯内置，无外部依赖）。"""
        if not document.pages:
            return 1.0
        if expected_pages <= 0:
            expected_pages = len(document.pages)

        est_heights = [self._estimate_page_content_height_pt(page) for page in document.pages]
        caps = [max(page.usable_height_pt * self._PAGE_FIT_HEADROOM, 1.0) for page in document.pages]

        def _est_pages(scale: float) -> int:
            total = 0
            for h, cap in zip(est_heights, caps):
                used = max(h * scale, 0.0)
                total += max(1, int(math.ceil(used / cap)))
            return total

        if _est_pages(1.0) <= expected_pages:
            return 1.0

        analytic = 1.0
        for h, cap in zip(est_heights, caps):
            if h > 0:
                analytic = min(analytic, cap / h)
        # 截断到两位小数以避免浮点误差
        analytic = max(
            self._PAGE_FIT_MIN_SCALE,
            min(1.0, round(math.floor(analytic * 100.0) / 100.0, 2)),
        )
        safety_margin = self._fit_safety_margin(document)
        if analytic < 1.0 and safety_margin > 0.0:
            analytic = round(max(self._PAGE_FIT_MIN_SCALE, analytic - safety_margin), 2)

        candidates = [s for s in self._PAGE_FIT_SCALES if self._PAGE_FIT_MIN_SCALE <= s <= analytic]
        if not candidates:
            candidates = [self._PAGE_FIT_MIN_SCALE]
        for scale in candidates:
            if _est_pages(scale) <= expected_pages:
                return scale
        return candidates[-1]

    @staticmethod
    def _fit_safety_margin(document: "Document") -> float:
        """为复杂版面选择更保守的单页适配安全余量。

        经验上，多栏布局表格（multi_col_table）渲染的 Word 表格
        会产生额外的行高/单元格开销，即使没有跨列区块也需要更多余量。
        """
        margin = 0.0
        for page in document.pages:
            for zone in getattr(page, "zones", []):
                col_count = max(int(getattr(zone, "col_count", 1) or 1), 1)
                has_spanned = any(
                    len(getattr(block, "spanned_cols", [])) > 1
                    for block in getattr(zone, "blocks", [])
                )
                if col_count >= 4 and has_spanned:
                    return 0.03
                if col_count >= 3 and has_spanned:
                    margin = max(margin, 0.025)
                    continue
                if col_count >= 3:
                    margin = max(margin, 0.02)
                if col_count >= 2:
                    # 2+ 栏的 multi_col_table 即使无跨列也有显著表格开销
                    margin = max(margin, 0.01)
        return margin

    def _estimate_page_content_height_pt(self, page: "Page") -> float:
        """估算页面渲染后的总内容高度（pt）。

        估算结果乘以安全系数，以涵盖区块间开销（节断点、分栏切换、
        节格式差异等）在单区域估算中未被计入的部分。
        """
        mapper = page.coord_mapper
        usable_w_pt = page.usable_width_pt
        total = 0.0
        for zone in page.zones:
            total += self._estimate_zone_height_pt(zone, page, mapper, usable_w_pt)
        # 布局表格 zone 额外开销：表格结构、单元格边距、行最小高度等
        # 在 Word 中，布局表格每行即使无内容也有约 8-12pt 的单元格开销，
        # 加上表格自身的顶部/底部间隙和列分隔线开销。
        # 此外，表格内段落的 space_before/space_after 在单元格中
        # 会有额外的渲染开销（比非表格上下文多出约 1-2pt/段）。
        for zone in page.zones:
            if zone.rendering_strategy == 'multi_col_table':
                n_blocks = sum(1 for _ in zone.blocks)
                total += 50.0 + max(n_blocks - 10, 0) * 3.0
        # 全局安全系数：覆盖 zone 间开销（节断点、分栏切换等）
        # 以及 Word 渲染引擎本身的额外开销（最小行距、单元格边距等）
        # multi_col_table 布局因表格结构开销，需要更高的安全系数
        _PAGE_ESTIMATE_SAFETY = 1.08
        return max(total * _PAGE_ESTIMATE_SAFETY, 1.0)

    def _estimate_zone_height_pt(self, zone: Zone, page: "Page", mapper, usable_w_pt: float) -> float:
        blocks = zone.blocks
        if not blocks:
            return 0.0

        if zone.rendering_strategy == 'single_col':
            return self._cap_zone_height(
                self._estimate_stream_height_pt(
                    blocks=blocks,
                    page=page,
                    mapper=mapper,
                    col_width_pt=usable_w_pt,
                    gap_cap_pt=18.0,
                ),
                blocks, mapper,
            )

        num_cols = max(zone.col_count, 1)
        layout_factor = self._column_layout_height_factor(zone)
        col_unit = usable_w_pt / num_cols

        # 检查是否存在真正跨列（跨越 >1 列）的区块
        has_spanned = any(
            len(getattr(b, "spanned_cols", [])) > 1 for b in blocks
        )

        if not has_spanned:
            # 无跨列：各列内容并行渲染，高度取最高列
            by_col = defaultdict(list)
            for b in blocks:
                by_col[b.col_index].append(b)
            return self._cap_zone_height(
                max(
                    self._estimate_stream_height_pt(
                        blocks=by_col.get(ci, []),
                        page=page,
                        mapper=mapper,
                        col_width_pt=col_unit,
                        gap_cap_pt=18.0,
                    )
                    for ci in range(num_cols)
                ) * layout_factor,
                blocks, mapper,
            )

        # 有跨列区块：需要按垂直顺序估算
        spanned_set = set()
        for b in blocks:
            if len(getattr(b, "spanned_cols", [])) > 1:
                spanned_set.update(b.spanned_cols)
        standalone_cols = [ci for ci in range(num_cols) if ci not in spanned_set]
        spanned_cols = sorted(spanned_set)
        standalone_max = 0.0
        for ci in standalone_cols:
            col_blocks = [b for b in blocks if b.col_index == ci and len(b.spanned_cols) == 1]
            standalone_max = max(
                standalone_max,
                self._estimate_stream_height_pt(
                    blocks=col_blocks,
                    page=page,
                    mapper=mapper,
                    col_width_pt=col_unit,
                    gap_cap_pt=18.0,
                ),
            )

        span_width = col_unit * len(spanned_cols)
        span_blocks = [b for b in blocks if len(b.spanned_cols) > 1]
        span_top = min((b.bbox.y1 for b in span_blocks), default=float("inf"))
        span_bottom = max((b.bbox.y2 for b in span_blocks), default=float("-inf"))
        merged_height = self._estimate_stream_height_pt(
            blocks=span_blocks,
            page=page,
            mapper=mapper,
            col_width_pt=span_width,
            gap_cap_pt=18.0,
        )

        above_max = 0.0
        below_max = 0.0
        for ci in spanned_cols:
            above_blocks = []
            below_blocks = []
            for b in blocks:
                if b.col_index != ci or len(b.spanned_cols) != 1:
                    continue
                if not span_blocks:
                    above_blocks.append(b)
                    continue
                if float(b.bbox.y2) <= float(span_top) + 8.0:
                    above_blocks.append(b)
                elif float(b.bbox.y1) >= float(span_bottom) - 8.0:
                    below_blocks.append(b)
                elif ((float(b.bbox.y1) + float(b.bbox.y2)) * 0.5) <= ((float(span_top) + float(span_bottom)) * 0.5):
                    above_blocks.append(b)
                else:
                    below_blocks.append(b)
            above_max = max(
                above_max,
                self._estimate_stream_height_pt(
                    blocks=above_blocks,
                    page=page,
                    mapper=mapper,
                    col_width_pt=col_unit,
                    gap_cap_pt=18.0,
                ),
            )
            below_max = max(
                below_max,
                self._estimate_stream_height_pt(
                    blocks=below_blocks,
                    page=page,
                    mapper=mapper,
                    col_width_pt=col_unit,
                    gap_cap_pt=18.0,
                ),
            )
        merged_total = above_max + merged_height + below_max
        est = max(standalone_max, merged_total) * layout_factor
        return self._cap_zone_height(est, blocks, mapper)

    def _cap_zone_height(self, est: float, blocks: list, mapper) -> float:
        """以区块的实际 bbox 垂直范围约束估算高度。

        防止间隙/段落间距累积导致估算过高，从而触发不必要的 fit_scale。

        但源 bbox 是 OCR 紧密边界框（无内边距），而 Word 渲染为每个元素
        增加额外开销（表格单元格间距、图片段落后间距、段落最小行距等）。
        因此乘数需要留有足够余量。
        """
        all_y1 = min((b.bbox.y1 for b in blocks), default=0.0)
        all_y2 = max((b.bbox.y2 for b in blocks), default=0.0)
        bbox_span_pt = mapper.h(max(all_y2 - all_y1, 0.0))
        if bbox_span_pt > 0:
            # 区块越多，累积开销越大；多栏布局乘数更大
            n = len(blocks)
            cap_mult = 1.12 + 0.008 * min(n, 20)  # 1.12 ~ 1.28，随区块数线性增长
            est = min(est, bbox_span_pt * cap_mult)
        return est

    @staticmethod
    def _column_layout_height_factor(zone: Zone) -> float:
        """为复杂多栏区添加保守高度余量。"""
        num_cols = max(int(getattr(zone, "col_count", 1) or 1), 1)
        if num_cols <= 1:
            return 1.0
        factor = 1.0 + 0.02 * max(0, num_cols - 1)  # 2 栏 1.02，3 栏 1.04，...
        if any(len(getattr(block, "spanned_cols", [])) > 1 for block in zone.blocks):
            factor += 0.02
        return factor

    def _estimate_stream_height_pt(
        self,
        blocks: list,
        page: "Page",
        mapper,
        col_width_pt: float,
        gap_cap_pt: float,
    ) -> float:
        if not blocks:
            return 0.0
        total = 0.0
        prev_y = None
        for block in sorted(blocks, key=lambda b: b.bbox.y1):
            if prev_y is not None:
                gap_px = max(0.0, block.bbox.y1 - prev_y)
                if gap_px > 2:
                    total += min(mapper.h(gap_px), gap_cap_pt)
            total += self._estimate_block_height_pt(
                block=block,
                page=page,
                mapper=mapper,
                col_width_pt=col_width_pt,
            )
            prev_y = block.bbox.y2
        return total

    def _estimate_block_height_pt(
        self,
        block: Block,
        page: "Page",
        mapper,
        col_width_pt: float,
    ) -> float:
        raw_h = max(mapper.h(max(block.bbox.height, 0.0)), 0.0)
        if isinstance(block, TextBlock):
            return self._estimate_text_block_height_pt(block, page, mapper, col_width_pt)
        if isinstance(block, TableBlock):
            factor = self._TABLE_HEIGHT_RISK if block.html else self._IMAGE_HEIGHT_RISK
            return raw_h * factor + 1.0
        if isinstance(block, (ImageBlock, EquationBlock)):
            if isinstance(block, EquationBlock) and not block.image_data:
                return self.config.default_font_size_pt * 1.6
            return raw_h * self._IMAGE_HEIGHT_RISK + 5.0
        return raw_h

    def _estimate_text_block_height_pt(
        self,
        block: TextBlock,
        page: "Page",
        mapper,
        col_width_pt: float,
    ) -> float:
        rtype = block.block_type
        page_defaults = getattr(page, "style_defaults", None)
        config_defaults = self._config_style_defaults()
        block_effective = resolve_textline_style(
            page_defaults=page_defaults,
            block_style=block.style,
            line_style=None,
            config_defaults=config_defaults,
        )
        est_h = self._estimate_text_content_height_pt(
            block=block,
            page=page,
            block_effective=block_effective,
            col_width_pt=col_width_pt,
        )
        bbox_h = mapper.h(max(block.bbox.height, 0.0))
        return max(est_h, bbox_h * 0.95)

    def _estimate_text_content_height_pt(
        self,
        block: TextBlock,
        page: "Page",
        block_effective: dict,
        col_width_pt: float,
    ) -> float:
        rtype = block.block_type
        mapper = page.coord_mapper

        text = block.full_text().strip()
        para_count = 1
        is_caption = rtype in _CAPTION_TYPES
        if block.paragraphs and len(block.paragraphs) > 1 and not is_caption:
            para_count = sum(1 for p in block.paragraphs if ''.join(ln.text.strip() for ln in p.lines).strip())
            para_count = max(1, para_count)

        font_size = float(block_effective.get("font_size_pt") or block.estimated_font_size_pt or self.config.default_font_size_pt)
        if rtype == BlockType.TITLE:
            font_size = self._resolve_title_font_size_pt(
                block=block,
                page=page,
                font_size_pt=font_size,
                alignment=str(block_effective.get("alignment") or "left"),
            )
        elif is_caption:
            font_size = max(font_size - 0.5, 6.0)
        elif rtype in (BlockType.HEADER, BlockType.FOOTER, BlockType.REFERENCE):
            font_size = max(font_size - 1.0, 6.0)
        else:
            font_size = self._resolve_body_font_size_pt(block, page, font_size)

        line_spacing = float(block_effective.get("line_spacing") or self.config.default_line_spacing)
        line_spacing = max(1.1, min(line_spacing, 2.2))
        line_height = max(font_size * line_spacing, font_size * 1.05)
        space_after = float(block_effective.get("space_after_pt", 1.0) or 1.0)

        base_lines = max(block.count_lines(), 1)
        source_width = max(mapper.w(max(block.bbox.width, 1.0)), 1.0)
        target_width = max(col_width_pt * 0.96, 1.0)
        wrap_ratio = self._effective_text_wrap_ratio(
            text=text,
            base_lines=base_lines,
            source_width=source_width,
            target_width=target_width,
            is_caption=is_caption,
        )

        est_lines = max(base_lines, int(math.ceil(base_lines * wrap_ratio)))
        return est_lines * line_height + para_count * space_after

    def _effective_text_wrap_ratio(
        self,
        text: str,
        base_lines: int,
        source_width: float,
        target_width: float,
        is_caption: bool,
    ) -> float:
        """估算目标容器相对源容器的换行放大倍数。"""
        if is_caption or base_lines <= 1:
            return 1.0

        width_ratio = max(1.0, source_width / max(target_width, 1.0))
        # 目标列宽未明显变窄时，保留 OCR 原始行数，避免为密集 CJK 版面过度缩放。
        if width_ratio <= 1.01:
            return 1.0

        wrap_risk = self._text_wrap_risk(text)
        if width_ratio > 1.25:
            wrap_risk += 0.02
        if len(text) / max(base_lines, 1) > 45:
            wrap_risk += 0.02
        return max(1.0, width_ratio * wrap_risk)

    @staticmethod
    def _visual_row_count(block: TextBlock) -> int:
        centers: List[float] = []
        heights: List[float] = []
        for line in block.lines or []:
            if line.y1 is None or line.y2 is None:
                continue
            y1 = float(line.y1)
            y2 = float(line.y2)
            centers.append((y1 + y2) * 0.5)
            heights.append(max(y2 - y1, 1.0))
        if not centers:
            return block.count_lines()

        heights_sorted = sorted(heights)
        median_h = heights_sorted[len(heights_sorted) // 2]
        threshold = max(median_h * 0.55, 2.0)
        rows: List[float] = []
        row_counts: List[int] = []
        for center in sorted(centers):
            if not rows or abs(center - rows[-1]) > threshold:
                rows.append(center)
                row_counts.append(1)
                continue
            count = row_counts[-1]
            rows[-1] = (rows[-1] * count + center) / (count + 1)
            row_counts[-1] = count + 1
        return max(1, len(rows))

    @classmethod
    def _is_wide_centered_single_line_title(cls, block: TextBlock, page: "Page", alignment: str) -> bool:
        if block.block_type != BlockType.TITLE:
            return False
        if (alignment or "").lower() != "center":
            return False
        if cls._visual_row_count(block) != 1:
            return False
        return float(block.bbox.width) >= max(float(page.image_width) * 0.55, 1.0)

    @classmethod
    def _is_masthead_title(cls, block: TextBlock, page: "Page", alignment: str) -> bool:
        text = (block.full_text() or "").strip()
        aspect = float(block.bbox.width) / max(float(block.bbox.height), 1.0)
        return (
            cls._is_wide_centered_single_line_title(block, page, alignment)
            and aspect >= 2.8
            and cls._cjk_ratio(text) >= 0.20
            and float(block.bbox.y1) <= max(float(page.image_height) * 0.12, 1.0)
        )

    @staticmethod
    def _title_level(block: TextBlock) -> Optional[int]:
        text = (block.full_text() or "").strip()
        match = _NUMBERED_TITLE_LEVEL_RE.match(text)
        if not match:
            return None
        numeric = match.group(1)
        if numeric:
            return numeric.count(".") + 1
        return 1

    @staticmethod
    def _is_numbered_section_title(block: TextBlock) -> bool:
        return DocxRenderer._title_level(block) is not None

    def _resolve_title_font_size_pt(
        self,
        block: TextBlock,
        page: "Page",
        font_size_pt: float,
        alignment: str,
    ) -> float:
        cfg = self.config
        title_level = self._title_level(block)
        text = (block.full_text() or "").strip()
        aspect = float(block.bbox.width) / max(float(block.bbox.height), 1.0)
        if len(text) >= 18 and aspect < 4.0:
            return min(font_size_pt, self.config.default_font_size_pt * 1.15)
        if self._is_masthead_title(block, page, alignment):
            layout_profile = str((getattr(page, "attributes", None) or {}).get("layout_profile", "") or "")
            if layout_profile == "newspaper_mixed":
                return min(font_size_pt * max(cfg.title_masthead_scale, 1.55), 58.0)
            return min(font_size_pt * cfg.title_masthead_scale, cfg.title_masthead_cap)
        if title_level is not None:
            return font_size_pt
        if self._is_wide_centered_single_line_title(block, page, alignment):
            return min(font_size_pt * cfg.title_wide_centered_scale, cfg.title_wide_centered_cap)
        if self._visual_row_count(block) >= 2:
            return min(font_size_pt * 1.06, cfg.title_default_cap)
        return min(font_size_pt * cfg.title_default_scale, cfg.title_default_cap)

    @staticmethod
    def _cjk_ratio(text: str) -> float:
        if not text:
            return 0.0
        cjk = sum(1 for ch in text if '\u4e00' <= ch <= '\u9fff')
        return cjk / max(len(text), 1)

    @classmethod
    def _visual_fragment_gap_text(cls, prev_line, curr_line, block: TextBlock) -> str:
        """Return source-geometry spaces between OCR fragments on the same row."""
        if prev_line is None or curr_line is None:
            return ""
        if block.block_type not in {
            BlockType.TITLE,
            BlockType.HEADER,
            BlockType.FOOTER,
            BlockType.PAGE_NUMBER,
            BlockType.TABLE_CAPTION,
            BlockType.FIGURE_CAPTION,
            BlockType.FORMULA_CAPTION,
            BlockType.TABLE_FOOTNOTE,
        }:
            return ""
        if (
            getattr(prev_line, "x1", None) is None
            or getattr(prev_line, "x2", None) is None
            or getattr(curr_line, "x1", None) is None
            or getattr(curr_line, "x2", None) is None
        ):
            return ""

        gap_px = float(curr_line.x1) - float(prev_line.x2)
        if gap_px <= 0:
            return ""

        widths: List[float] = []
        for line in (prev_line, curr_line):
            line_text = (getattr(line, "text", "") or "").strip()
            char_count = max(len(line_text), 1)
            width = max(float(line.x2) - float(line.x1), 0.0)
            if width > 0:
                widths.append(width / char_count)
        if not widths:
            return ""

        widths.sort()
        avg_char_px = widths[len(widths) // 2]
        threshold = avg_char_px * (0.35 if block.block_type == BlockType.TITLE else 0.55)
        if gap_px < max(threshold, 2.0):
            return ""

        combined = f"{getattr(prev_line, 'text', '')}{getattr(curr_line, 'text', '')}"
        if cls._cjk_ratio(combined) >= 0.30:
            count = max(1, min(6, int(round(gap_px / max(avg_char_px, 1.0)))))
            return "\u3000" * count

        count = max(1, min(12, int(round(gap_px / max(avg_char_px * 0.45, 1.0)))))
        return " " * count

    @staticmethod
    def _should_render_visual_rows_as_paragraphs(
        block: TextBlock,
        visual_rows: List[List[object]],
    ) -> bool:
        if len(visual_rows) <= 1:
            return False
        if block.block_type not in {
            BlockType.FOOTNOTE,
            BlockType.FIGURE_CAPTION,
            BlockType.TABLE_CAPTION,
            BlockType.TABLE_FOOTNOTE,
            BlockType.FORMULA_CAPTION,
        }:
            return False
        text_len = len((block.full_text() or "").strip())
        return len(visual_rows) <= 4 and text_len <= 160

    @staticmethod
    def _alignment_for_visual_row(row: List[object], block: TextBlock, fallback):
        xs = [
            float(getattr(ln, "x1"))
            for ln in row
            if getattr(ln, "x1", None) is not None
        ]
        xe = [
            float(getattr(ln, "x2"))
            for ln in row
            if getattr(ln, "x2", None) is not None
        ]
        if not xs or not xe:
            return fallback

        row_left = min(xs)
        row_right = max(xe)
        row_center = (row_left + row_right) * 0.5
        block_left = float(block.bbox.x1)
        block_right = float(block.bbox.x2)
        block_width = max(block_right - block_left, 1.0)
        block_center = (block_left + block_right) * 0.5
        left_gap = max(0.0, row_left - block_left)
        right_gap = max(0.0, block_right - row_right)

        if right_gap <= block_width * 0.04 and left_gap >= block_width * 0.18:
            return WD_ALIGN_PARAGRAPH.RIGHT
        if abs(row_center - block_center) <= block_width * 0.08:
            return WD_ALIGN_PARAGRAPH.CENTER
        if left_gap <= block_width * 0.04 and right_gap >= block_width * 0.18:
            return WD_ALIGN_PARAGRAPH.LEFT
        return fallback

    def _resolve_body_font_size_pt(
        self,
        block: TextBlock,
        page: "Page",
        font_size_pt: float,
    ) -> float:
        """Use geometric font size, with page-level guards for short body blocks."""
        page_body = self._page_body_font_size_pt(page)
        if page_body is None:
            return font_size_pt

        text = (block.full_text() or "").strip()
        is_short = block.count_lines() <= 2 or len(text) <= 120
        is_listish = any(self._is_listish_line(getattr(line, "text", "")) for line in block.lines or [])
        if is_short or is_listish:
            return min(font_size_pt, page_body * 1.12)
        if block.block_type == BlockType.TEXT and font_size_pt > page_body * 1.28:
            return min(font_size_pt, page_body * 1.18)
        return font_size_pt

    @staticmethod
    def _looks_like_question_or_option_block(block: TextBlock) -> bool:
        lines = [(getattr(line, "text", "") or "").strip() for line in block.lines or []]
        lines = [line for line in lines if line]
        if not lines:
            return False
        hits = sum(1 for line in lines if _QUESTION_OR_OPTION_RE.match(line))
        if hits >= 2:
            return True
        text = re.sub(r"\s+", " ", block.full_text() or "").strip()
        option_hits = len(re.findall(r"(?:^|\s)[A-H][\.\)]\s*\S", text))
        numbered_prompt = bool(re.match(r"^\s*\d{1,2}\s*[A-Z][a-z]", text))
        return numbered_prompt and option_hits >= 2

    @staticmethod
    def _page_body_font_size_pt(page: "Page") -> Optional[float]:
        values: List[float] = []
        for zone in getattr(page, "zones", []) or []:
            for block in getattr(zone, "blocks", []) or []:
                if not isinstance(block, TextBlock) or block.block_type != BlockType.TEXT:
                    continue
                style = getattr(block, "style", None)
                size = getattr(style, "font_size_pt", None) if style is not None else None
                if size is None:
                    continue
                text_len = len((block.full_text() or "").strip())
                if text_len < 24 and block.count_lines() <= 2:
                    continue
                width_ratio = float(block.bbox.width) / max(float(getattr(page, "image_width", 0) or 0), 1.0)
                if width_ratio <= 0.06:
                    continue
                values.extend([float(size)] * max(1, min(5, text_len // 80 + 1)))
        if not values:
            return None
        values.sort()
        return values[len(values) // 2]

    def _text_wrap_risk(self, text: str) -> float:
        if not text:
            return self._TEXT_WRAP_RISK_LATIN
        cjk = sum(1 for ch in text if '\u4e00' <= ch <= '\u9fff')
        ratio = cjk / max(len(text), 1)
        return self._TEXT_WRAP_RISK_CJK if ratio >= 0.5 else self._TEXT_WRAP_RISK_LATIN

    # ------------------------------------------------------------------
    # 区块渲染
    # ------------------------------------------------------------------

    def _render_block(self, container, block: Block, ctx: RenderContext,
                      space_before: float = 0) -> Tuple[Optional[object], bool]:
        """渲染单个区块。返回 (last_paragraph, ended_sentence)。"""
        if isinstance(block, ImageBlock):
            return self._render_image_block(container, block, ctx, space_before), True
        elif isinstance(block, EquationBlock):
            return self._render_equation_block(container, block, ctx, space_before), True
        elif isinstance(block, TableBlock):
            self._render_table_block(container, block, ctx, space_before)
            return None, True
        elif isinstance(block, TextBlock):
            return self._render_text_block(container, block, ctx, space_before)
        else:
            # 未知区块类型：尝试作为图片处理
            if hasattr(block, 'image_data') and block.image_data:
                return self._render_image_block(container, block, ctx, space_before), True
            return None, True

    def _render_image_block(self, container, block, ctx: RenderContext,
                            space_before: float) -> Optional[object]:
        """将图片/图像区块渲染为居中内联图片。"""
        image_data = getattr(block, 'image_data', None)
        if not image_data:
            return None

        pw = self._visual_block_width_pt(block, ctx, apply_fit_scale=False)

        if space_before > self._scale(3):
            add_spacing_para(container, min(space_before, self._scale(12)))

        p = container.add_paragraph()
        reset_paragraph_format(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run()
        try:
            run.add_picture(io.BytesIO(image_data), width=Pt(pw))
        except Exception:
            pass
        set_paragraph_spacing(p, space_after=self._scale(3))
        return p

    def _visual_block_width_pt(
        self,
        block: Block,
        ctx: RenderContext,
        *,
        apply_fit_scale: bool,
        max_ratio: float = 0.98,
        min_width_pt: float = 12.0,
    ) -> float:
        """Map a visual block's source bbox width into the current DOCX container."""
        bbox = block.bbox
        local_col_px = max(float(ctx.col_right_px) - float(ctx.col_left_px), 1.0)
        source_ratio = max(float(bbox.width), 1.0) / local_col_px
        width_pt = ctx.col_width_pt * min(max(source_ratio, 0.0), max_ratio)

        mapper_width = ctx.coord_mapper.w(max(float(bbox.width), 1.0))
        if not getattr(ctx, "in_table_cell", False):
            width_pt = min(width_pt, mapper_width, ctx.col_width_pt * max_ratio)
        else:
            width_pt = min(width_pt, ctx.col_width_pt * max_ratio)
            if (
                block.block_type in {BlockType.FIGURE, BlockType.TABLE}
                and float(getattr(ctx, "span_gap_pt", 0.0) or 0.0) > 0
                and float(getattr(ctx.page, "usable_width_pt", 0.0) or 0.0) > 0
            ):
                page_ratio_width = (
                    float(ctx.page.usable_width_pt)
                    * max(float(bbox.width), 1.0)
                    / max(float(ctx.page.image_width), 1.0)
                )
                width_pt = min(
                    max(width_pt, page_ratio_width),
                    ctx.col_width_pt + float(ctx.span_gap_pt),
                    float(ctx.page.usable_width_pt) * max_ratio,
                )

        effective_max_ratio = max_ratio
        if (
            bool(getattr(ctx, "local_visual_band", False))
            and block.block_type in {BlockType.FIGURE, BlockType.TABLE}
        ):
            effective_max_ratio = min(effective_max_ratio, 0.9)

        cap_width_pt = ctx.col_width_pt * effective_max_ratio
        if (
            getattr(ctx, "in_table_cell", False)
            and block.block_type in {BlockType.FIGURE, BlockType.TABLE}
            and float(getattr(ctx, "span_gap_pt", 0.0) or 0.0) > 0
        ):
            cap_width_pt = min(
                ctx.col_width_pt + float(ctx.span_gap_pt),
                float(getattr(ctx.page, "usable_width_pt", ctx.col_width_pt) or ctx.col_width_pt) * max_ratio,
            )

        if float(bbox.height) > 0 and float(bbox.width) > 0:
            src_aspect = float(bbox.width) / max(float(bbox.height), 1.0)
            if src_aspect > 1.8:
                height_width_pt = ctx.coord_mapper.h(float(bbox.height)) * src_aspect
                width_pt = min(width_pt, height_width_pt)

        width_pt = max(min_width_pt, min(width_pt, cap_width_pt))
        return self._scale(width_pt) if apply_fit_scale else width_pt

    def _render_equation_block(self, container, block: EquationBlock,
                               ctx: RenderContext, space_before: float) -> Optional[object]:
        """渲染公式：优先使用图片；若携带编号，则同段右对齐显示。"""
        image_data = block.image_data
        formula_number = str((getattr(block, "attributes", None) or {}).get("formula_number_text", "") or "").strip()
        raw_label = str((getattr(block, "attributes", None) or {}).get("raw_layout_label", "") or "")
        if raw_label == "formula_number" and formula_number and _FORMULA_NUMBER_TEXT_RE.match(formula_number):
            if space_before > self._scale(3):
                add_spacing_para(container, min(space_before, self._scale(8)))
            p = container.add_paragraph()
            reset_paragraph_format(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(formula_number)
            set_run_font(
                run,
                font_size=self._scale_font(self.config.default_font_size_pt),
                east_asia=self.config.default_cjk_font,
                font_name=self.config.default_font,
            )
            set_paragraph_spacing(p, space_after=self._scale(2))
            return p
        if image_data and formula_number:
            if space_before > self._scale(3):
                add_spacing_para(container, min(space_before, self._scale(12)))
            p = container.add_paragraph()
            reset_paragraph_format(p)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            tab_stop = p.paragraph_format.tab_stops.add_tab_stop(
                Pt(max(ctx.col_width_pt - self._scale(6), 24.0)),
                alignment=WD_TAB_ALIGNMENT.RIGHT,
                leader=WD_TAB_LEADER.SPACES,
            )
            del tab_stop
            run = p.add_run()
            try:
                formula_width = self._visual_block_width_pt(
                    block,
                    ctx,
                    apply_fit_scale=True,
                    max_ratio=0.72,
                    min_width_pt=18.0,
                )
                run.add_picture(io.BytesIO(image_data), width=Pt(formula_width))
            except Exception:
                pass
            p.add_run("\t")
            num_run = p.add_run(formula_number)
            set_run_font(
                num_run,
                font_size=self._scale_font(self.config.default_font_size_pt),
                east_asia=self.config.default_cjk_font,
                font_name=self.config.default_font,
            )
            set_paragraph_spacing(p, space_after=self._scale(3))
            return p
        if image_data:
            return self._render_image_block(container, block, ctx, space_before)
        # 若有 LaTeX 但无图片，则以文本方式渲染
        if block.latex:
            p = container.add_paragraph()
            reset_paragraph_format(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(block.latex)
            set_run_font(run, font_size=self._scale_font(self.config.default_font_size_pt), italic=True)
            return p
        # 无图片、无 LaTeX：插入占位符
        p = container.add_paragraph()
        reset_paragraph_format(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[公式]")
        from docx.shared import RGBColor
        set_run_font(run, font_size=self._scale_font(self.config.default_font_size_pt), italic=True,
                     color_rgb=RGBColor(160, 160, 160))
        return p

    def _render_table_block(self, container, block: TableBlock,
                            ctx: RenderContext, space_before: float) -> None:
        """渲染表格：优先使用 HTML 结构，其次回退到图片。"""
        if space_before > self._scale(MIN_LINE_SPACING_PT):
            add_spacing_para(container, space_before)

        rendered = False
        if block.html:
            try:
                parser = HtmlToDocx()
                parser.table_style = "Table Grid"
                parser.handle_table(
                    block.html,
                    container,
                    table_image_data=block.image_data,
                    table_width_pt=ctx.col_width_pt,
                )
                if hasattr(container, 'tables') and container.tables:
                    t = container.tables[-1]
                    t.alignment = WD_TABLE_ALIGNMENT.CENTER
                    # 缩放列宽，防止超出可用列宽
                    fit_table_to_width(t, ctx.col_width_pt)
                    for row in t.rows:
                        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                rendered = True
            except Exception as e:
                logger.warning("Table HTML rendering failed, falling back to plain text: %s", e)

        if not rendered and block.image_data:
            self._render_image_block(container, block, ctx, 0.0)
            rendered = True

        if not rendered and block.html:
            plain = re.sub(r"<[^>]+>", " ", block.html)
            plain = re.sub(r"\s+", " ", plain).strip()
            if plain:
                p = container.add_paragraph()
                reset_paragraph_format(p)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(plain)
                set_run_font(run, font_size=self._scale(self.config.default_font_size_pt))
                rendered = True

        if not rendered:
            p = container.add_paragraph()
            reset_paragraph_format(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("[表格]")
            set_run_font(run, font_size=self._scale(self.config.default_font_size_pt))

    def _render_text_block(self, container, block: TextBlock,
                           ctx: RenderContext,
                           space_before: float = 0) -> Tuple[Optional[object], bool]:
        """渲染文本/标题/说明区块。

        完全由 ``block.style`` 驱动，不做任何样式推断。
        所有样式字段均已由管线的 :mod:`docflow.layout.style_inferrer` 填充。
        """
        rtype = block.block_type
        bs = block.style          # style_inferrer 保证非 None
        page_defaults = getattr(ctx.page, "style_defaults", None)
        config_defaults = self._config_style_defaults()
        block_effective = resolve_textline_style(
            page_defaults=page_defaults,
            block_style=bs,
            line_style=None,
            config_defaults=config_defaults,
        )
        is_title = rtype == BlockType.TITLE
        is_caption = rtype in _CAPTION_TYPES
        preserve_line_breaks = bool(
            getattr(self.config, "docx_preserve_visual_line_breaks", True)
        )

        # ── 文本内容 ─────────────────────────────────────────────────
        text = block.full_text().strip()
        if not text:
            return None, True

        mapper = ctx.coord_mapper
        layout_profile = str((getattr(ctx.page, "attributes", None) or {}).get("layout_profile", "") or "")
        compact_multicol_text = (
            layout_profile in {"newspaper_mixed", "generic_complex"}
            and int(getattr(block, "col_count", 1) or 1) >= 3
            and rtype == BlockType.TEXT
        )

        # ── 从 block.style 读取全部样式（无推断逻辑） ────────────────
        alignment = _ALIGN_MAP.get(
            str(block_effective.get("alignment", "left")).lower(),
            WD_ALIGN_PARAGRAPH.LEFT,
        )
        sp_after = self._scale(max(float(block_effective.get("space_after_pt", 1.0) or 1.0) - self._corr_space_after_pt, 0))
        if compact_multicol_text:
            sp_after = 0.0
        line_spacing = block_effective.get("line_spacing")
        if line_spacing is not None:
            min_line_spacing = 1.0 if compact_multicol_text else 1.1
            line_spacing = max(min_line_spacing, float(line_spacing))
        preserve_breaks_on_ambiguous_justify = bool(
            getattr(self.config, "docx_preserve_breaks_on_ambiguous_justify", True)
        )
        ambiguous_justify = (
            alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
            and self._is_ambiguous_justify(block, ctx.col_left_px, ctx.col_right_px)
        )
        if (
            alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
            and self._should_downgrade_justify_to_left(block, ctx)
        ):
            alignment = WD_ALIGN_PARAGRAPH.LEFT
            ambiguous_justify = True
        if (
            alignment == WD_ALIGN_PARAGRAPH.CENTER
            and rtype == BlockType.TEXT
            and self._looks_like_question_or_option_block(block)
        ):
            alignment = WD_ALIGN_PARAGRAPH.LEFT
        local_visual_band = bool(getattr(ctx, "local_visual_band", False))
        if local_visual_band and is_title:
            alignment = self._local_visual_title_alignment(block, ctx, alignment)
        elif local_visual_band and rtype in {BlockType.TEXT, BlockType.ABSTRACT, BlockType.REFERENCE}:
            alignment = self._local_visual_text_alignment(block, ctx, alignment)
        bbox_h_pt = mapper.h(max(block.bbox.height, 0.0))
        content_h_pt = self._estimate_text_content_height_pt(
            block=block,
            page=ctx.page,
            block_effective=block_effective,
            col_width_pt=ctx.col_width_pt,
        )
        reserve_bbox_after_pt = 0.0
        if (
            rtype == BlockType.TEXT
            and bbox_h_pt > content_h_pt * 1.12
            and self._cjk_ratio(text) >= 0.55
            and float(block.bbox.width) / max(float(ctx.page.image_width), 1.0) >= 0.62
        ):
            reserve_bbox_after_pt = min(max(0.0, bbox_h_pt * 0.95 - content_h_pt), 28.0)
        if (
            rtype == BlockType.TEXT
            and float(block.bbox.width) <= max(float(ctx.page.image_width) * 0.12, 140.0)
            and block.count_lines() <= 3
            and len(text) <= 24
        ):
            alignment = WD_ALIGN_PARAGRAPH.CENTER if float(block.bbox.width) <= max(float(ctx.page.image_width) * 0.08, 110.0) else WD_ALIGN_PARAGRAPH.LEFT
            line_spacing = max(1.0, float(line_spacing or 1.0))
        # ── 段落渲染内部函数 ─────────────────────────────────────────
        def _line_effective_style(line) -> dict:
            return resolve_textline_style(
                page_defaults=page_defaults,
                block_style=bs,
                line_style=getattr(line, "style", None),
                config_defaults=config_defaults,
            )

        def _write_one_run(p, run_text: str, run_style: dict) -> None:
            if not run_text:
                return
            color = self._parse_css_hex(run_style.get("color"))
            font_size = self._scale_font(float(run_style.get("font_size_pt") or self.config.default_font_size_pt))
            bold = bool(run_style.get("bold", False))
            italic = bool(run_style.get("italic", False))
            underline = bool(run_style.get("underline", False))
            strike = bool(run_style.get("strikethrough", False))
            superscript = bool(run_style.get("superscript", False))
            subscript = bool(run_style.get("subscript", False))
            east_asia = str(run_style.get("font_family") or self.config.default_cjk_font)
            font_name = str(run_style.get("font_family_western") or self.config.default_font)

            run = p.add_run(run_text)
            if is_title:
                title_alignment = self._paragraph_alignment_name(alignment)
                title_font_size = self._resolve_title_font_size_pt(
                    block=block,
                    page=ctx.page,
                    font_size_pt=float(run_style.get("font_size_pt") or self.config.default_font_size_pt),
                    alignment=title_alignment,
                )
                title_east_asia = self.config.title_font
                if bs is not None and getattr(bs, "font_family", None):
                    title_east_asia = east_asia
                set_run_font(
                    run,
                    font_size=self._scale_title_font(block, ctx.page, title_font_size),
                    bold=bold,
                    italic=italic,
                    underline=underline,
                    strikethrough=strike,
                    superscript=superscript,
                    subscript=subscript,
                    east_asia=title_east_asia,
                    font_name=font_name,
                    color_rgb=color,
                )
                return
            if is_caption:
                font_size = max(font_size - 0.5, 6.0)
            elif rtype in (BlockType.HEADER, BlockType.FOOTER, BlockType.REFERENCE):
                font_size = max(font_size - 1.0, 6.0)
            elif (
                rtype == BlockType.TEXT
                and float(block.bbox.width) <= max(float(ctx.page.image_width) * 0.12, 140.0)
                and len(text) <= 24
            ):
                font_size = min(font_size, self._scale_font(self.config.default_font_size_pt * 1.12))
            else:
                body_font_size = self._resolve_body_font_size_pt(
                    block=block,
                    page=ctx.page,
                    font_size_pt=float(run_style.get("font_size_pt") or self.config.default_font_size_pt),
                )
                font_size = self._scale_font(body_font_size)
                if (
                    bool(getattr(ctx, "local_visual_band", False))
                    and rtype == BlockType.TEXT
                    and block.count_lines() <= 2
                    and len(text) <= 120
                ):
                    font_size = min(font_size, self._scale_font(self.config.default_font_size_pt * 1.05))

            set_run_font(
                run,
                font_size=font_size,
                bold=bold,
                italic=italic,
                underline=underline,
                strikethrough=strike,
                superscript=superscript,
                subscript=subscript,
                east_asia=east_asia,
                font_name=font_name,
                color_rgb=color,
            )

        def _write_runs(p, txt: str, para_lines=None) -> None:
            if para_lines:
                clean_lines = [ln for ln in para_lines if ln.text and ln.text.strip()]
                if not clean_lines:
                    return
                visual_rows = self._group_lines_into_visual_rows(clean_lines)
                keep_visual_breaks = self._should_keep_visual_breaks(
                    block=block,
                    alignment=alignment,
                    preserve_line_breaks=preserve_line_breaks,
                    preserve_breaks_on_ambiguous_justify=preserve_breaks_on_ambiguous_justify,
                    ambiguous_justify=ambiguous_justify,
                    visual_rows=visual_rows,
                    render_mode=str(getattr(ctx, "render_mode", "") or ""),
                    in_table_cell=bool(getattr(ctx, "in_table_cell", False)),
                    force_table_breaks=bool(getattr(ctx, "preserve_visual_breaks_in_table", False)),
                    local_visual_band=local_visual_band,
                )
                if compact_multicol_text:
                    keep_visual_breaks = False
                prev_row_text = ""
                for ri, row in enumerate(visual_rows):
                    prev_text = ""
                    prev_line = None
                    for ln in row:
                        curr_text = ln.text.strip()
                        if not curr_text:
                            continue
                        if (
                            not keep_visual_breaks
                            and not prev_text
                            and prev_row_text
                            and should_insert_space(prev_row_text, curr_text)
                        ):
                            p.add_run(" ")
                        if prev_text:
                            gap_text = self._visual_fragment_gap_text(prev_line, ln, block)
                            if gap_text:
                                p.add_run(gap_text)
                            elif should_insert_space(prev_text, curr_text):
                                p.add_run(" ")
                        _write_one_run(p, curr_text, _line_effective_style(ln))
                        prev_text = curr_text
                        prev_line = ln
                        prev_row_text = curr_text
                    if ri < len(visual_rows) - 1 and keep_visual_breaks:
                        p.add_run().add_break()
                return
            _write_one_run(p, txt, block_effective)

        def _make_para(txt: str, indent_pt: float, is_first: bool, para_lines=None) -> object:
            p = container.add_paragraph()
            reset_paragraph_format(p)
            set_paragraph_spacing(p,
                                  space_before=space_before if is_first else 0,
                                  space_after=sp_after)
            p.alignment = alignment
            if indent_pt > 0 and not is_title and not is_caption:
                p.paragraph_format.first_line_indent = Pt(indent_pt)
            if is_title:
                p.paragraph_format.keep_with_next = True
            if line_spacing is not None:
                p.paragraph_format.line_spacing = float(line_spacing)
            _write_runs(p, txt, para_lines=para_lines)
            return p

        def _make_visual_row_para(row: List[object], is_first: bool, is_last: bool) -> object:
            row_text = "".join((getattr(ln, "text", "") or "").strip() for ln in row).strip()
            p = container.add_paragraph()
            reset_paragraph_format(p)
            set_paragraph_spacing(
                p,
                space_before=space_before if is_first else 0,
                space_after=sp_after if is_last else 0,
            )
            p.alignment = self._alignment_for_visual_row(row, block, alignment)
            if line_spacing is not None:
                p.paragraph_format.line_spacing = float(line_spacing)
            _write_runs(p, row_text, para_lines=row)
            return p

        # ── 多段落 vs 单段落渲染 ─────────────────────────────────────
        # 多段落时：block.paragraphs 已由 pipeline 设置，
        # 各段的 first_line_indent_px 已由 style_inferrer 填充
        # Caption 类块始终走单段落路径，以便用软换行保留行间结构
        last_p = None
        visual_rows = self._group_lines_into_visual_rows(block.lines or [])
        if self._should_render_visual_rows_as_paragraphs(block, visual_rows):
            for ri, row in enumerate(visual_rows):
                last_p = _make_visual_row_para(
                    row,
                    is_first=(ri == 0),
                    is_last=(ri == len(visual_rows) - 1),
                )
        elif block.paragraphs and len(block.paragraphs) > 1 and not is_caption:
            for pi, para in enumerate(block.paragraphs):
                para_text = para.text.strip()
                if not para_text.strip():
                    continue
                indent_pt = 0.0 if is_title else (
                    round(mapper.w(para.first_line_indent_px) * 2) / 2.0
                    if para.first_line_indent_px > 0 else 0.0
                )
                indent_pt = self._scale(indent_pt)
                last_p = _make_para(
                    para_text,
                    indent_pt,
                    is_first=(pi == 0),
                    para_lines=para.lines,
                )
        else:
            indent_pt = 0.0 if is_title else self._scale((bs.first_line_indent_pt or 0.0) if bs else 0.0)
            para_lines = block.lines if block.lines else None
            last_p = _make_para(text, indent_pt, is_first=True, para_lines=para_lines)

        if last_p is not None and reserve_bbox_after_pt > 0:
            set_paragraph_spacing(last_p, space_after=sp_after + self._scale(reserve_bbox_after_pt))

        ended = bool(text) and text[-1] in '。！？…；.!?'
        return last_p, is_title or ended

    @staticmethod
    def _line_y_range(line) -> Optional[Tuple[float, float]]:
        region = getattr(line, "text_region", None)
        if not region:
            return None
        ys = [
            float(pt[1])
            for pt in region
            if isinstance(pt, (list, tuple)) and len(pt) >= 2
        ]
        if not ys:
            return None
        return min(ys), max(ys)

    @classmethod
    def _group_lines_into_visual_rows(cls, lines: List[object]) -> List[List[object]]:
        """按 y 轴重叠将 OCR 文本片段聚合为视觉行。"""
        if not lines:
            return []

        y_ranges = [cls._line_y_range(ln) for ln in lines]
        heights: List[float] = []
        for yr in y_ranges:
            if yr is None:
                continue
            y1, y2 = yr
            heights.append(max(y2 - y1, 0.0))
        avg_h = (sum(heights) / len(heights)) if heights else 0.0
        overlap_threshold = max(avg_h * 0.35, 1.0) if avg_h > 0 else None

        rows: List[List[object]] = []
        current_row: List[object] = []
        current_range: Optional[Tuple[float, float]] = None

        for line, y_range in zip(lines, y_ranges):
            if not current_row:
                current_row = [line]
                current_range = y_range
                continue

            same_row = False
            if current_range is not None and y_range is not None and overlap_threshold is not None:
                overlap = min(current_range[1], y_range[1]) - max(current_range[0], y_range[0])
                row_center = (current_range[0] + current_range[1]) * 0.5
                line_center = (y_range[0] + y_range[1]) * 0.5
                center_delta = abs(line_center - row_center)
                center_threshold = max(avg_h * 0.48, 2.0) if avg_h > 0 else 2.0
                if overlap > overlap_threshold and center_delta <= center_threshold:
                    same_row = True

            if same_row:
                current_row.append(line)
                if current_range is not None and y_range is not None:
                    current_range = (
                        min(current_range[0], y_range[0]),
                        max(current_range[1], y_range[1]),
                    )
            else:
                rows.append(current_row)
                current_row = [line]
                current_range = y_range

        if current_row:
            rows.append(current_row)

        for row in rows:
            row.sort(
                key=lambda ln: (
                    float("inf")
                    if getattr(ln, "x1", None) is None
                    else float(getattr(ln, "x1"))
                )
            )
        return rows

    def _is_ambiguous_justify(self, block: TextBlock, col_left: float, col_right: float) -> bool:
        """判断当前 justify 块是否缺乏足够几何证据。"""
        edges: List[Tuple[float, float]] = []
        for ln in block.lines:
            if ln.x1 is None or ln.x2 is None:
                continue
            edges.append((float(ln.x1), float(ln.x2)))
        if not edges:
            return True

        col_w = max(col_right - col_left, 1.0)
        thresh = col_w * 0.10
        min_lines = int(max(1, getattr(self.config, "align_justify_min_lines", 3)))

        left_hit = sum(1 for x1, _ in edges if abs(x1 - col_left) <= thresh) / len(edges)
        body = edges[:-1] if len(edges) >= 3 else edges
        right_hit = sum(1 for _, x2 in body if abs(col_right - x2) <= thresh) / max(len(body), 1)
        ragged = sum(1 for _, x2 in body if (col_right - x2) > thresh) / max(len(body), 1)

        if len(edges) < min_lines:
            return not (left_hit >= 0.85 and right_hit >= 0.85 and ragged <= 0.25)
        return left_hit < 0.70 or right_hit < 0.70 or ragged > 0.40

    def _should_downgrade_justify_to_left(self, block: TextBlock, ctx: RenderContext) -> bool:
        if block.block_type != BlockType.TEXT:
            return False
        text = block.full_text().strip()
        if not text or self._cjk_ratio(text) >= 0.20:
            return False
        page_w = max(float(getattr(ctx.page, "image_width", 0) or 0), 1.0)
        width_ratio = float(block.bbox.width) / page_w
        if width_ratio > 0.38:
            return False
        if any(self._is_listish_line(getattr(line, "text", "")) for line in getattr(block, "lines", []) or []):
            return False
        return block.count_lines() >= 3 or len(text) >= 80

    @staticmethod
    def _paragraph_alignment_name(alignment) -> str:
        if alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return "center"
        if alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            return "right"
        if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            return "justify"
        return "left"

    def _local_visual_title_alignment(self, block: TextBlock, ctx: RenderContext, fallback):
        edges: List[Tuple[float, float]] = []
        for line in block.lines or []:
            if line.x1 is None or line.x2 is None:
                continue
            edges.append((float(line.x1), float(line.x2)))
        if not edges:
            return fallback

        col_left = float(getattr(ctx, "col_left_px", 0.0) or 0.0)
        col_right = float(getattr(ctx, "col_right_px", 0.0) or 0.0)
        col_w = max(col_right - col_left, 1.0)
        left_gap = min(max(0.0, x1 - col_left) for x1, _ in edges)
        right_gap = min(max(0.0, col_right - x2) for _, x2 in edges)
        left_edge_span = max(x1 for x1, _ in edges) - min(x1 for x1, _ in edges)
        right_edge_span = max(x2 for _, x2 in edges) - min(x2 for _, x2 in edges)
        centers = [((x1 + x2) * 0.5) for x1, x2 in edges]
        center_delta = abs((sum(centers) / len(centers)) - ((col_left + col_right) * 0.5))

        if (
            len(edges) >= 2
            and left_gap <= col_w * 0.08
            and left_edge_span <= col_w * 0.035
            and right_edge_span >= col_w * 0.12
        ):
            return WD_ALIGN_PARAGRAPH.LEFT
        if left_gap <= col_w * 0.08 and right_gap >= max(left_gap + col_w * 0.04, col_w * 0.035):
            return WD_ALIGN_PARAGRAPH.LEFT
        if right_gap <= col_w * 0.08 and left_gap >= max(right_gap + col_w * 0.04, col_w * 0.035):
            return WD_ALIGN_PARAGRAPH.RIGHT
        if center_delta <= col_w * 0.08:
            return WD_ALIGN_PARAGRAPH.CENTER
        return fallback

    def _local_visual_text_alignment(self, block: TextBlock, ctx: RenderContext, fallback):
        edges: List[Tuple[float, float]] = []
        for line in block.lines or []:
            if line.x1 is None or line.x2 is None:
                continue
            edges.append((float(line.x1), float(line.x2)))
        if not edges:
            return fallback

        col_left = float(getattr(ctx, "col_left_px", 0.0) or 0.0)
        col_right = float(getattr(ctx, "col_right_px", 0.0) or 0.0)
        col_w = max(col_right - col_left, 1.0)
        left_gap = sum(max(0.0, x1 - col_left) for x1, _ in edges) / len(edges)
        right_gap = sum(max(0.0, col_right - x2) for _, x2 in edges) / len(edges)
        centers = [((x1 + x2) * 0.5) for x1, x2 in edges]
        center_delta = abs((sum(centers) / len(centers)) - ((col_left + col_right) * 0.5))

        if left_gap <= col_w * 0.08 and right_gap >= max(left_gap + col_w * 0.04, col_w * 0.035):
            return WD_ALIGN_PARAGRAPH.LEFT
        if right_gap <= col_w * 0.08 and left_gap >= max(right_gap + col_w * 0.04, col_w * 0.035):
            return WD_ALIGN_PARAGRAPH.RIGHT
        if center_delta <= col_w * 0.08:
            return WD_ALIGN_PARAGRAPH.CENTER
        return fallback

    @staticmethod
    def _is_listish_line(text: str) -> bool:
        return bool(_LISTISH_RE.match((text or "").strip()))

    def _should_keep_visual_breaks(
        self,
        block: TextBlock,
        alignment,
        preserve_line_breaks: bool,
        preserve_breaks_on_ambiguous_justify: bool,
        ambiguous_justify: bool,
        visual_rows: List[List[object]],
        render_mode: str = "",
        in_table_cell: bool = False,
        force_table_breaks: bool = False,
        local_visual_band: bool = False,
    ) -> bool:
        if block.block_type in _CAPTION_TYPES:
            return True
        if not preserve_line_breaks:
            return False
        if (
            local_visual_band
            and block.block_type == BlockType.TITLE
            and alignment == WD_ALIGN_PARAGRAPH.LEFT
            and self._cjk_ratio(block.full_text()) < 0.20
        ):
            return False
        if (
            force_table_breaks
            and in_table_cell
            and block.block_type == BlockType.TEXT
            and self._cjk_ratio(block.full_text()) < 0.20
        ):
            return True

        row_count = len(visual_rows)
        text_len = len(block.full_text().strip())
        has_listish = any(
            row and self._is_listish_line(getattr(row[0], "text", ""))
            for row in visual_rows
        )

        if self._should_preserve_field_like_visual_breaks(block, visual_rows):
            return True

        if render_mode == "reflow" and block.block_type == BlockType.TEXT:
            return False

        if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            if not preserve_breaks_on_ambiguous_justify:
                return False
            if not ambiguous_justify:
                return False
            # 仅对短文本/列表维持视觉换行，避免长段正文被拆碎。
            return row_count <= 3 or has_listish or text_len <= 120

        if block.block_type in {
            BlockType.TITLE,
            BlockType.HEADER,
            BlockType.FOOTER,
            BlockType.PAGE_NUMBER,
            BlockType.TABLE_CAPTION,
            BlockType.FIGURE_CAPTION,
            BlockType.FORMULA_CAPTION,
            BlockType.TABLE_FOOTNOTE,
        }:
            return True

        if has_listish:
            return True
        if (
            alignment in {WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT}
            and self._looks_like_sparse_latin_linebreak_block(block, visual_rows)
        ):
            return True
        return row_count <= 2

    def _should_preserve_field_like_visual_breaks(
        self,
        block: TextBlock,
        visual_rows: List[List[object]],
    ) -> bool:
        if block.block_type not in {
            BlockType.TEXT,
            BlockType.HEADER,
            BlockType.FOOTER,
            BlockType.PAGE_NUMBER,
            BlockType.REFERENCE,
            BlockType.FOOTNOTE,
        }:
            return False
        if len(visual_rows) < 2:
            return False

        row_texts = [
            " ".join((getattr(ln, "text", "") or "").strip() for ln in row).strip()
            for row in visual_rows
        ]
        row_texts = [text for text in row_texts if text]
        if len(row_texts) < 2:
            return False

        field_hits = sum(1 for text in row_texts if _FIELD_LIKE_LINE_RE.search(text))
        colon_hits = sum(1 for text in row_texts if re.match(r"^\s*[\w\s]{1,32}:", text))
        if field_hits >= 1 and (field_hits + colon_hits) >= 2:
            return True

        if block.block_type in {BlockType.HEADER, BlockType.FOOTER, BlockType.PAGE_NUMBER}:
            return True

        page_h = max(float(getattr(block, "bbox", None).y2 if getattr(block, "bbox", None) else 0.0), 1.0)
        page = getattr(block, "page", None)
        if page is not None:
            page_h = max(float(getattr(page, "image_height", 0) or 0), page_h)
        page_like_top = float(block.bbox.y1) <= max(96.0, page_h * 0.07)
        if not page_like_top:
            return False

        row_width_ratios = self._visual_row_width_ratios(block, visual_rows)
        if len(row_width_ratios) < 2:
            return False
        width_spread = max(row_width_ratios) - min(row_width_ratios)
        shortish_rows = sum(1 for text in row_texts if len(text) <= 72)
        return width_spread >= 0.35 and shortish_rows >= max(1, len(row_texts) // 2)

    @staticmethod
    def _visual_row_width_ratios(block: TextBlock, visual_rows: List[List[object]]) -> List[float]:
        base_width = max(float(block.bbox.width), 1.0)
        ratios: List[float] = []
        for row in visual_rows:
            xs = [float(getattr(ln, "x1")) for ln in row if getattr(ln, "x1", None) is not None]
            xe = [float(getattr(ln, "x2")) for ln in row if getattr(ln, "x2", None) is not None]
            if xs and xe:
                ratios.append((max(xe) - min(xs)) / base_width)
        return ratios

    def _looks_like_sparse_latin_linebreak_block(
        self,
        block: TextBlock,
        visual_rows: List[List[object]],
    ) -> bool:
        """识别应保留 OCR 视觉换行的稀疏左对齐拉丁正文。"""
        if block.block_type != BlockType.TEXT:
            return False
        if len(visual_rows) < 4:
            return False
        text = block.full_text().strip()
        if not text or self._cjk_ratio(text) >= 0.20:
            return False

        row_widths = self._visual_row_width_ratios(block, visual_rows)

        if len(row_widths) < 4:
            return False

        avg_ratio = sum(row_widths) / len(row_widths)
        min_ratio = min(row_widths)
        max_ratio = max(row_widths)
        short_ratio = sum(1 for ratio in row_widths if ratio <= 0.60) / len(row_widths)
        long_ratio = sum(1 for ratio in row_widths if ratio >= 0.88) / len(row_widths)

        return (
            avg_ratio <= 0.78
            and (max_ratio - min_ratio) >= 0.40
            and short_ratio >= 0.25
            and long_ratio >= 0.25
        )

    # ------------------------------------------------------------------
    # 跨列区域的布局表格
    # ------------------------------------------------------------------

    def _render_native_columns_zone(self, doc, zone: Zone, page: "Page",
                                    col_px: dict) -> None:
        """使用 Word 原生分栏渲染无跨列区域。"""
        mapper = page.coord_mapper
        usable_w_pt = page.usable_width_pt
        img_w = page.image_width
        num_cols = max(zone.col_count, 1)
        col_widths = self._column_widths_pt(num_cols, col_px, img_w, usable_w_pt)

        for ci in range(num_cols):
            col_blks = sorted(
                [b for b in zone.blocks if b.col_index == ci and len(b.spanned_cols) <= 1],
                key=lambda b: b.bbox.y1,
            )
            cl = col_px.get(ci, [0, img_w])[0]
            cr = col_px.get(ci, [0, img_w])[1]
            ctx = RenderContext(
                coord_mapper=mapper,
                page=page,
                col_width_pt=col_widths[ci],
                col_left_px=cl,
                col_right_px=cr,
            )
            prev_y = 0
            for block in col_blks:
                gap = max(0, block.bbox.y1 - prev_y)
                sp = self._scale(max(min(mapper.h(gap), 12) - self._corr_gap_pt, 0)) if (prev_y > 0 and gap > 2) else 0
                self._render_block(doc, block, ctx, space_before=sp)
                prev_y = block.bbox.y2

            # 显式列断，避免后续列内容继续填充上一列剩余空间。这里仍是
            # Word 原生分栏，图片在所在栏内直接内联插入，没有使用布局表格。
            if ci < num_cols - 1:
                p = doc.add_paragraph()
                reset_paragraph_format(p)
                set_paragraph_spacing(p, space_before=0, space_after=0)
                p.add_run().add_break(WD_BREAK.COLUMN)

    @staticmethod
    def _paragraph_xml_has_content(p_el) -> bool:
        texts = p_el.findall('.//' + qn('w:t'))
        drawings = p_el.findall('.//' + qn('a:blip'),
                                {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
        return bool(texts or drawings)

    @classmethod
    def _prune_leading_empty_cell_paragraphs(cls, cell) -> None:
        """移除布局单元格开头的占位空段落，避免内容整体下沉。"""
        tc = cell._tc
        while True:
            children = list(tc)
            if len(children) <= 2:
                return
            lead_idx = 1 if children and children[0].tag == qn('w:tcPr') else 0
            if lead_idx >= len(children) - 1:
                return
            lead = children[lead_idx]
            if lead.tag != qn('w:p') or cls._paragraph_xml_has_content(lead):
                return
            tc.remove(lead)

    def _render_layout_table_zone(self, doc, zone: Zone, page: "Page",
                                  col_px: dict) -> None:
        """为含跨列区块的区域创建无边框布局表格。"""
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

        mapper = page.coord_mapper
        usable_w_pt = page.usable_width_pt
        img_w = page.image_width
        blocks = zone.blocks
        if not blocks:
            return

        num_cols = zone.col_count
        local_visual_band = self._is_local_visual_zone(zone, page)

        # 找出跨列集合
        spanned_set = set()
        for b in blocks:
            cols = self._layout_block_cols(b, num_cols, img_w)
            if len(cols) > 1:
                spanned_set.update(cols)

        visual_gap = min(12.0, usable_w_pt * 0.015)
        if local_visual_band:
            visual_gap = max(18.0, min(28.0, usable_w_pt * 0.035))
        col_widths = self._column_widths_pt(num_cols, col_px, img_w, usable_w_pt)
        if zone.region_kind == "decorative_sidecar" and num_cols == 2:
            label_width_px = max(
                (
                    float(block.bbox.width)
                    for block in blocks
                    if isinstance(block, TextBlock)
                    and (getattr(block, "attributes", None) or {}).get("docx_decorative_role") == "left_sidecar"
                ),
                default=0.0,
            )
            side_ratio = min(0.20, max(0.12, (label_width_px / max(float(img_w), 1.0)) + 0.035))
            col_widths = [usable_w_pt * side_ratio, usable_w_pt * (1.0 - side_ratio)]
            visual_gap = min(8.0, usable_w_pt * 0.01)
        if local_visual_band:
            col_widths = self._local_visual_band_col_widths_pt(
                zone,
                page,
                total_width_pt=usable_w_pt,
                visual_gap_pt=visual_gap,
            )

        if not spanned_set:
            # 无跨列区块：在简单的 N 列布局表格中渲染
            # gutter 仅作为视觉分隔（cell右边距），不从列宽中扣除，
            # 保证内容区宽度接近原始列宽，避免文字过度换行导致列高膨胀。
            tbl = doc.add_table(rows=1, cols=num_cols)
            tbl.autofit = False
            clear_table_borders(tbl)
            set_table_col_widths(tbl, col_widths)

            row = tbl.rows[0]
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                reset_paragraph_format(cell.paragraphs[0])
                set_paragraph_spacing(cell.paragraphs[0],
                                      line_spacing=self._scale(MIN_LINE_SPACING_PT), exact=True)

            for ci in range(num_cols):
                cell = row.cells[ci]
                if ci < num_cols - 1:
                    set_cell_right_margin(cell, visual_gap)
                col_blks = sorted([b for b in blocks if b.col_index == ci],
                                  key=lambda b: b.bbox.y1)
                if not col_blks:
                    continue
                cl = min(b.bbox.x1 for b in col_blks)
                cr = max(b.bbox.x2 for b in col_blks)
                ctx = RenderContext(coord_mapper=mapper, page=page,
                                   col_width_pt=col_widths[ci],
                                   col_left_px=cl, col_right_px=cr,
                                   in_table_cell=True,
                                   local_visual_band=local_visual_band)
                prev_y = 0
                for block in col_blks:
                    gap = max(0, block.bbox.y1 - prev_y)
                    sp = self._scale(max(min(mapper.h(gap), 12) - self._corr_gap_pt, 0)) if (prev_y > 0 and gap > 2) else 0
                    self._render_block(cell, block, ctx, space_before=sp)
                    prev_y = block.bbox.y2
                self._prune_leading_empty_cell_paragraphs(cell)
            return

        self._render_spanned_layout_table_zone(
            doc,
            zone,
            page,
            col_px,
            col_widths=col_widths,
            visual_gap=visual_gap,
        )
        return

    def _render_spanned_layout_table_zone(
        self,
        doc,
        zone: Zone,
        page: "Page",
        col_px: dict,
        *,
        col_widths: List[float],
        visual_gap: float,
    ) -> None:
        """Render mixed single-column and spanned blocks with parallel segments.

        The previous row-per-block table kept Word XML simple, but it serialized
        blocks from different columns vertically. That made visually parallel
        content add up in height and pushed many single-page sources onto a
        second Word page. This segment table keeps independent columns running
        side by side while still allowing a spanned segment to occupy multiple
        source columns.
        """
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

        mapper = page.coord_mapper
        img_w = page.image_width
        blocks = self._visual_order_for_spanned_zone(zone)
        num_cols = max(int(zone.col_count or 1), 1)
        if not blocks or num_cols <= 1:
            return

        def _init_cell(cell, add_gap: bool) -> None:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            reset_paragraph_format(cell.paragraphs[0])
            set_paragraph_spacing(
                cell.paragraphs[0],
                line_spacing=self._scale(MIN_LINE_SPACING_PT),
                exact=True,
            )
            if add_gap:
                set_cell_right_margin(cell, visual_gap)

        def _block_cols(block) -> List[int]:
            return self._layout_block_cols(block, num_cols, img_w)

        def _render_stream(
            cell,
            stream_blocks: List[Block],
            cols: List[int],
            width_pt: float,
            *,
            base_y: Optional[float] = None,
            gap_cap_pt: float = 4.0,
        ) -> None:
            if not stream_blocks:
                return
            left_px = min((col_px.get(ci, [0, img_w])[0] for ci in cols), default=0)
            right_px = max((col_px.get(ci, [0, img_w])[1] for ci in cols), default=img_w)
            ctx = RenderContext(
                coord_mapper=mapper,
                page=page,
                col_width_pt=width_pt,
                col_left_px=left_px,
                col_right_px=right_px,
                in_table_cell=True,
                span_gap_pt=max(0, len(cols) - 1) * visual_gap,
            )
            setattr(ctx, "render_mode", "reflow")
            prev_y = float(base_y or 0.0)
            for block in sorted(stream_blocks, key=lambda b: (float(b.bbox.y1), float(b.bbox.x1))):
                gap = max(0, block.bbox.y1 - prev_y)
                sp = self._scale(max(min(mapper.h(gap), gap_cap_pt) - self._corr_gap_pt, 0)) if (prev_y > 0 and gap > 2) else 0
                self._render_block(cell, block, ctx, space_before=sp)
                prev_y = block.bbox.y2
            self._prune_leading_empty_cell_paragraphs(cell)

        def _single_col_bands(segment_blocks: List[Block], segment_cols: set[int]) -> List[List[Block]]:
            singles = [
                block for block in segment_blocks
                if len(set(_block_cols(block)) & segment_cols) <= 1
            ]
            bands: List[List[Block]] = []
            band_start: Optional[float] = None
            band_bottom: Optional[float] = None
            band_window = max(220.0, min(800.0, float(page.image_height) * 0.12))
            for block in sorted(singles, key=lambda b: (float(b.bbox.y1), float(b.bbox.x1))):
                y1 = float(block.bbox.y1)
                y2 = float(block.bbox.y2)
                if (
                    not bands
                    or band_bottom is None
                    or band_start is None
                    or y1 > band_bottom + 24.0
                    or y1 > band_start + band_window
                ):
                    bands.append([block])
                    band_start = y1
                    band_bottom = y2
                    continue
                bands[-1].append(block)
                band_bottom = max(band_bottom, y2)
            return bands

        def _overlap_y(a: Block, b: Block) -> float:
            return max(0.0, min(float(a.bbox.y2), float(b.bbox.y2)) - max(float(a.bbox.y1), float(b.bbox.y1)))

        def _is_side_block(block: Block, span: Block, span_cols: set[int]) -> bool:
            cols = _block_cols(block)
            if not cols or cols[0] in span_cols:
                return False
            overlap = _overlap_y(block, span)
            if overlap <= 0:
                return False
            ref_h = max(1.0, min(float(block.bbox.height), float(span.bbox.height)))
            return overlap >= max(18.0, ref_h * 0.12)

        layout_span_types = {
            BlockType.FIGURE,
            BlockType.TABLE,
            BlockType.EQUATION,
        }
        all_spans = [block for block in blocks if len(_block_cols(block)) > 1]
        non_layout_spans = [block for block in all_spans if block.block_type not in layout_span_types]
        span_blocks = [block for block in all_spans]
        single_blocks = [block for block in blocks if len(_block_cols(block)) <= 1]
        consumed: set[int] = set()

        first_single_top = min((float(block.bbox.y1) for block in single_blocks), default=float("inf"))
        leading_spans = [
            block for block in sorted(non_layout_spans, key=lambda b: (float(b.bbox.y1), float(b.bbox.x1)))
            if float(block.bbox.y1) <= first_single_top + 8.0
        ]

        def _render_standalone_span(block: Block) -> None:
            cols = _block_cols(block)
            if not cols:
                return
            tbl = doc.add_table(rows=1, cols=num_cols)
            tbl.autofit = False
            clear_table_borders(tbl)
            set_table_col_widths(tbl, col_widths)
            row = tbl.rows[0]
            for ci, cell in enumerate(row.cells):
                _init_cell(cell, add_gap=(ci < num_cols - 1))
            start = min(cols)
            end = max(cols)
            cell = row.cells[start]
            if end > start:
                cell = cell.merge(row.cells[end])
            _render_stream(
                cell,
                [block],
                cols,
                sum(col_widths[start:end + 1]),
                base_y=float(block.bbox.y1),
                gap_cap_pt=10.0,
            )
            fit_table_to_width(tbl, sum(col_widths))

        for block in leading_spans:
            _render_standalone_span(block)
            consumed.add(id(block))

        body_single_top = min(
            (
                float(block.bbox.y1)
                for block in single_blocks
                if block.block_type not in {BlockType.HEADER, BlockType.FOOTER, BlockType.PAGE_NUMBER}
            ),
            default=first_single_top,
        )
        leading_layout_spans = [
            block for block in sorted(span_blocks, key=lambda b: (float(b.bbox.y1), float(b.bbox.x1)))
            if id(block) not in consumed
            and block.block_type in layout_span_types
            and float(block.bbox.y1) <= body_single_top
            and float(block.bbox.y2) <= body_single_top + max(18.0, float(page.image_height) * 0.015)
        ]
        for block in leading_layout_spans:
            _render_standalone_span(block)
            consumed.add(id(block))

        layout_spans = [
            block for block in span_blocks
            if id(block) not in consumed and block.block_type in layout_span_types
        ]
        spanned_col_set = sorted({
            ci
            for block in layout_spans
            for ci in _block_cols(block)
        })
        segments = self._column_segments(num_cols, spanned_col_set)
        outer_widths = [sum(col_widths[ci] for ci in cols) for _, cols in segments]

        tbl = doc.add_table(rows=1, cols=len(segments))
        tbl.autofit = False
        clear_table_borders(tbl)
        set_table_col_widths(tbl, outer_widths)
        row = tbl.rows[0]

        remaining = [block for block in blocks if id(block) not in consumed]
        segment_sets = [set(cols) for _, cols in segments]
        assigned: set[int] = set()

        def _render_segment_cell(cell, seg_cols: List[int], seg_blocks: List[Block], width_pt: float) -> None:
            if not seg_blocks:
                return
            if len(seg_cols) <= 1:
                _render_stream(cell, seg_blocks, seg_cols, width_pt, gap_cap_pt=10.0)
                return

            local_widths = [col_widths[ci] for ci in seg_cols]
            sub_tbl = cell.add_table(rows=0, cols=len(seg_cols))
            sub_tbl.autofit = False
            clear_table_borders(sub_tbl)

            local_index = {ci: idx for idx, ci in enumerate(seg_cols)}
            span_items = [
                block for block in seg_blocks
                if len(set(_block_cols(block)) & set(seg_cols)) > 1
            ]
            single_items = [
                block for block in seg_blocks
                if len(set(_block_cols(block)) & set(seg_cols)) <= 1
            ]
            sub_units: List[Tuple[float, str, object]] = []
            assigned_single_ids: set[int] = set()
            for span in sorted(span_items, key=lambda b: (float(b.bbox.y1), float(b.bbox.x1))):
                before = [
                    block for block in single_items
                    if id(block) not in assigned_single_ids
                    and float(block.bbox.y2) <= float(span.bbox.y1) + 8.0
                ]
                if before:
                    sub_units.append((min(float(b.bbox.y1) for b in before), "band", before))
                    assigned_single_ids.update(id(block) for block in before)
                span_cols = set(_block_cols(span))
                side_blocks = [
                    block for block in single_items
                    if id(block) not in assigned_single_ids
                    and not (set(_block_cols(block)) & span_cols)
                    and _is_side_block(block, span, span_cols)
                ]
                assigned_single_ids.update(id(block) for block in side_blocks)
                sub_units.append((float(span.bbox.y1), "span", (span, side_blocks)))
            after = [block for block in single_items if id(block) not in assigned_single_ids]
            if after:
                sub_units.append((min(float(b.bbox.y1) for b in after), "band", after))

            for row_top, kind, payload in sorted(sub_units, key=lambda item: item[0]):
                sub_row = sub_tbl.add_row()
                for local_ci, sub_cell in enumerate(sub_row.cells):
                    _init_cell(sub_cell, add_gap=(local_ci < len(seg_cols) - 1))
                if kind == "span":
                    block, side_blocks = payload
                    cols = [ci for ci in _block_cols(block) if ci in local_index]
                    start = min(local_index[ci] for ci in cols)
                    end = max(local_index[ci] for ci in cols)
                    sub_cell = sub_row.cells[start]
                    if end > start:
                        sub_cell = sub_cell.merge(sub_row.cells[end])
                    _render_stream(
                        sub_cell,
                        [block],
                        cols,
                        sum(local_widths[start:end + 1]),
                        base_y=row_top,
                        gap_cap_pt=14.0,
                    )
                    side_by_col: dict[int, List[Block]] = defaultdict(list)
                    for side_block in side_blocks:
                        side_cols = [ci for ci in _block_cols(side_block) if ci in local_index]
                        if not side_cols:
                            continue
                        side_by_col[side_cols[0]].append(side_block)
                    occupied = set(cols)
                    for ci, side_stream in side_by_col.items():
                        if ci in occupied:
                            continue
                        _render_stream(
                            sub_row.cells[local_index[ci]],
                            side_stream,
                            [ci],
                            local_widths[local_index[ci]],
                            base_y=row_top,
                            gap_cap_pt=10.0,
                        )
                    continue

                by_col: dict[int, List[Block]] = defaultdict(list)
                for block in payload:
                    cols = _block_cols(block)
                    if cols:
                        by_col[cols[0]].append(block)
                for ci in seg_cols:
                    _render_stream(
                        sub_row.cells[local_index[ci]],
                        by_col.get(ci, []),
                        [ci],
                        local_widths[local_index[ci]],
                        base_y=row_top,
                        gap_cap_pt=10.0,
                    )

            set_table_col_widths(sub_tbl, local_widths)
            fit_table_to_width(sub_tbl, width_pt)
            self._prune_leading_empty_cell_paragraphs(cell)

        for seg_idx, (_, cols) in enumerate(segments):
            cell = row.cells[seg_idx]
            _init_cell(cell, add_gap=(seg_idx < len(segments) - 1))
            col_set = segment_sets[seg_idx]
            stream: List[Block] = []
            for block in remaining:
                if id(block) in assigned:
                    continue
                block_cols = set(_block_cols(block))
                if len(block_cols) > 1:
                    if block_cols <= col_set:
                        stream.append(block)
                        assigned.add(id(block))
                    continue
                if block_cols and next(iter(block_cols)) in col_set:
                    stream.append(block)
                    assigned.add(id(block))
            _render_segment_cell(cell, cols, stream, outer_widths[seg_idx])

        leftovers = [block for block in remaining if id(block) not in assigned]
        if leftovers:
            _render_stream(row.cells[-1], leftovers, list(range(num_cols)), outer_widths[-1], gap_cap_pt=10.0)

        fit_table_to_width(tbl, sum(col_widths))
        return

    @staticmethod
    def _visual_order_for_spanned_zone(zone: Zone) -> List[Block]:
        blocks = list(getattr(zone, "blocks", []) or [])
        if not blocks:
            return []
        has_flow = any(
            bool((getattr(block, "attributes", None) or {}).get("flow_id"))
            for block in blocks
        )
        has_cross_media = any(
            block.block_type in {BlockType.FIGURE, BlockType.TABLE}
            and len(getattr(block, "spanned_cols", []) or []) > 1
            for block in blocks
        )
        if has_flow or not has_cross_media:
            return sorted(blocks, key=lambda b: (float(b.bbox.y1), float(b.bbox.x1)))

        def _band_key(block: Block) -> tuple[int, int, float, float]:
            y1 = float(block.bbox.y1)
            x1 = float(block.bbox.x1)
            col = int(getattr(block, "col_index", 0) or 0)
            band = int(y1 // 900.0)
            if block.block_type in {BlockType.TITLE, BlockType.FIGURE, BlockType.TABLE}:
                band = max(0, band - 1)
            return (band, col, y1, x1)

        return sorted(blocks, key=_band_key)

    # ------------------------------------------------------------------
    # 工具方法
    # ------------------------------------------------------------------

    @staticmethod
    def _cleanup_trailing_paragraphs(doc: DocxDocument) -> None:
        """移除尾部空段落并将不可避免的空段落缩小，
        以避免产生额外的空白页。"""
        from lxml import etree

        body = doc.element.body
        paras = body.findall(qn('w:p'))
        last_content_idx = -1
        for idx, p_el in enumerate(paras):
            runs = p_el.findall('.//' + qn('w:t'))
            blips = p_el.findall('.//' + qn('a:blip'),
                                 {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
            if runs or blips:
                last_content_idx = idx

        to_remove = []
        for idx, p_el in enumerate(paras):
            if idx > last_content_idx:
                has_sect = p_el.find(qn('w:pPr') + '/' + qn('w:sectPr')) is not None
                if not has_sect and idx > last_content_idx + 1:
                    to_remove.append(p_el)
        for p_el in to_remove:
            body.remove(p_el)

        # 若最后一个非 sectPr 元素是表格，Word 会自动插入一个
        # 段落标记，可能导致空白第二页。插入一个显式的近零高度
        # 段落来控制这种情况。
        children = list(body)
        sectPr_el = None
        for ch in reversed(children):
            if ch.tag == qn('w:sectPr'):
                sectPr_el = ch
            elif ch.tag == qn('w:tbl'):
                # 最后一个内容元素是表格 – 添加微小段落
                tiny_p = etree.SubElement(body, qn('w:p'))
                body.remove(tiny_p)  # SubElement 会追加；我们需要插入到指定位置
                if sectPr_el is not None:
                    sectPr_el.addprevious(tiny_p)
                else:
                    body.append(tiny_p)
                break
            else:
                break

        # 将尾部空段落缩小到近零高度，避免 Word 必须的文档结束标记
        # 导致产生“空白最后一页”。
        for p_el in body.findall(qn('w:p')):
            runs = p_el.findall('.//' + qn('w:t'))
            blips = p_el.findall('.//' + qn('a:blip'),
                                 {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
            if runs or blips:
                continue
            # 检查是否位于最后一个内容元素之后
            is_trailing = True
            for sibling in p_el.itersiblings():
                tag = sibling.tag
                if tag == qn('w:p'):
                    s_runs = sibling.findall('.//' + qn('w:t'))
                    s_blips = sibling.findall('.//' + qn('a:blip'),
                                              {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
                    if s_runs or s_blips:
                        is_trailing = False
                        break
                elif tag == qn('w:tbl'):
                    is_trailing = False
                    break
            if not is_trailing:
                continue

            pPr = p_el.find(qn('w:pPr'))
            if pPr is None:
                pPr = etree.SubElement(p_el, qn('w:pPr'))
                p_el.insert(0, pPr)
            spacing = pPr.find(qn('w:spacing'))
            if spacing is None:
                spacing = etree.SubElement(pPr, qn('w:spacing'))
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '0')
            spacing.set(qn('w:line'), '1')
            spacing.set(qn('w:lineRule'), 'exact')
            rPr = pPr.find(qn('w:rPr'))
            if rPr is None:
                rPr = etree.SubElement(pPr, qn('w:rPr'))
            sz = rPr.find(qn('w:sz'))
            if sz is None:
                sz = etree.SubElement(rPr, qn('w:sz'))
            sz.set(qn('w:val'), '2')  # 1pt = 2 half-points
