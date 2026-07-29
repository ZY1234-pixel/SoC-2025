"""Mechanical DOCX emission from a complete ReflowLayoutPlan."""

from __future__ import annotations

import base64
import io
import math
from dataclasses import replace
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from docflow.model.stages import FlowKind, Rect, ReflowLayoutPlan
from docflow.planning.text_metrics import (
    estimate_text_units,
    estimate_wrapped_lines,
    fit_font_size_to_lines,
    infer_occupancy_line_height,
)
from docflow.renderer.docx_utils.html_table import (
    get_table_cell_placements,
    get_table_column_weights,
    get_table_dimensions,
)
from docflow.renderer.docx_utils.table_fmt import (
    clear_table_borders,
    set_cell_margins,
    set_horizontal_table_borders,
    set_table_col_widths,
)
_ALIGNMENT = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

class ReflowDocxRenderer:
    def render(self, plan: ReflowLayoutPlan, output_path: str) -> None:
        document = self.build(plan)
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        temporary = output.with_name(f".{output.name}.tmp")
        document.save(temporary)
        temporary.replace(output)

    def build(self, plan: ReflowLayoutPlan):
        document = Document()
        roles = {role.role_id: role for role in plan.roles}
        for page_number, page in enumerate(plan.pages):
            section = document.sections[0] if page_number == 0 else document.add_section(WD_SECTION.NEW_PAGE)
            if page_number > 0:
                self._collapse_section_break(document.paragraphs[-1])
            self._set_page_geometry(section, page.geometry)
            elements = {element.element_id: element for element in page.elements}
            self._set_furniture_distances(section, page, elements)
            usable_width = self._usable_width(page.geometry)
            self._render_furniture(section.header, page.header_element_ids, elements, roles, page.fit_scale, usable_width)
            self._render_furniture(section.footer, page.footer_element_ids, elements, roles, page.fit_scale, usable_width)
            body = document
            for flow in page.sections:
                if flow.kind == FlowKind.SINGLE:
                    for identifier in flow.element_ids:
                        self._render_element(body, elements[identifier], roles, page.fit_scale, usable_width)
                elif flow.kind == FlowKind.WRAPPED:
                    self._render_wrapped(body, flow, elements, roles, page.fit_scale, usable_width)
                elif flow.kind == FlowKind.SEQUENTIAL_COLUMNS:
                    self._render_columns(body, flow, elements, roles, page.fit_scale)
                else:
                    self._render_grid(body, flow, elements, roles, page.fit_scale)
        if plan.pages:
            self._collapse_section_break(document.add_paragraph())
        return document

    @staticmethod
    def _collapse_section_break(paragraph) -> None:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(1)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        properties = paragraph._p.get_or_add_pPr()
        spacing = properties.find(qn("w:spacing"))
        spacing.set(qn("w:line"), "1")
        mark_properties = properties.find(qn("w:rPr"))
        if mark_properties is None:
            mark_properties = OxmlElement("w:rPr")
            properties.append(mark_properties)
        mark_properties.append(OxmlElement("w:vanish"))
        if not paragraph.runs:
            paragraph.add_run()
        for run in paragraph.runs:
            run.font.size = Pt(0.5)
            run._r.get_or_add_rPr().append(OxmlElement("w:vanish"))

    @staticmethod
    def _set_page_geometry(section, geometry) -> None:
        for grid in section._sectPr.xpath("./w:docGrid"):
            section._sectPr.remove(grid)
        section.page_width = Pt(geometry.width_pt)
        section.page_height = Pt(geometry.height_pt)
        section.top_margin = Pt(geometry.margin_top_pt)
        section.right_margin = Pt(geometry.margin_right_pt)
        section.bottom_margin = Pt(geometry.margin_bottom_pt)
        section.left_margin = Pt(geometry.margin_left_pt)
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

    def _render_furniture(self, container, identifiers, elements, roles, fit_scale, usable_width) -> None:
        self._clear_container(container)
        if not identifiers:
            return
        if all(elements[identifier].payload.get("image_base64") for identifier in identifiers):
            for identifier in identifiers:
                element = elements[identifier]
                data = self._decode_image(element.payload.get("image_base64"))
                if not data:
                    continue
                bbox = element.payload.get("source_bbox") or (0, 0, 1, 1)
                source_scale = float(element.payload.get("source_scale", 1.0))
                paragraph = container.add_paragraph()
                picture = paragraph.add_run().add_picture(
                    io.BytesIO(data),
                    width=Pt(max((float(bbox[2]) - float(bbox[0])) * source_scale, 0.5)),
                )
                self._anchor_picture(
                    picture._inline,
                    float(bbox[0]) * source_scale,
                    float(bbox[1]) * source_scale,
                )
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = Pt(1)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            return
        if len(identifiers) == 1:
            element = elements[identifiers[0]]
            if element.text_structure.orientation == "vertical":
                self._write_vertical_text(container, element, roles, fit_scale)
                return
            paragraph = container.add_paragraph()
            self._write_text(paragraph, element, roles, fit_scale)
            return
        table = container.add_table(rows=1, cols=3, width=Pt(usable_width))
        self._format_layout_table(table, (usable_width / 3,) * 3)
        for cell in table.rows[0].cells:
            self._clear_container(cell)
        for identifier in identifiers:
            element = elements[identifier]
            bbox = element.payload.get("source_bbox") or (0, 0, 1, 1)
            page_width_px = max(float(element.payload.get("page_width_px", 1.0)), 1.0)
            column = min(int(((float(bbox[0]) + float(bbox[2])) / 2.0) / page_width_px * 3), 2)
            cell = table.cell(0, column)
            paragraph = cell.paragraphs[-1] if cell.paragraphs else cell.add_paragraph()
            paragraph.alignment = (WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT)[column]
            if element.text_structure.orientation == "vertical":
                self._write_vertical_text(cell, element, roles, fit_scale)
                continue
            if element.payload.get("image_base64"):
                data = self._decode_image(element.payload.get("image_base64"))
                if data:
                    width = (float(bbox[2]) - float(bbox[0])) / page_width_px * usable_width * fit_scale
                    paragraph.add_run().add_picture(io.BytesIO(data), width=Pt(max(width, 0.5)))
                continue
            self._write_text(paragraph, element, roles, fit_scale)
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.right_indent = Pt(0)
            paragraph.alignment = (WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT)[column]

    def _render_columns(self, container, flow, elements, roles, fit_scale) -> None:
        table = container.add_table(rows=1, cols=len(flow.column_widths_pt))
        self._format_layout_table(table, flow.column_widths_pt, flow.gutter_pt)
        if flow.row_heights_pt:
            table.rows[0].height = Pt(flow.row_heights_pt[0] * fit_scale)
            table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for column, cell in enumerate(table.rows[0].cells):
            self._clear_container(cell)
            for identifier in flow.element_ids:
                element = elements[identifier]
                if int(element.payload.get("column", 0)) == column:
                    self._render_element(cell, element, roles, fit_scale, flow.column_widths_pt[column])
        self._collapse_trailing_paragraph(container)

    def _render_wrapped(self, container, flow, elements, roles, fit_scale, container_width) -> None:
        floating = elements[flow.floating_element_id]
        width = flow.floating_width_pt * fit_scale
        table = container.add_table(rows=1, cols=1)
        self._format_layout_table(table, (width,))
        self._float_table(
            table,
            flow.floating_side,
            flow.floating_offset_x_pt * fit_scale,
            flow.floating_offset_y_pt * fit_scale,
            flow.gutter_pt * fit_scale,
        )
        cell = table.cell(0, 0)
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
        self._clear_container(cell)
        data = self._decode_image(floating.payload.get("image_base64"))
        if floating.kind == "figure_group" and data:
            paragraph = cell.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.add_run().add_picture(io.BytesIO(data), width=Pt(max(width, 0.5)))
            self._write_caption(
                cell,
                floating.payload.get("caption"),
                roles,
                fit_scale,
                floating.payload.get("caption_alignment"),
                roles.get(floating.role_id),
                floating.payload.get("caption_font_size_pt"),
            )
        else:
            self._render_element(cell, floating, roles, fit_scale, width)
        self._collapse_trailing_paragraph(cell)
        self._collapse_trailing_paragraph(container)
        for identifier in flow.element_ids:
            if identifier != flow.floating_element_id:
                self._render_element(container, elements[identifier], roles, fit_scale, container_width)

    @staticmethod
    def _float_table(table, side: str, offset_x_pt: float, offset_y_pt: float, gutter_pt: float) -> None:
        properties = table._tbl.tblPr
        positioning = OxmlElement("w:tblpPr")
        distance = str(round(max(gutter_pt / 2.0, 2.0) * 20))
        positioning.set(qn("w:leftFromText"), distance)
        positioning.set(qn("w:rightFromText"), distance)
        positioning.set(qn("w:topFromText"), "0")
        positioning.set(qn("w:bottomFromText"), "0")
        positioning.set(qn("w:vertAnchor"), "text")
        positioning.set(qn("w:horzAnchor"), "text")
        if side == "left":
            positioning.set(qn("w:tblpX"), str(round(offset_x_pt * 20)))
        else:
            positioning.set(qn("w:tblpXSpec"), side)
        positioning.set(qn("w:tblpY"), str(round(offset_y_pt * 20)))
        properties.append(positioning)

    @staticmethod
    def _anchor_picture(inline, x_pt: float, y_pt: float) -> None:
        inline.tag = qn("wp:anchor")
        for name, value in {
            "distT": "0",
            "distB": "0",
            "distL": "0",
            "distR": "0",
            "simplePos": "0",
            "relativeHeight": "0",
            "behindDoc": "0",
            "locked": "0",
            "layoutInCell": "1",
            "allowOverlap": "1",
        }.items():
            inline.set(name, value)
        simple_position = OxmlElement("wp:simplePos")
        simple_position.set("x", "0")
        simple_position.set("y", "0")
        inline.insert(0, simple_position)
        for index, (axis, offset) in enumerate((("H", x_pt), ("V", y_pt)), start=1):
            position = OxmlElement(f"wp:position{axis}")
            position.set("relativeFrom", "page")
            value = OxmlElement("wp:posOffset")
            value.text = str(round(offset * 12700))
            position.append(value)
            inline.insert(index, position)
        inline.insert(4, OxmlElement("wp:wrapNone"))

    def _render_grid(self, container, flow, elements, roles, fit_scale) -> None:
        row_count = max(cell.row + cell.row_span for cell in flow.grid_cells)
        column_count = len(flow.column_widths_pt)
        table = container.add_table(rows=row_count, cols=column_count)
        self._format_layout_table(table, flow.column_widths_pt, flow.gutter_pt)
        # Word needs room for the closing paragraph mark of a vertically merged text cell.
        terminal_text_ids = tuple(
            identifier
            for cell in flow.grid_cells
            if cell.row + cell.row_span == row_count and cell.row_span > 1
            for identifier in cell.element_ids
            if elements[identifier].kind != "figure_group"
        )
        terminal_line_reserve = max(
            (
                float(
                    elements[identifier].payload.get("line_height_pt")
                    or (
                        roles[elements[identifier].role_id].font_size_pt
                        * roles[elements[identifier].role_id].line_spacing
                        if elements[identifier].role_id in roles
                        else 0.0
                    )
                )
                * fit_scale
                * 0.5
                for identifier in terminal_text_ids
            ),
            default=0.0,
        )
        for index, (row, height) in enumerate(zip(table.rows, flow.row_heights_pt)):
            row.height = Pt(
                height * fit_scale
                + (terminal_line_reserve if index == row_count - 1 else 0.0)
            )
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        target_cells = {}
        for grid_cell in flow.grid_cells:
            cell = table.cell(grid_cell.row, grid_cell.column)
            floats_across_rows = grid_cell.row_span > 1 and any(
                elements[identifier].kind == "figure_group" for identifier in grid_cell.element_ids
            )
            if grid_cell.row_span > 1 or grid_cell.column_span > 1:
                cell = cell.merge(
                    table.cell(
                        grid_cell.row if floats_across_rows else grid_cell.row + grid_cell.row_span - 1,
                        grid_cell.column + grid_cell.column_span - 1,
                    )
                )
            half_gutter_twips = int(max(flow.gutter_pt, 0.0) * 10)
            set_cell_margins(
                cell,
                top=0,
                bottom=0,
                start=half_gutter_twips if grid_cell.column > 0 else 0,
                end=(
                    half_gutter_twips
                    if grid_cell.column + grid_cell.column_span < column_count
                    else 0
                ),
            )
            target_cells[(grid_cell.row, grid_cell.column)] = cell
        for row in table.rows:
            for cell in row.cells:
                if cell._tc.xpath("./w:p | ./w:tbl"):
                    self._clear_container(cell)
        for grid_cell in flow.grid_cells:
            cell = target_cells[(grid_cell.row, grid_cell.column)]
            container_width = (
                sum(flow.column_widths_pt[grid_cell.column : grid_cell.column + grid_cell.column_span])
                + flow.gutter_pt * (grid_cell.column_span - 1)
            )
            for identifier in grid_cell.element_ids:
                element = elements[identifier]
                if element.kind == "figure_group" and grid_cell.row_span > 1:
                    element = replace(element, payload={**element.payload, "float_in_grid": True})
                self._render_element(cell, element, roles, fit_scale, container_width)
        for row in table.rows:
            for cell in row.cells:
                if not cell.paragraphs and not cell.tables:
                    cell.add_paragraph()
                    self._collapse_trailing_paragraph(cell)
        self._collapse_trailing_paragraph(container)

    def _render_element(self, container, element, roles, fit_scale: float, container_width: float) -> None:
        if element.kind in {"heading", "paragraph_group", "caption"}:
            if element.text_structure.orientation == "vertical":
                self._write_vertical_text(container, element, roles, fit_scale)
                return
            split_text_rows = element.payload.get("split_text_rows") or ()
            if split_text_rows:
                for index, row in enumerate(split_text_rows):
                    payload = dict(element.payload)
                    payload.update(
                        split_text_rows=(row,),
                        lines=tuple(row),
                        space_before_pt=payload.get("space_before_pt", 0.0) if index == 0 else 0.0,
                    )
                    paragraph = container.add_paragraph()
                    self._write_text(
                        paragraph,
                        replace(element, text=" ".join(row), payload=payload),
                        roles,
                        fit_scale,
                        container_width,
                    )
                return
            paragraph = container.add_paragraph()
            self._write_text(paragraph, element, roles, fit_scale, container_width)
            return
        self._write_block_spacing(container, element, fit_scale)
        if element.kind == "figure_group":
            if element.payload.get("caption_position") == "before":
                self._write_caption(container, element.payload.get("caption"), roles, fit_scale, element.payload.get("caption_alignment"), roles.get(element.role_id), element.payload.get("caption_font_size_pt"), container_width)
            self._write_image(container, element, fit_scale, container_width)
            if element.payload.get("caption_position") != "before":
                self._write_caption(container, element.payload.get("caption"), roles, fit_scale, element.payload.get("caption_alignment"), roles.get(element.role_id), element.payload.get("caption_font_size_pt"), container_width)
            return
        if element.kind == "equation_group":
            self._write_equation(container, element, roles, fit_scale, container_width)
            return
        if element.kind == "table_group":
            if element.payload.get("caption_position") == "before":
                self._write_caption(container, element.payload.get("caption"), roles, fit_scale, element.payload.get("caption_alignment"), roles.get(element.role_id), element.payload.get("caption_font_size_pt"), container_width)
            self._write_native_table(container, element, roles, fit_scale, container_width)
            if element.payload.get("caption_position") != "before":
                self._write_caption(container, element.payload.get("caption"), roles, fit_scale, element.payload.get("caption_alignment"), roles.get(element.role_id), element.payload.get("caption_font_size_pt"), container_width)

    def _write_vertical_text(self, container, element, roles, fit_scale: float) -> None:
        bbox = element.payload.get("source_bbox") or (0, 0, 1, 1)
        width = max((float(bbox[2]) - float(bbox[0])) * float(element.payload.get("source_scale", 1.0)), 1.0)
        if hasattr(container, "_tc"):
            self._shade(container._tc.get_or_add_tcPr(), element.payload.get("background_color"))
        table = container.add_table(rows=1, cols=1)
        self._format_layout_table(table, (width,))
        cell = table.cell(0, 0)
        self._shade(cell._tc.get_or_add_tcPr(), element.payload.get("background_color"))
        paragraph = cell.paragraphs[0]
        self._write_text(paragraph, element, roles, fit_scale, width)
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)
        self._collapse_trailing_paragraph(container)

    def _write_text(self, paragraph, element, roles, fit_scale: float, container_width: float | None = None) -> None:
        self._write_paragraph_geometry(paragraph, element, fit_scale)
        if element.text_structure.is_list:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if element.payload.get("background_color") and container_width:
            left_indent = paragraph.paragraph_format.left_indent.pt if paragraph.paragraph_format.left_indent else 0.0
            visual_width = container_width * float(element.payload.get("width_fraction", 1.0)) * fit_scale
            paragraph.paragraph_format.right_indent = Pt(max(container_width - left_indent - visual_width, 0.0))
        role = roles.get(element.role_id) or self._body_role(roles)
        font_size = max(round(role.font_size_pt * fit_scale * 2) / 2.0, 0.5) if role else None
        source_lines = element.payload.get("lines") or ()
        split_text_rows = element.payload.get("split_text_rows") or ()
        if split_text_rows and container_width:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            right_indent = paragraph.paragraph_format.right_indent.pt if paragraph.paragraph_format.right_indent else 0.0
            paragraph.paragraph_format.tab_stops.add_tab_stop(
                Pt(max(container_width - right_indent, 1.0)),
                WD_TAB_ALIGNMENT.RIGHT,
            )
        if font_size and container_width and len(source_lines) > 1 and not split_text_rows:
            left_indent = paragraph.paragraph_format.left_indent.pt if paragraph.paragraph_format.left_indent else 0.0
            right_indent = paragraph.paragraph_format.right_indent.pt if paragraph.paragraph_format.right_indent else 0.0
            first_indent = paragraph.paragraph_format.first_line_indent.pt if paragraph.paragraph_format.first_line_indent else 0.0
            line_widths = (container_width - left_indent - right_indent - first_indent,) + (
                container_width - left_indent - right_indent,
            ) * (len(source_lines) - 1)
            font_size = fit_font_size_to_lines(
                font_size,
                tuple(str(line) for line in source_lines),
                line_widths,
                0.90
                if element.kind == "heading" or element.text_structure.preserve_source_lines
                else 0.99,
            )
        elif (
            font_size
            and container_width
            and len(source_lines) == 1
            and element.text_structure.orientation != "vertical"
        ):
            units = estimate_text_units(element.text)
            visual_width = container_width * float(element.payload.get("width_fraction", 1.0))
            font_size = min(font_size, visual_width * 0.90 / max(units, 1.0))
        line_height = element.payload.get("line_height_pt")
        if role and line_height:
            rendered_line_height = max(float(line_height) * fit_scale, font_size * 1.05)
            if (
                container_width
                and element.kind == "paragraph_group"
                and not element.text_structure.preserve_source_lines
                and element.text_structure.orientation == "horizontal"
            ):
                left_indent = paragraph.paragraph_format.left_indent.pt if paragraph.paragraph_format.left_indent else 0.0
                right_indent = paragraph.paragraph_format.right_indent.pt if paragraph.paragraph_format.right_indent else 0.0
                width = max(container_width - left_indent - right_indent, 1.0)
                content_bbox = self._content_bbox(element)
                lines = estimate_wrapped_lines(
                    element.text,
                    font_size,
                    width,
                    int(element.payload.get("visual_line_count") or len(source_lines)),
                    content_bbox.width * float(element.payload.get("source_scale", 1.0)),
                    fit_scale,
                )
                rendered_line_height = infer_occupancy_line_height(
                    font_size,
                    rendered_line_height,
                    content_bbox.height * float(element.payload.get("source_scale", 1.0)) * fit_scale,
                    lines,
                )
            paragraph.paragraph_format.line_spacing = Pt(rendered_line_height)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        elif role:
            paragraph.paragraph_format.line_spacing = role.line_spacing
        run = paragraph.add_run(self._visual_text(element))
        self._style_run(run, role, 1.0, font_size_pt=font_size)
        self._shade(paragraph._p.get_or_add_pPr(), element.payload.get("background_color"))
        paragraph.paragraph_format.widow_control = False
        paragraph.paragraph_format.keep_together = False
        paragraph.paragraph_format.keep_with_next = False

    @staticmethod
    def _write_paragraph_geometry(paragraph, element, fit_scale: float) -> None:
        paragraph.paragraph_format.space_before = Pt(float(element.payload.get("space_before_pt", 0.0)) * fit_scale)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(float(element.payload.get("first_line_indent_pt", 0.0)) * fit_scale)
        paragraph.paragraph_format.left_indent = Pt(float(element.payload.get("left_indent_pt", 0.0)))
        paragraph.paragraph_format.right_indent = Pt(float(element.payload.get("right_indent_pt", 0.0)))
        paragraph.alignment = _ALIGNMENT.get(element.payload.get("alignment"), WD_ALIGN_PARAGRAPH.LEFT)

    @staticmethod
    def _set_furniture_distances(section, page, elements) -> None:
        if page.header_element_ids:
            section.header_distance = Pt(0 if all(
                elements[identifier].payload.get("image_base64") for identifier in page.header_element_ids
            ) else (
                min(float(elements[identifier].payload["source_bbox"][1]) for identifier in page.header_element_ids)
                * float(elements[page.header_element_ids[0]].payload.get("source_scale", 1.0))
            ))
        if page.footer_element_ids:
            reference = elements[page.footer_element_ids[0]]
            page_height = float(reference.payload.get("page_height_px", 1.0))
            bottom = max(float(elements[identifier].payload["source_bbox"][3]) for identifier in page.footer_element_ids)
            section.footer_distance = Pt((page_height - bottom) * float(reference.payload.get("source_scale", 1.0)))

    @staticmethod
    def _write_block_spacing(container, element, fit_scale: float) -> None:
        spacing = float(element.payload.get("space_before_pt", 0.0)) * fit_scale
        if spacing <= 0:
            return
        paragraph = container.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(max(spacing - 1.0, 0.0))
        paragraph.paragraph_format.line_spacing = Pt(1)

    def _write_caption(self, container, text, roles, fit_scale: float, alignment=None, fallback_role=None, font_size_pt=None, container_width=None) -> None:
        if not text:
            return
        role = fallback_role or next((role for role_id, role in roles.items() if role_id.startswith("caption_")), None) or self._body_role(roles)
        paragraph = container.add_paragraph()
        paragraph.alignment = _ALIGNMENT.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        if "\t" in str(text) and container_width:
            paragraph.paragraph_format.tab_stops.add_tab_stop(Pt(container_width / 2.0), WD_TAB_ALIGNMENT.CENTER)
        run = paragraph.add_run(str(text))
        font_size = max(round(float(font_size_pt) * fit_scale * 2) / 2.0, 0.5) if font_size_pt else None
        self._style_run(run, role, fit_scale if font_size is None else 1.0, font_size_pt=font_size)
        rendered_size = font_size or max(round(role.font_size_pt * fit_scale * 2) / 2.0, 0.5)
        paragraph.paragraph_format.line_spacing = Pt(rendered_size * 1.1)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    def _write_image(self, container, element, fit_scale: float, container_width: float) -> None:
        data = self._decode_image(element.payload.get("image_base64"))
        if not data:
            return
        bbox = element.payload.get("primary_bbox") or element.payload.get("source_bbox") or (0, 0, 1, 1)
        source_width = (float(bbox[2]) - float(bbox[0])) * float(element.payload.get("source_scale", 1.0))
        width = min(container_width * float(element.payload.get("width_fraction", 1.0)), source_width) * (
            1.0 if element.kind == "figure_group" else fit_scale
        )
        paragraph = container.add_paragraph()
        paragraph.alignment = _ALIGNMENT.get(element.payload.get("alignment"), WD_ALIGN_PARAGRAPH.CENTER)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        picture = paragraph.add_run().add_picture(io.BytesIO(data), width=Pt(max(width, 0.5)))
        if element.payload.get("float_in_grid"):
            self._float_picture(picture._inline, element.payload.get("alignment"))
            paragraph.paragraph_format.line_spacing = Pt(1)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    @classmethod
    def _float_picture(cls, inline, alignment) -> None:
        cls._anchor_picture(inline, 0.0, 0.0)
        horizontal = inline.find(qn("wp:positionH"))
        horizontal.set("relativeFrom", "column")
        position = horizontal[0]
        position.tag = qn("wp:align")
        position.text = alignment if alignment in {"left", "center", "right"} else "left"
        inline.find(qn("wp:positionV")).set("relativeFrom", "paragraph")

    def _write_equation(self, container, element, roles, fit_scale: float, container_width: float) -> None:
        number = str(element.payload.get("number") or "")
        if not number:
            self._write_image(container, element, fit_scale, container_width)
            if not element.payload.get("image_base64") and element.payload.get("latex"):
                paragraph = container.add_paragraph(str(element.payload["latex"]))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return
        table = container.add_table(rows=1, cols=2)
        widths = (container_width * 0.88, container_width * 0.12)
        self._format_layout_table(table, widths)
        body = table.cell(0, 0)
        data = self._decode_image(element.payload.get("image_base64"))
        if data:
            paragraph = body.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            width = widths[0] * float(element.payload.get("width_fraction", 1.0)) * fit_scale
            paragraph.add_run().add_picture(io.BytesIO(data), width=Pt(max(width, 0.5)))
        elif element.payload.get("latex"):
            body.paragraphs[0].add_run(str(element.payload["latex"]))
        number_cell = table.cell(0, 1)
        number_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = number_cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        role = roles.get(element.role_id) or self._body_role(roles)
        self._style_run(paragraph.add_run(number), role, fit_scale)
        self._collapse_trailing_paragraph(container)

    def _write_native_table(self, container, element, roles, fit_scale: float, container_width: float) -> None:
        html = element.payload.get("html")
        soup = BeautifulSoup(str(html or "<table><tr><td></td></tr></table>"), "html.parser")
        source = soup.find("table")
        if source is None:
            source = BeautifulSoup("<table><tr><td></td></tr></table>", "html.parser").find("table")
        rows, columns = get_table_dimensions(source)
        rows, columns = max(rows, 1), max(columns, 1)
        table = container.add_table(rows=rows, cols=columns)
        native_marker = OxmlElement("w:tblCaption")
        native_marker.set(qn("w:val"), "docflow-native-table")
        table._tbl.tblPr.append(native_marker)
        rule_style = element.payload.get("table_rule_style") or "grid"
        if rule_style == "grid":
            table.style = "Table Grid"
        elif rule_style == "horizontal":
            head = source.find("thead")
            set_horizontal_table_borders(table, len(head.find_all("tr", recursive=False)) if head else 0)
        else:
            clear_table_borders(table)
        left_indent = float(element.payload.get("left_indent_pt", 0.0))
        if left_indent > 0:
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table_indent = OxmlElement("w:tblInd")
            table_indent.set(qn("w:w"), str(round(left_indent * 20)))
            table_indent.set(qn("w:type"), "dxa")
            table._tbl.tblPr.append(table_indent)
        else:
            table.alignment = {
                "left": WD_TABLE_ALIGNMENT.LEFT,
                "right": WD_TABLE_ALIGNMENT.RIGHT,
            }.get(element.payload.get("alignment"), WD_TABLE_ALIGNMENT.CENTER)
        column_weights = get_table_column_weights(source)
        table_width = container_width * float(element.payload.get("width_fraction", 1.0))
        column_widths = [table_width * weight / sum(column_weights) for weight in column_weights]
        set_table_col_widths(table, column_widths)
        placements = get_table_cell_placements(source)
        role = roles.get(element.role_id) or self._body_role(roles)
        base_size = float(element.payload.get("table_font_size_pt") or (role.font_size_pt if role else 10.5))
        minimum_size = float(element.payload.get("table_min_font_size_pt", 0.5))
        font_size = max(round(base_size * fit_scale * 2) / 2.0, minimum_size)
        fit_size = min(
            (
                (sum(column_widths[column : column + span]) - 2.0)
                * 0.96
                / max(estimate_text_units(cell.get_text(" ", strip=True)), 1.0)
                for _row, column, _row_span, span, cell in placements
            ),
            default=font_size,
        )
        if fit_size >= minimum_size:
            font_size = min(font_size, fit_size)
        source_row_height = float(element.payload.get("table_height_pt") or 0.0) * fit_scale / rows
        row_heights = [max(source_row_height, font_size * 1.2 + 2.0) for _ in range(rows)]
        for row_index, column, row_span, span, cell in placements:
            cell_width = sum(column_widths[column : column + span])
            available_width = max((cell_width - 2.0) * 0.96, 1.0)
            lines = max(1, math.ceil(estimate_text_units(cell.get_text(" ", strip=True)) * font_size / available_width))
            required_per_row = (lines * font_size * 1.2 + 2.0) / max(row_span, 1)
            for target_row in range(row_index, min(row_index + row_span, rows)):
                row_heights[target_row] = max(row_heights[target_row], required_per_row)
        for row, row_height in zip(table.rows, row_heights):
            row.height = Pt(row_height)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row_styles = {
            int(row): (fill, text_color)
            for row, fill, text_color in element.payload.get("table_row_styles", ())
        }
        for row_index, column_index, row_span, column_span, cell_source in placements:
            cell = table.cell(row_index, column_index)
            if row_span > 1 or column_span > 1:
                cell = cell.merge(table.cell(row_index + row_span - 1, column_index + column_span - 1))
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = Pt(font_size * 1.2)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.paragraph_format.widow_control = False
            paragraph.paragraph_format.keep_together = False
            paragraph.paragraph_format.keep_with_next = False
            set_cell_margins(cell, top=20, bottom=20, start=20, end=20)
            run = paragraph.add_run(cell_source.get_text(" ", strip=True))
            self._style_run(
                run,
                role,
                1.0,
                font_size_pt=font_size,
                font_family=element.payload.get("font_family"),
            )
            if row_index in row_styles:
                fill, text_color = row_styles[row_index]
                self._shade(cell._tc.get_or_add_tcPr(), fill)
                run.font.color.rgb = RGBColor.from_string(text_color.lstrip("#"))
        self._collapse_trailing_paragraph(container)

    @staticmethod
    def _format_layout_table(table, widths, gutter_pt: float = 0.0) -> None:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        clear_table_borders(table)
        column_count = len(widths)
        layout_widths = tuple(
            width + gutter_pt * ((column > 0) + (column < column_count - 1)) / 2.0
            for column, width in enumerate(widths)
        )
        set_table_col_widths(table, layout_widths)
        for row in table.rows:
            for column, cell in enumerate(row.cells):
                half_gutter_twips = int(max(gutter_pt, 0.0) * 10)
                set_cell_margins(
                    cell,
                    top=0,
                    bottom=0,
                    start=half_gutter_twips if column > 0 else 0,
                    end=half_gutter_twips if column < len(row.cells) - 1 else 0,
                )

    @staticmethod
    def _style_run(run, role, fit_scale: float, font_size_pt=None, font_family=None) -> None:
        if role is None:
            return
        run.font.name = role.western_font_family
        size = float(font_size_pt) if font_size_pt is not None else role.font_size_pt
        run.font.size = Pt(max(round(size * fit_scale * 2) / 2.0, 0.5))
        run.bold = role.bold
        run.italic = role.italic
        run.font.color.rgb = RGBColor.from_string(role.color.lstrip("#"))
        fonts = run._element.get_or_add_rPr().find(qn("w:rFonts"))
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            run._element.get_or_add_rPr().insert(0, fonts)
        fonts.set(qn("w:eastAsia"), font_family or role.font_family)

    @staticmethod
    def _shade(properties, color) -> None:
        if not color:
            return
        shading = properties.find(qn("w:shd"))
        if shading is None:
            shading = OxmlElement("w:shd")
            properties.append(shading)
        shading.set(qn("w:fill"), str(color).lstrip("#"))

    @staticmethod
    def _body_role(roles):
        return next((role for role_id, role in roles.items() if role_id.startswith("body_")), next(iter(roles.values()), None))

    @staticmethod
    def _visual_text(element) -> str:
        if element.text_structure.orientation == "vertical":
            return "\n".join(character for character in element.text if not character.isspace())
        split_text_rows = element.payload.get("split_text_rows") or ()
        if split_text_rows:
            return "\n".join("\t".join(str(value) for value in row) for row in split_text_rows)
        lines = element.payload.get("lines") or ()
        if element.text_structure.preserve_source_lines:
            return "\n".join(str(line) for line in lines)
        tops = element.payload.get("line_tops_px") or ()
        heights = element.payload.get("line_heights_px") or ()
        if element.kind != "heading" or len(lines) < 2 or len(tops) != len(lines) or len(heights) != len(lines):
            return element.text
        output = str(lines[0])
        row_bottom = float(tops[0]) + float(heights[0])
        for line, top, height in zip(lines[1:], tops[1:], heights[1:]):
            output += ("\n" if float(top) >= row_bottom - min(float(height), row_bottom - float(tops[0])) * 0.10 else " ") + str(line)
            row_bottom = max(row_bottom, float(top) + float(height))
        return output

    @staticmethod
    def _content_bbox(element) -> Rect:
        if element.content_bbox is not None:
            return element.content_bbox
        return Rect.from_sequence(element.payload.get("source_bbox") or (0, 0, 1, 1))

    @staticmethod
    def _decode_image(value):
        try:
            return base64.b64decode(value) if value else None
        except (ValueError, TypeError):
            return None

    @staticmethod
    def _usable_width(geometry) -> float:
        return geometry.width_pt - geometry.margin_left_pt - geometry.margin_right_pt

    @staticmethod
    def _clear_container(container) -> None:
        for paragraph in container.paragraphs:
            paragraph._element.getparent().remove(paragraph._element)

    @staticmethod
    def _collapse_trailing_paragraph(container) -> None:
        if not container.paragraphs or container.paragraphs[-1].text:
            return
        paragraph = container.paragraphs[-1]
        following = paragraph._p.getnext()
        if following is not None and following.tag != qn("w:sectPr"):
            return
        ReflowDocxRenderer._collapse_section_break(paragraph)
