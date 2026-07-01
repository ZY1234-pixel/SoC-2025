"""python-docx 图片插入辅助函数。

提供向段落添加内联图片和创建居中图片段落的便捷函数。
"""
import io
import logging

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docflow.renderer.docx_utils.paragraph_fmt import reset_paragraph_format

logger = logging.getLogger(__name__)


def add_image_to_paragraph(paragraph, image_bytes, width_pt):
    """向已有段落添加一个内联图片。

    图片作为新的 ``Run`` 追加到 *paragraph* 中。

    Args:
        paragraph: python-docx Paragraph 对象。
        image_bytes: 原始图片字节（PNG、JPEG 等）。
        width_pt: 显示宽度（磅）。

    Returns:
        成功返回 True，失败返回 False。
    """
    try:
        run = paragraph.add_run()
        run.add_picture(io.BytesIO(image_bytes), width=Pt(width_pt))
        return True
    except Exception:
        logger.warning("Failed to add inline image to paragraph.", exc_info=True)
        return False


def add_centered_image(container, image_bytes, width_pt):
    """添加一个包含居中内联图片的新段落。

    在 *container* 中创建新段落，重置格式，设置居中对齐，
    行距设为 1.0，然后插入图片。

    Args:
        container: 支持 ``add_paragraph`` 的 python-docx 容器
            （Document body 或 table Cell）。
        image_bytes: 原始图片字节。
        width_pt: 显示宽度（磅）。

    Returns:
        新创建的 Paragraph，失败时返回 None。
    """
    try:
        p = container.add_paragraph()
        reset_paragraph_format(p)
        pf = p.paragraph_format
        pf.line_spacing = 1.0
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        run.add_picture(io.BytesIO(image_bytes), width=Pt(width_pt))
        return p
    except Exception:
        logger.warning("Failed to add centred image.", exc_info=True)
        return None
