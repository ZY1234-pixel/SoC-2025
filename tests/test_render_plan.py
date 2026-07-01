from pathlib import Path
import sys
from io import BytesIO
from zipfile import ZipFile

from lxml import etree
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "Code" / "docflow_src"))

from docflow.model.base import BBox, BlockType
from docflow.model.blocks.equation_block import EquationBlock
from docflow.model.blocks.image_block import ImageBlock
from docflow.model.blocks.text_block import TextBlock, TextLine
from docflow.model.page import Document, Page
from docflow.model.zone import Zone
from docflow.layout.style_inferrer import infer_block_styles
from docflow.pipeline import RecoveryPipeline
from docflow.renderer.context import RenderContext
from docflow.renderer.docx_renderer import DocxRenderer
from docflow.schema.models import BlockStyle
from docflow.utils.render_plan import build_render_plan


def test_render_plan_contains_layout_profile_and_zone_strategy():
    block = TextBlock(
        bbox=BBox(0, 0, 100, 40),
        block_type=BlockType.TEXT,
        lines=[TextLine(text="hello world")],
        col_count=1,
        col_index=0,
        spanned_cols=[0],
    )
    page = Page(index=0, image_width=1000, image_height=1400)
    page.attributes = {"layout_profile": "single_column", "rule_stats": {"category_fix_count": 1}}
    page.zones = [Zone(col_count=1, blocks=[block], has_spanned=False)]
    document = Document(pages=[page], metadata={})

    plan = build_render_plan(document, output_format="docx")
    assert plan["summary"]["page_count"] == 1
    assert plan["pages"][0]["layout_profile"] == "single_column"
    assert plan["pages"][0]["zones"][0]["rendering_strategy"] == "single_col"


def test_render_plan_exposes_region_metadata():
    block = TextBlock(
        bbox=BBox(0, 0, 100, 40),
        block_type=BlockType.TEXT,
        lines=[TextLine(text="hello world")],
        col_count=3,
        col_index=1,
        spanned_cols=[1],
        attributes={"xycutpp_proto": {"region_id": "local_parallel_1", "region_kind": "local_parallel_text_band"}},
    )
    page = Page(index=0, image_width=1000, image_height=1400)
    page.attributes = {"layout_profile": "generic_complex", "rule_stats": {}}
    page.zones = [
        Zone(
            col_count=3,
            blocks=[block],
            has_spanned=False,
            region_id="local_parallel_1",
            region_kind="local_parallel_text_band",
        )
    ]
    document = Document(pages=[page], metadata={})

    plan = build_render_plan(document, output_format="docx")
    zone = plan["pages"][0]["zones"][0]

    assert zone["region_id"] == "local_parallel_1"
    assert zone["region_kind"] == "local_parallel_text_band"


def test_font_size_estimation_uses_full_page_scale_not_margin_scale():
    page = Page(index=0, image_width=800, image_height=1132)
    page.detect_page_size()
    page.margin_top_pt = 120
    page.margin_bottom_pt = 120
    line = TextLine(
        text="字号估计",
        text_region=[[80, 100], [260, 100], [260, 124], [80, 124]],
    )
    block = TextBlock(
        bbox=BBox(80, 96, 260, 128),
        block_type=BlockType.TEXT,
        lines=[line],
    )

    block.estimate_font_size(page.full_coord_mapper)

    assert block.estimated_font_size_pt == 14.0


def test_weak_multicolumn_evidence_profiles_as_single_column():
    page = Page(index=0, image_width=800, image_height=1132)
    text_blocks = [
        TextBlock(
            bbox=BBox(90, 80, 710, 120),
            block_type=BlockType.TITLE,
            lines=[TextLine(text="通用标题文本")],
            col_count=2,
            col_index=0,
            spanned_cols=[0, 1],
        ),
        TextBlock(
            bbox=BBox(95, 170, 705, 360),
            block_type=BlockType.TEXT,
            lines=[TextLine(text="这是一个跨越页面主体宽度的正文段落，用来模拟普通单栏文档中的长段落。")],
            col_count=2,
            col_index=0,
            spanned_cols=[0, 1],
        ),
        TextBlock(
            bbox=BBox(96, 390, 704, 620),
            block_type=BlockType.TEXT,
            lines=[TextLine(text="第二个正文块同样占据大部分页面宽度，不形成左右并行的连续阅读流。")],
            col_count=2,
            col_index=0,
            spanned_cols=[0, 1],
        ),
        TextBlock(
            bbox=BBox(100, 660, 250, 690),
            block_type=BlockType.TEXT,
            lines=[TextLine(text="短小左侧块")],
            col_count=2,
            col_index=0,
            spanned_cols=[0],
        ),
        TextBlock(
            bbox=BBox(520, 710, 690, 740),
            block_type=BlockType.TEXT,
            lines=[TextLine(text="短小右侧块")],
            col_count=2,
            col_index=1,
            spanned_cols=[1],
        ),
    ]
    page.zones = [Zone(col_count=2, blocks=text_blocks, has_spanned=True)]

    assert RecoveryPipeline._infer_layout_profile(page, text_blocks) == "single_column"


def test_weak_multicolumn_collapse_clears_column_metadata():
    page = Page(index=0, image_width=800, image_height=1132)
    block = TextBlock(
        bbox=BBox(520, 320, 690, 350),
        block_type=BlockType.TEXT,
        lines=[TextLine(text="短小右侧块")],
        col_count=2,
        col_index=1,
        spanned_cols=[1],
    )
    page.zones = [Zone(col_count=2, blocks=[block], has_spanned=True)]

    RecoveryPipeline._collapse_to_single_column(page, [block])

    assert page.zones[0].col_count == 1
    assert block.col_count == 1
    assert block.col_index == 0
    assert block.spanned_cols == [0]


def test_weak_title_span_anchors_to_main_column():
    title = TextBlock(
        bbox=BBox(798, 471, 1368, 501),
        block_type=BlockType.TITLE,
        lines=[TextLine(text="4 Numerical Solution and its Validation")],
        col_count=2,
        col_index=0,
        spanned_cols=[0, 1],
    )
    col_bounds = [(144.0, 813.0), (836.0, 1507.0)]

    changed = RecoveryPipeline._collapse_weak_title_span_to_anchor_column(title, col_bounds)

    assert changed
    assert title.col_index == 1
    assert title.spanned_cols == [1]


def test_multiline_header_with_stable_wide_lines_is_centered():
    page = Page(index=0, image_width=800, image_height=1132)
    block = TextBlock(
        bbox=BBox(121, 166, 725, 276),
        block_type=BlockType.HEADER,
        lines=[
            TextLine(text="广东省自然资源厅", text_region=[[126, 166], [611, 166], [611, 218], [126, 218]]),
            TextLine(text="文件", text_region=[[618, 179], [725, 179], [725, 257], [618, 257]]),
            TextLine(text="广东省农业农村厅", text_region=[[127, 222], [610, 222], [610, 274], [127, 274]]),
        ],
    )
    page.zones = [Zone(col_count=1, blocks=[block], has_spanned=False)]

    infer_block_styles(
        page.zones,
        page.coord_mapper,
        page_width_px=page.image_width,
    )

    assert block.style.alignment == "center"


def test_multiline_title_with_stable_centered_lines_is_centered():
    page = Page(index=0, image_width=800, image_height=1132)
    block = TextBlock(
        bbox=BBox(113, 419, 667, 535),
        block_type=BlockType.TITLE,
        lines=[
            TextLine(text="广东省自然资源厅广东省农业农村厅关于", text_region=[[118, 419], [662, 419], [662, 452], [118, 452]]),
            TextLine(text="妥善处理农村不动产登记历史", text_region=[[199, 458], [583, 458], [583, 492], [199, 492]]),
            TextLine(text="遗留问题的若干意见", text_region=[[257, 499], [525, 499], [525, 535], [257, 535]]),
        ],
    )
    page.zones = [Zone(col_count=1, blocks=[block], has_spanned=False)]

    infer_block_styles(
        page.zones,
        page.coord_mapper,
        page_width_px=page.image_width,
    )

    assert block.style.alignment == "center"


def test_local_image_sidecar_does_not_make_whole_page_two_column():
    page = Page(index=0, image_width=1705, image_height=2203)
    figure = ImageBlock(
        bbox=BBox(58, 296, 399, 741),
        block_type=BlockType.FIGURE,
        col_count=2,
        col_index=0,
        spanned_cols=[0],
    )
    side_title = TextBlock(
        bbox=BBox(471, 293, 1551, 385),
        block_type=BlockType.TITLE,
        lines=[TextLine(text="Archives of Environmental Health: An International Journal")],
        col_count=2,
        col_index=1,
        spanned_cols=[1],
    )
    side_text = TextBlock(
        bbox=BBox(471, 394, 1522, 462),
        block_type=BlockType.TEXT,
        lines=[TextLine(text="Publication details, including instructions for authors")],
        col_count=2,
        col_index=1,
        spanned_cols=[1],
    )
    wide_tail = TextBlock(
        bbox=BBox(51, 908, 1634, 1011),
        block_type=BlockType.TEXT,
        lines=[TextLine(text="To cite this article: Shunichi Araki and colleagues")],
        col_count=1,
        col_index=0,
        spanned_cols=[0],
    )
    page.zones = [Zone(col_count=2, blocks=[figure, side_title, side_text, wide_tail], has_spanned=True)]

    assert RecoveryPipeline._has_weak_multicolumn_evidence(page, [figure, side_title, side_text, wide_tail])
    assert RecoveryPipeline._infer_layout_profile(page, [figure, side_title, side_text, wide_tail]) == "single_column"


def test_textbook_mixed_profile_uses_reflow_render_mode():
    assert RecoveryPipeline._render_mode_for_profile("textbook_mixed") == "reflow"

    block = TextBlock(
        bbox=BBox(120, 120, 680, 180),
        block_type=BlockType.TITLE,
        lines=[TextLine(text="附子")],
        col_count=2,
        col_index=1,
        spanned_cols=[1],
    )
    page = Page(index=0, image_width=800, image_height=1132)
    page.attributes = {"layout_profile": "textbook_mixed", "rule_stats": {}}
    page.zones = [Zone(col_count=2, blocks=[block], has_spanned=False)]
    document = Document(pages=[page], metadata={})

    plan = build_render_plan(document, output_format="docx")

    assert plan["pages"][0]["render_mode"] == "reflow"
    assert plan["pages"][0]["zones"][0]["rendering_strategy"] == "single_col"


def test_textbook_side_note_plus_two_body_columns_gets_three_tracks():
    side_a = TextBlock(
        bbox=BBox(57, 642, 435, 1037),
        block_type=BlockType.FOOTNOTE,
        block_id="side_a",
        lines=[TextLine(text="Left side note content with several lines of explanatory material.") for _ in range(5)],
    )
    side_b = TextBlock(
        bbox=BBox(59, 1128, 424, 1365),
        block_type=BlockType.FOOTNOTE,
        block_id="side_b",
        lines=[TextLine(text="More side note content for the same annotation rail.") for _ in range(4)],
    )
    mid_a = TextBlock(
        bbox=BBox(582, 641, 1042, 788),
        block_type=BlockType.TEXT,
        block_id="mid_a",
        lines=[TextLine(text="Middle body column text with enough content.") for _ in range(3)],
    )
    mid_b = TextBlock(
        bbox=BBox(587, 792, 1061, 1080),
        block_type=BlockType.TEXT,
        block_id="mid_b",
        lines=[TextLine(text="Another middle body paragraph line.") for _ in range(4)],
    )
    right_a = TextBlock(
        bbox=BBox(1090, 618, 1578, 805),
        block_type=BlockType.TEXT,
        block_id="right_a",
        lines=[TextLine(text="Right body column content continues here.") for _ in range(3)],
    )
    right_b = TextBlock(
        bbox=BBox(1092, 817, 1574, 1040),
        block_type=BlockType.TEXT,
        block_id="right_b",
        lines=[TextLine(text="More right body column content.") for _ in range(4)],
    )
    blocks = [side_a, side_b, mid_a, mid_b, right_a, right_b]

    bounds = RecoveryPipeline._side_note_three_column_bounds(blocks, 1684)

    assert len(bounds) == 3
    assert bounds[0][1] < bounds[1][0] < bounds[2][0]


def test_reflow_title_alignment_uses_page_container_width():
    title = TextBlock(
        bbox=BBox(523, 233, 601, 261),
        block_type=BlockType.TITLE,
        lines=[TextLine(text="附子", text_region=[[523, 233], [601, 233], [601, 261], [523, 261]])],
        col_count=2,
        col_index=1,
        spanned_cols=[1],
    )
    numbered = TextBlock(
        bbox=BBox(136, 172, 296, 202),
        block_type=BlockType.TITLE,
        lines=[TextLine(text="七、温里药", text_region=[[136, 172], [296, 172], [296, 202], [136, 202]])],
        col_count=2,
        col_index=0,
        spanned_cols=[0],
    )
    zone = Zone(col_count=2, blocks=[numbered, title], has_spanned=False)

    infer_block_styles(
        [zone],
        Page(index=0, image_width=1102, image_height=1631).coord_mapper,
        page_width_px=1102,
        reflow_title_page_width_px=1102,
    )

    assert title.style.alignment == "center"
    assert numbered.style.alignment == "left"


def test_reflow_text_block_merges_visual_ocr_lines():
    lines = [
        TextLine(text="以习近平新时代中国特色社会主义思想为指导，全面贯彻党", text_region=[[90, 100], [690, 100], [690, 124], [90, 124]]),
        TextLine(text="的二十大精神，坚持以人民为中心的发展思想，以切实保障农民", text_region=[[90, 130], [690, 130], [690, 154], [90, 154]]),
        TextLine(text="住宅财产权和宅基地使用权为目标。", text_region=[[90, 160], [500, 160], [500, 184], [90, 184]]),
    ]
    block = TextBlock(
        bbox=BBox(90, 100, 690, 184),
        block_type=BlockType.TEXT,
        lines=lines,
        style=BlockStyle(font_size_pt=18.0, alignment="justify"),
    )
    page = Page(index=0, image_width=800, image_height=1132)
    page.attributes = {"layout_profile": "single_column", "render_mode": "reflow"}
    page.zones = [Zone(col_count=1, blocks=[block], has_spanned=False)]
    doc = Document(pages=[page], metadata={})

    docx_bytes = DocxRenderer().render_bytes(doc, enforce_single_page=False)
    with ZipFile(BytesIO(docx_bytes)) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")

    assert "<w:br" not in xml
    assert "全面贯彻党" in xml
    assert "的二十大精神" in xml


def test_chinese_numbered_same_level_titles_share_font_size():
    page = Page(index=0, image_width=800, image_height=1132)
    first = TextBlock(
        bbox=BBox(120, 200, 280, 230),
        block_type=BlockType.TITLE,
        lines=[TextLine(text="一、总体要求")],
        estimated_font_size_pt=24.0,
    )
    second = TextBlock(
        bbox=BBox(120, 500, 280, 526),
        block_type=BlockType.TITLE,
        lines=[TextLine(text="二、适用范围")],
        estimated_font_size_pt=18.0,
    )
    zone = Zone(col_count=1, blocks=[first, second], has_spanned=False)

    infer_block_styles(
        [zone],
        page.coord_mapper,
        page_width_px=page.image_width,
        font_mapper=page.full_coord_mapper,
    )

    assert first.style.font_size_pt == second.style.font_size_pt
    assert first.style.font_size_pt == 21.0


def test_split_visual_rows_title_alignment_uses_merged_row_edges():
    page = Page(index=0, image_width=800, image_height=1132)
    masthead = TextBlock(
        bbox=BBox(122, 165, 723, 280),
        block_type=BlockType.TITLE,
        lines=[
            TextLine(text="广东省自然资源", text_region=[[122, 167], [561, 165], [561, 222], [123, 224]]),
            TextLine(text="厅", text_region=[[554, 170], [611, 170], [611, 216], [554, 216]]),
            TextLine(text="文件", text_region=[[623, 182], [723, 182], [723, 254], [623, 254]]),
            TextLine(text="广东省农业农村", text_region=[[122, 215], [564, 211], [564, 276], [123, 280]]),
            TextLine(text="厅", text_region=[[557, 217], [613, 217], [613, 271], [557, 271]]),
        ],
        estimated_font_size_pt=28.0,
    )
    zone = Zone(col_count=1, blocks=[masthead], has_spanned=False)

    infer_block_styles(
        [zone],
        page.coord_mapper,
        page_width_px=page.image_width,
        font_mapper=page.full_coord_mapper,
    )

    assert masthead.style.alignment == "center"


def test_multiline_title_uses_modest_font_scale():
    page = Page(index=0, image_width=800, image_height=1132)
    title = TextBlock(
        bbox=BBox(117, 421, 663, 536),
        block_type=BlockType.TITLE,
        lines=[
            TextLine(text="广东省自然资源厅广东省农业农村厅关于", text_region=[[117, 424], [663, 421], [663, 455], [117, 457]]),
            TextLine(text="妥善处理农村不动产登记历史", text_region=[[199, 462], [583, 460], [583, 492], [199, 494]]),
            TextLine(text="遗留问题的若干意见", text_region=[[259, 501], [526, 499], [526, 533], [259, 536]]),
        ],
        estimated_font_size_pt=21.0,
    )

    resolved = DocxRenderer()._resolve_title_font_size_pt(
        block=title,
        page=page,
        font_size_pt=21.0,
        alignment="center",
    )

    assert resolved == 22.26


def test_split_masthead_title_does_not_use_single_line_scale():
    page = Page(index=0, image_width=800, image_height=1132)
    masthead = TextBlock(
        bbox=BBox(122, 165, 723, 280),
        block_type=BlockType.TITLE,
        lines=[
            TextLine(text="广东省自然资源", text_region=[[122, 167], [561, 165], [561, 222], [123, 224]]),
            TextLine(text="厅", text_region=[[554, 170], [611, 170], [611, 216], [554, 216]]),
            TextLine(text="文件", text_region=[[623, 182], [723, 182], [723, 254], [623, 254]]),
            TextLine(text="广东省农业农村", text_region=[[122, 215], [564, 211], [564, 276], [123, 280]]),
            TextLine(text="厅", text_region=[[557, 217], [613, 217], [613, 271], [557, 271]]),
        ],
        estimated_font_size_pt=28.0,
    )

    resolved = DocxRenderer()._resolve_title_font_size_pt(
        block=masthead,
        page=page,
        font_size_pt=28.0,
        alignment="center",
    )

    assert resolved == 28.0


def test_title_visual_fragments_preserve_source_gap_in_docx():
    title = TextBlock(
        bbox=BBox(125, 28, 1386, 138),
        block_type=BlockType.TITLE,
        lines=[
            TextLine(text="助力双方交往", text_region=[[125, 28], [719, 28], [719, 138], [125, 138]]),
            TextLine(text="搭建友谊桥梁", text_region=[[794, 36], [1386, 36], [1386, 135], [794, 135]]),
        ],
        style=BlockStyle(font_size_pt=35.0, alignment="center"),
    )
    page = Page(index=0, image_width=1524, image_height=1368)
    page.attributes = {"layout_profile": "newspaper_mixed"}
    page.zones = [Zone(col_count=1, blocks=[title], has_spanned=False)]
    doc = Document(pages=[page], metadata={})

    docx_bytes = DocxRenderer().render_bytes(doc, enforce_single_page=False)
    with ZipFile(BytesIO(docx_bytes)) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    text = "".join(root.xpath(".//w:t/text()", namespaces=ns))
    assert "助力双方交往\u3000搭建友谊桥梁" in text


def test_newspaper_masthead_title_resists_overaggressive_page_fit():
    renderer = DocxRenderer()
    renderer._fit_scale = 0.70
    renderer._font_floor = 6.5
    page = Page(index=0, image_width=1524, image_height=1368)
    page.attributes = {"layout_profile": "newspaper_mixed"}
    title = TextBlock(
        bbox=BBox(125, 28, 1386, 138),
        block_type=BlockType.TITLE,
        lines=[
            TextLine(text="助力双方交往", text_region=[[125, 28], [719, 28], [719, 138], [125, 138]]),
            TextLine(text="搭建友谊桥梁", text_region=[[794, 36], [1386, 36], [1386, 135], [794, 135]]),
        ],
        style=BlockStyle(font_size_pt=35.0, alignment="center"),
    )

    scaled = renderer._scale_title_font(title, page, 42.0)

    assert round(scaled, 1) == 37.0


def test_multiline_footnote_rows_keep_independent_alignment():
    footnote = TextBlock(
        bbox=BBox(808, 701, 1486, 750),
        block_type=BlockType.FOOTNOTE,
        lines=[
            TextLine(
                text="在厄立特里亚不久前举办的第六届中国风筝文化节上，当地小学生体验风筝制作。",
                text_region=[[813, 705], [1448, 705], [1448, 725], [813, 725]],
            ),
            TextLine(
                text="中国驻厄立特里亚大使馆供图",
                text_region=[[1248, 728], [1486, 728], [1486, 749], [1248, 749]],
            ),
        ],
        style=BlockStyle(font_size_pt=9.5, alignment="center"),
    )
    page = Page(index=0, image_width=1524, image_height=1368)
    page.zones = [Zone(col_count=1, blocks=[footnote], has_spanned=False)]
    doc = Document(pages=[page], metadata={})

    docx_bytes = DocxRenderer().render_bytes(doc, enforce_single_page=False)
    with ZipFile(BytesIO(docx_bytes)) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    seen = {}
    for paragraph in root.findall(".//w:p", ns):
        text = "".join(paragraph.xpath(".//w:t/text()", namespaces=ns))
        if not text:
            continue
        jc = paragraph.find("./w:pPr/w:jc", ns)
        seen[text] = jc.get(f"{{{ns['w']}}}val") if jc is not None else ""

    assert seen["在厄立特里亚不久前举办的第六届中国风筝文化节上，当地小学生体验风筝制作。"] == "center"
    assert seen["中国驻厄立特里亚大使馆供图"] == "right"


def test_newspaper_multicol_body_merges_ocr_lines_into_paragraph():
    body = TextBlock(
        bbox=BBox(8, 198, 360, 342),
        block_type=BlockType.TEXT,
        lines=[
            TextLine(text="身着中国传统民族服装的厄立特里亚青", text_region=[[8, 198], [360, 198], [360, 220], [8, 220]]),
            TextLine(text="年依次登台表演中国民族舞、现代舞、扇子舞", text_region=[[8, 222], [360, 222], [360, 244], [8, 244]]),
            TextLine(text="等，曼妙的舞姿赢得现场观众阵阵掌声。", text_region=[[8, 246], [330, 246], [330, 268], [8, 268]]),
        ],
        style=BlockStyle(font_size_pt=10.0, alignment="justify"),
        col_count=4,
        col_index=0,
        spanned_cols=[0],
    )
    page = Page(index=0, image_width=1524, image_height=1368)
    page.attributes = {"layout_profile": "newspaper_mixed"}
    page.zones = [Zone(col_count=4, blocks=[body], has_spanned=False)]
    doc = Document(pages=[page], metadata={})

    docx_bytes = DocxRenderer().render_bytes(doc, enforce_single_page=False)
    with ZipFile(BytesIO(docx_bytes)) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    paragraphs = [
        paragraph
        for paragraph in root.findall(".//w:p", ns)
        if "身着中国传统民族服装" in "".join(paragraph.xpath(".//w:t/text()", namespaces=ns))
    ]

    assert len(paragraphs) == 1
    assert not paragraphs[0].findall(".//w:br", ns)


def test_title_block_does_not_force_bold_without_style_signal():
    title = TextBlock(
        bbox=BBox(120, 200, 320, 230),
        block_type=BlockType.TITLE,
        lines=[TextLine(text="一、总体要求")],
        style=BlockStyle(font_size_pt=16.0, alignment="left", bold=False),
    )
    page = Page(index=0, image_width=800, image_height=1132)
    page.zones = [Zone(col_count=1, blocks=[title], has_spanned=False)]
    doc = Document(pages=[page], metadata={})

    docx_bytes = DocxRenderer().render_bytes(doc, enforce_single_page=False)
    with ZipFile(BytesIO(docx_bytes)) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")

    assert "一、总体要求" in xml
    assert '<w:b w:val="0"/>' in xml
    assert '<w:b/>' not in xml
    assert '<w:b w:val="1"/>' not in xml


def test_title_block_preserves_explicit_bold_style():
    title = TextBlock(
        bbox=BBox(120, 200, 320, 230),
        block_type=BlockType.TITLE,
        lines=[TextLine(text="显式加粗标题")],
        style=BlockStyle(font_size_pt=16.0, alignment="left", bold=True),
    )
    page = Page(index=0, image_width=800, image_height=1132)
    page.zones = [Zone(col_count=1, blocks=[title], has_spanned=False)]
    doc = Document(pages=[page], metadata={})

    docx_bytes = DocxRenderer().render_bytes(doc, enforce_single_page=False)
    with ZipFile(BytesIO(docx_bytes)) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")

    assert "显式加粗标题" in xml
    assert "<w:b" in xml


def test_renderer_merges_adjacent_visual_text_band_for_docx_layout():
    page = Page(index=0, image_width=1200, image_height=1600)
    figure = ImageBlock(
        bbox=BBox(80, 1000, 500, 1300),
        block_type=BlockType.FIGURE,
        col_count=2,
        col_index=0,
        spanned_cols=[0],
    )
    caption = TextBlock(
        bbox=BBox(180, 940, 1020, 990),
        block_type=BlockType.TEXT,
        lines=[TextLine(text="图下说明文字应跟随左侧图片，而不是独立撑开一整行。")],
        col_count=1,
        col_index=0,
        spanned_cols=[0],
    )
    side_text = TextBlock(
        bbox=BBox(570, 1010, 1080, 1320),
        block_type=BlockType.TEXT,
        lines=[TextLine(text="右侧正文与图片垂直重叠，应合并到同一个局部双栏图文带。")],
        col_count=2,
        col_index=1,
        spanned_cols=[1],
    )
    zones = [
        Zone(col_count=2, blocks=[figure], has_spanned=False),
        Zone(col_count=1, blocks=[caption], has_spanned=False),
        Zone(col_count=2, blocks=[side_text], has_spanned=False),
    ]

    merged = DocxRenderer()._merge_adjacent_visual_text_zones(zones, page)

    assert len(merged) == 1
    assert merged[0].col_count == 2
    assert merged[0].blocks == [figure, caption, side_text]
    assert figure.spanned_cols == [0]
    assert caption.spanned_cols == [0]
    assert side_text.spanned_cols == [1]


def test_renderer_splits_embedded_visual_text_band_inside_single_zone():
    page = Page(index=0, image_width=1000, image_height=1500)
    intro = TextBlock(
        bbox=BBox(120, 620, 520, 660),
        block_type=BlockType.TEXT,
        block_id="intro",
        lines=[TextLine(text="前一小节收尾文字。")],
        col_count=1,
        col_index=0,
        spanned_cols=[0],
    )
    title = TextBlock(
        bbox=BBox(500, 720, 590, 760),
        block_type=BlockType.TITLE,
        block_id="title",
        lines=[TextLine(text="肉桂")],
        col_count=1,
        col_index=0,
        spanned_cols=[0],
    )
    source = TextBlock(
        bbox=BBox(130, 780, 430, 815),
        block_type=BlockType.TEXT,
        block_id="source",
        lines=[TextLine(text="【来源】樟科植物肉桂的树皮。")],
        col_count=1,
        col_index=0,
        spanned_cols=[0],
    )
    figure = ImageBlock(
        bbox=BBox(650, 780, 930, 1020),
        block_type=BlockType.FIGURE,
        block_id="figure",
        col_count=1,
        col_index=0,
        spanned_cols=[0],
    )
    harvest = TextBlock(
        bbox=BBox(130, 830, 430, 865),
        block_type=BlockType.TEXT,
        block_id="harvest",
        lines=[TextLine(text="【采收加工】一般在秋季剥取。")],
        col_count=1,
        col_index=0,
        spanned_cols=[0],
    )
    caption = TextBlock(
        bbox=BBox(700, 1035, 880, 1070),
        block_type=BlockType.FIGURE_CAPTION,
        block_id="caption",
        lines=[TextLine(text="图 2-169 肉桂")],
        col_count=1,
        col_index=0,
        spanned_cols=[0],
    )
    tail = TextBlock(
        bbox=BBox(130, 1120, 520, 1180),
        block_type=BlockType.TEXT,
        block_id="tail",
        lines=[TextLine(text="后续贮藏要求。")],
        col_count=1,
        col_index=0,
        spanned_cols=[0],
    )
    zone = Zone(col_count=1, blocks=[intro, title, source, figure, harvest, caption, tail], has_spanned=False)

    split = DocxRenderer()._split_embedded_visual_text_bands([zone], page)

    assert [len(item.blocks) for item in split] == [1, 1, 4, 1]
    assert split[1].col_count == 1
    assert [block.block_id for block in split[1].blocks] == [title.block_id]
    assert split[2].col_count == 2
    assert [block.block_id for block in split[2].blocks] == [source.block_id, figure.block_id, harvest.block_id, caption.block_id]
    assert figure.spanned_cols == [1]
    assert source.spanned_cols == [0]
    assert harvest.spanned_cols == [0]
    assert caption.spanned_cols == [1]


def test_local_visual_band_short_title_renders_above_sidecar_band():
    page = Page(index=0, image_width=1102, image_height=1631)
    title = TextBlock(
        bbox=BBox(524, 233, 601, 261),
        block_type=BlockType.TITLE,
        block_id="title",
        lines=[TextLine(text="附子")],
    )
    section = TextBlock(
        bbox=BBox(137, 173, 296, 202),
        block_type=BlockType.TITLE,
        block_id="section",
        lines=[TextLine(text="七、温里药")],
    )
    figure = ImageBlock(
        bbox=BBox(713, 312, 911, 523),
        block_type=BlockType.FIGURE,
        block_id="figure",
    )
    source = TextBlock(
        bbox=BBox(172, 274, 547, 300),
        block_type=BlockType.TEXT,
        block_id="source",
        lines=[TextLine(text="【来源】毛茛科植物乌头的子根的加工品。")],
    )
    caption = TextBlock(
        bbox=BBox(730, 539, 893, 564),
        block_type=BlockType.FIGURE_CAPTION,
        block_id="caption",
        lines=[TextLine(text="图 2-168 黑顺片")],
    )
    zone = Zone(col_count=1, blocks=[section, title, source, figure, caption], has_spanned=False)

    split = DocxRenderer()._split_embedded_visual_text_bands([zone], page)

    assert [block.block_id for block in split[0].blocks] == [section.block_id]
    assert [block.block_id for block in split[1].blocks] == [title.block_id]
    assert split[1].col_count == 1
    assert [block.block_id for block in split[2].blocks] == [source.block_id, figure.block_id, caption.block_id]
    assert figure.col_index == 1
    assert title.col_index == 0
    assert source.col_index == 0


def test_single_column_page_with_local_sidecar_keeps_prefix_above_band():
    page = Page(index=0, image_width=1705, image_height=2203)
    prefix = TextBlock(
        bbox=BBox(53, 89, 1608, 216),
        block_type=BlockType.TEXT,
        block_id="prefix",
        lines=[TextLine(text="On: 27 December 2014, At: 18:34")],
    )
    figure = ImageBlock(
        bbox=BBox(58, 296, 399, 741),
        block_type=BlockType.FIGURE,
        block_id="figure",
    )
    side_title = TextBlock(
        bbox=BBox(471, 293, 1551, 385),
        block_type=BlockType.TITLE,
        block_id="side_title",
        lines=[TextLine(text="Archives of Environmental Health: An International Journal")],
    )
    side_text = TextBlock(
        bbox=BBox(471, 394, 1522, 462),
        block_type=BlockType.TEXT,
        block_id="side_text",
        lines=[TextLine(text="Publication details, including instructions for authors")],
    )
    side_author = TextBlock(
        bbox=BBox(472, 646, 1327, 686),
        block_type=BlockType.TEXT,
        block_id="side_author",
        lines=[TextLine(text="Shunichi Araki, Katsuyuki Murata")],
    )
    tail = TextBlock(
        bbox=BBox(51, 908, 1634, 1011),
        block_type=BlockType.TEXT,
        block_id="tail",
        lines=[TextLine(text="To cite this article: Shunichi Araki")],
    )
    zone = Zone(col_count=1, blocks=[figure, prefix, side_title, side_text, side_author, tail], has_spanned=False)

    split = DocxRenderer()._split_embedded_visual_text_bands([zone], page)

    assert [block.block_id for block in split[0].blocks] == ["prefix"]
    assert split[1].col_count == 2
    assert [block.block_id for block in split[1].blocks] == ["figure", "side_title", "side_text", "side_author"]
    assert [block.block_id for block in split[2].blocks] == ["tail"]
    assert side_title.col_index == 1


def test_local_visual_band_caps_short_body_font_size():
    text = TextBlock(
        bbox=BBox(473, 734, 955, 771),
        block_type=BlockType.TEXT,
        lines=[TextLine(text="b Tokyo Rosai Hospital, Tokyo, Japan")],
        style=BlockStyle(font_size_pt=16.0, alignment="left"),
    )
    page = Page(index=0, image_width=1705, image_height=2203)
    ctx = RenderContext(
        coord_mapper=page.coord_mapper,
        page=page,
        col_width_pt=300.0,
        col_left_px=468.0,
        col_right_px=1551.0,
        in_table_cell=True,
        local_visual_band=True,
    )
    doc = Document(pages=[page], metadata={})
    page.zones = [Zone(col_count=1, blocks=[text], has_spanned=False)]

    renderer = DocxRenderer()
    docx_bytes = renderer.render_bytes(doc, enforce_single_page=False)
    with ZipFile(BytesIO(docx_bytes)) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    sizes = [
        int(val)
        for val in root.xpath(".//w:sz/@w:val", namespaces=ns)
    ]

    # Direct page rendering has no local band context, so verify the helper path directly too.
    from docx import Document as DocxDocument

    tmp_doc = DocxDocument()
    renderer._render_text_block(tmp_doc, text, ctx, 0.0)
    buf = BytesIO()
    tmp_doc.save(buf)
    with ZipFile(BytesIO(buf.getvalue())) as archive:
        local_root = etree.fromstring(archive.read("word/document.xml"))
    local_sizes = [
        int(val)
        for val in local_root.xpath(".//w:sz/@w:val", namespaces=ns)
    ]

    assert max(sizes) >= 32
    assert max(local_sizes) <= 22


def test_edge_decorative_text_is_rendered_as_sidecar_not_body_flow():
    page = Page(index=0, image_width=1700, image_height=2200)
    note = TextBlock(
        bbox=BBox(0, 497, 116, 623),
        block_type=BlockType.TEXT,
        block_id="note",
        lines=[TextLine(text="星学")],
    )
    body = TextBlock(
        bbox=BBox(478, 500, 1485, 599),
        block_type=BlockType.TEXT,
        block_id="body",
        lines=[TextLine(text="Price Includes: meals in local restaurants; return flights or railway")],
    )
    zone = Zone(col_count=1, blocks=[note, body], has_spanned=False)

    zones = DocxRenderer()._prepare_edge_decorative_text_zones([zone], page)

    assert len(zones) == 1
    assert zones[0].region_kind == "decorative_sidecar"
    assert zones[0].col_count == 2
    assert [(block.block_id, block.col_index) for block in zones[0].blocks] == [("note", 0), ("body", 1)]
    assert zones[0].blocks[0].attributes["docx_decorative_role"] == "left_sidecar"


def test_left_side_note_label_is_rendered_as_sidecar():
    page = Page(index=0, image_width=1700, image_height=2200)
    note = TextBlock(
        bbox=BBox(183, 219, 402, 258),
        block_type=BlockType.TEXT,
        block_id="note",
        lines=[TextLine(text="随堂笔记")],
    )
    body = TextBlock(
        bbox=BBox(476, 231, 1486, 383),
        block_type=BlockType.TEXT,
        block_id="body",
        lines=[TextLine(text="A regular English reading passage with enough content to form the main body rail.")],
    )
    zone = Zone(col_count=1, blocks=[note, body], has_spanned=False)

    zones = DocxRenderer()._prepare_edge_decorative_text_zones([zone], page)

    assert zones[0].region_kind == "decorative_sidecar"
    assert [(block.block_id, block.col_index) for block in zones[0].blocks] == [("note", 0), ("body", 1)]


def test_short_body_font_size_is_capped_by_page_body_size():
    page = Page(index=0, image_width=1700, image_height=2200)
    normal = TextBlock(
        bbox=BBox(470, 650, 1480, 900),
        block_type=BlockType.TEXT,
        lines=[
            TextLine(text="A regular paragraph line with enough content."),
            TextLine(text="Another regular paragraph line with enough content."),
            TextLine(text="A third regular paragraph line with enough content."),
        ],
        style=BlockStyle(font_size_pt=10.0),
    )
    short = TextBlock(
        bbox=BBox(478, 392, 1486, 493),
        block_type=BlockType.TEXT,
        lines=[TextLine(text="● Two days from only £ 599 per person")],
        style=BlockStyle(font_size_pt=18.0),
    )
    page.zones = [Zone(col_count=1, blocks=[normal, short], has_spanned=False)]

    capped = DocxRenderer()._resolve_body_font_size_pt(short, page, 18.0)

    assert round(capped, 1) == 11.2


def test_local_visual_band_visual_blocks_keep_inner_gutter():
    page = Page(index=0, image_width=1705, image_height=2203)
    figure = ImageBlock(
        bbox=BBox(58, 296, 399, 741),
        block_type=BlockType.FIGURE,
    )
    ctx = RenderContext(
        coord_mapper=page.coord_mapper,
        page=page,
        col_width_pt=137.25,
        col_left_px=58.0,
        col_right_px=399.0,
        in_table_cell=True,
        local_visual_band=True,
    )

    width = DocxRenderer()._visual_block_width_pt(
        figure,
        ctx,
        apply_fit_scale=False,
    )

    assert width <= 137.25 * 0.9


def test_local_visual_band_title_with_stable_left_edge_is_left_aligned():
    page = Page(index=0, image_width=1705, image_height=2203)
    title = TextBlock(
        bbox=BBox(471, 489, 1645, 638),
        block_type=BlockType.TITLE,
        lines=[
            TextLine(
                text="Dose-Response Relationship between Tobacco",
                text_region=[[477, 496], [1429, 496], [1429, 526], [477, 526]],
            ),
            TextLine(
                text="Consumption and Melanin Pigmentation in the Attached",
                text_region=[[477, 546], [1640, 546], [1640, 578], [477, 578]],
            ),
            TextLine(
                text="Gingiva",
                text_region=[[475, 597], [632, 597], [632, 638], [475, 638]],
            ),
        ],
    )
    ctx = RenderContext(
        coord_mapper=page.coord_mapper,
        page=page,
        col_width_pt=390.0,
        col_left_px=471.0,
        col_right_px=1645.0,
        in_table_cell=True,
        local_visual_band=True,
    )

    alignment = DocxRenderer()._local_visual_title_alignment(
        title,
        ctx,
        WD_ALIGN_PARAGRAPH.CENTER,
    )

    assert alignment == WD_ALIGN_PARAGRAPH.LEFT


def test_local_visual_band_reflows_left_title_and_author_line():
    page = Page(index=0, image_width=1705, image_height=2203)
    title = TextBlock(
        bbox=BBox(471, 489, 1645, 638),
        block_type=BlockType.TITLE,
        lines=[
            TextLine(
                text="Dose-Response Relationship between Tobacco",
                text_region=[[477, 496], [1429, 496], [1429, 526], [477, 526]],
            ),
            TextLine(
                text="Consumption and Melanin Pigmentation in the Attached",
                text_region=[[477, 546], [1640, 546], [1640, 578], [477, 578]],
            ),
            TextLine(
                text="Gingiva",
                text_region=[[475, 597], [632, 597], [632, 638], [475, 638]],
            ),
        ],
        style=BlockStyle(alignment="center", font_size_pt=12.0),
    )
    author = TextBlock(
        bbox=BBox(472, 646, 1327, 686),
        block_type=BlockType.TEXT,
        lines=[
            TextLine(
                text="Shunichi Araki a , Katsuyuki Murata a , Koichi Ushiob & Ryoji Sakai",
                text_region=[[473, 650], [1323, 650], [1323, 680], [473, 680]],
            )
        ],
        style=BlockStyle(alignment="center", font_size_pt=10.5),
    )
    ctx = RenderContext(
        coord_mapper=page.coord_mapper,
        page=page,
        col_width_pt=390.0,
        col_left_px=471.0,
        col_right_px=1645.0,
        in_table_cell=True,
        local_visual_band=True,
    )
    renderer = DocxRenderer()

    from docx import Document as DocxDocument

    tmp_doc = DocxDocument()
    renderer._render_text_block(tmp_doc, title, ctx, 0.0)
    renderer._render_text_block(tmp_doc, author, ctx, 0.0)
    buf = BytesIO()
    tmp_doc.save(buf)
    with ZipFile(BytesIO(buf.getvalue())) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    paragraphs = []
    for para in root.findall(".//w:p", ns):
        text = "".join(t.text or "" for t in para.findall(".//w:t", ns))
        if not text:
            continue
        jc = para.find(".//w:jc", ns)
        paragraphs.append((
            text,
            jc.get("{%s}val" % ns["w"]) if jc is not None else "left",
            len(para.findall(".//w:br", ns)),
        ))

    title_para = next(item for item in paragraphs if item[0].startswith("Dose-Response"))
    author_para = next(item for item in paragraphs if item[0].startswith("Shunichi Araki"))
    assert title_para[1] == "left"
    assert title_para[2] == 0
    assert author_para[1] == "left"


def test_field_like_top_text_preserves_visual_breaks_in_reflow():
    page = Page(index=0, image_width=1705, image_height=2203)
    block = TextBlock(
        bbox=BBox(53, 62, 1608, 216),
        block_type=BlockType.TEXT,
        lines=[
            TextLine(
                text="This article was downloaded by: [130.132.123.28]",
                text_region=[[53, 62], [621, 62], [621, 88], [53, 88]],
            ),
            TextLine(
                text="On: 27 December 2014, At: 18:34",
                text_region=[[53, 94], [488, 94], [488, 120], [53, 120]],
            ),
            TextLine(
                text="Publisher: Routledge",
                text_region=[[53, 126], [318, 126], [318, 152], [53, 152]],
            ),
            TextLine(
                text="Informa Ltd Registered in England and Wales Registered Number: 1072954 Registered office: Mortimer House,",
                text_region=[[53, 158], [1608, 158], [1608, 184], [53, 184]],
            ),
            TextLine(
                text="37-41 Mortimer Street, London W1T 3JH, UK",
                text_region=[[53, 190], [625, 190], [625, 216], [53, 216]],
            ),
        ],
        style=BlockStyle(alignment="left", font_size_pt=10.0),
    )
    page.attributes = {"render_mode": "reflow"}
    ctx = RenderContext(
        coord_mapper=page.coord_mapper,
        page=page,
        col_width_pt=520.0,
        col_left_px=53.0,
        col_right_px=1608.0,
    )
    ctx.render_mode = "reflow"

    from docx import Document as DocxDocument

    tmp_doc = DocxDocument()
    DocxRenderer()._render_text_block(tmp_doc, block, ctx, 0.0)
    buf = BytesIO()
    tmp_doc.save(buf)
    with ZipFile(BytesIO(buf.getvalue())) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    para = next(
        p for p in root.findall(".//w:p", ns)
        if "This article was downloaded by" in "".join(t.text or "" for t in p.findall(".//w:t", ns))
    )

    assert len(para.findall(".//w:br", ns)) == 4


def test_renderer_narrow_strip_block_does_not_create_spanned_segment():
    header = TextBlock(
        bbox=BBox(820, 30, 1080, 70),
        block_type=BlockType.HEADER,
        lines=[TextLine(text="The Economist November 11th 2023")],
        col_count=3,
        col_index=2,
        spanned_cols=[0, 1, 2],
    )

    cols = DocxRenderer._layout_block_cols(header, num_cols=3, page_width_px=1200)

    assert cols == [2]


def test_long_narrow_title_uses_body_sized_cap():
    page = Page(index=0, image_width=1386, image_height=1859)
    title = TextBlock(
        bbox=BBox(67, 159, 108, 337),
        block_type=BlockType.TITLE,
        lines=[TextLine(text="GUSHIZHONGDEKEXUE 故事中的科学")],
        estimated_font_size_pt=30.0,
    )

    resolved = DocxRenderer()._resolve_title_font_size_pt(
        block=title,
        page=page,
        font_size_pt=30.0,
        alignment="center",
    )

    assert resolved == 12.075


def test_multicolumn_latin_body_keeps_inferred_font_size():
    page = Page(index=0, image_width=5000, image_height=6567)
    block = TextBlock(
        bbox=BBox(300, 3000, 1720, 4260),
        block_type=BlockType.TEXT,
        lines=[TextLine(text="A long Latin magazine paragraph should not stay at the tiny fit floor.")],
        col_count=3,
        col_index=0,
        spanned_cols=[0],
    )

    resolved = DocxRenderer()._resolve_body_font_size_pt(
        block=block,
        page=page,
        font_size_pt=8.5,
    )

    assert resolved == 8.5


def test_visual_block_width_uses_source_container_ratio_for_spanned_figures():
    page = Page(index=0, image_width=1524, image_height=1368)
    figure = ImageBlock(
        bbox=BBox(775, 201, 1501, 685),
        block_type=BlockType.FIGURE,
        image_data=b"fake",
    )
    ctx = RenderContext(
        coord_mapper=page.coord_mapper,
        page=page,
        col_width_pt=260.0,
        col_left_px=760.0,
        col_right_px=1505.0,
        in_table_cell=True,
    )

    width = DocxRenderer()._visual_block_width_pt(figure, ctx, apply_fit_scale=False)

    assert round(width, 1) == 253.4


def test_spanned_figure_width_can_use_full_page_source_ratio():
    page = Page(index=0, image_width=1524, image_height=1368)
    page.page_width_pt = 841.9
    page.page_height_pt = 595.3
    page.orientation = "landscape"
    page.margin_left_pt = 18.0
    page.margin_right_pt = 18.0
    figure = ImageBlock(
        bbox=BBox(775, 201, 1501, 685),
        block_type=BlockType.FIGURE,
        image_data=b"fake",
    )
    ctx = RenderContext(
        coord_mapper=page.coord_mapper,
        page=page,
        col_width_pt=330.0,
        col_left_px=769.0,
        col_right_px=1505.0,
        in_table_cell=True,
        span_gap_pt=20.0,
    )
    conservative_ctx = RenderContext(
        coord_mapper=page.coord_mapper,
        page=page,
        col_width_pt=330.0,
        col_left_px=769.0,
        col_right_px=1505.0,
        in_table_cell=True,
    )

    width = DocxRenderer()._visual_block_width_pt(figure, ctx, apply_fit_scale=False)
    conservative_width = DocxRenderer()._visual_block_width_pt(
        figure,
        conservative_ctx,
        apply_fit_scale=False,
    )

    assert width > conservative_width
    assert round(width, 1) == 350.0


def test_formula_width_uses_bbox_ratio_instead_of_fixed_column_fraction():
    page = Page(index=0, image_width=1653, image_height=2206)
    formula = EquationBlock(
        bbox=BBox(145, 360, 238, 421),
        block_type=BlockType.EQUATION,
        image_data=b"fake",
    )
    ctx = RenderContext(
        coord_mapper=page.coord_mapper,
        page=page,
        col_width_pt=150.0,
        col_left_px=144.0,
        col_right_px=812.0,
        in_table_cell=False,
    )

    width = DocxRenderer()._visual_block_width_pt(
        formula,
        ctx,
        apply_fit_scale=True,
        max_ratio=0.72,
        min_width_pt=18.0,
    )

    assert round(width, 1) == 20.9


def test_single_page_fit_search_retries_with_lower_font_floor(monkeypatch):
    renderer = DocxRenderer()
    seen: list[tuple[float, float]] = []

    def fake_build(document, **_options):
        seen.append((renderer._fit_scale, renderer._font_floor))
        return object()

    def fake_overflow(_doc, _expected_pages):
        return renderer._font_floor > 7.0 or renderer._fit_scale > 0.82

    monkeypatch.setattr(renderer, "_build_docx", fake_build)
    monkeypatch.setattr(renderer, "_check_overflow", fake_overflow)

    doc = renderer._render_largest_fitting_scale(
        Document(pages=[], metadata={}),
        expected_pages=1,
        build_options={},
        min_scale=0.70,
        font_floors=(8.5, 7.0),
    )

    assert doc is not None
    assert any(floor == 8.5 for _scale, floor in seen)
    assert any(floor == 7.0 for _scale, floor in seen)
    assert renderer._font_floor == 7.0
