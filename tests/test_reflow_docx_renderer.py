from __future__ import annotations

import base64
import io

import pytest
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from PIL import Image
from bs4 import BeautifulSoup

from docflow.model.stages import (
    AnalysisPage,
    DocumentAnalysis,
    FlowKind,
    FlowSection,
    GridCell,
    PlannedElement,
    Rect,
    SemanticElement,
    TextStructure,
    TypographicRole,
)
from docflow.planning import ReflowPlanner
from docflow.renderer.reflow_docx_renderer import ReflowDocxRenderer
from docflow.renderer.docx_utils.html_table import estimate_text_units, get_table_column_weights


def _png_base64() -> str:
    image = Image.new("RGB", (120, 40), "white")
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    return base64.b64encode(buffer.getvalue()).decode("ascii")


def test_table_column_weights_follow_cell_content() -> None:
    table = BeautifulSoup("<table><tr><td>近12个月最高/最低（元）</td><td>15.06/9.98</td></tr></table>", "html.parser").table

    left, right = get_table_column_weights(table)

    assert left > right * 2


def test_text_width_units_match_latin_and_cjk_font_metrics() -> None:
    assert estimate_text_units("AAAA") == pytest.approx(1.68)
    assert estimate_text_units("中文") == 2.0


def test_single_line_text_is_sized_to_its_source_width() -> None:
    role = TypographicRole("heading", "黑体", "Times New Roman", 16, 1.0)
    element = SemanticElement(
        "heading",
        "heading",
        Rect(100, 100, 350, 150),
        1,
        ("source",),
        "投资评级：买入（维持）",
        "heading",
        payload={"lines": ("投资评级：买入（维持）",), "background_color": "#4671A6"},
    )
    plan = ReflowPlanner().plan(DocumentAnalysis((AnalysisPage(0, 1000, 1400, (element,)),), (role,)))

    document = ReflowDocxRenderer().build(plan)
    paragraph = document.tables[0].cell(0, 0).paragraphs[0]
    run = paragraph.runs[0]

    assert run.font.size.pt < role.font_size_pt * plan.pages[0].fit_scale
    assert paragraph._p.pPr.find(qn("w:shd")) is not None


def test_sparse_grid_keeps_valid_empty_cells() -> None:
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)
    elements = (
        SemanticElement("top-left", "paragraph_group", Rect(0, 0, 400, 200), 1, ("a",), "A", "body"),
        SemanticElement("bottom-right", "paragraph_group", Rect(600, 300, 1000, 500), 2, ("b",), "B", "body"),
        SemanticElement("top-right", "paragraph_group", Rect(600, 0, 1000, 200), 3, ("c",), "C", "body"),
        SemanticElement("bottom-left", "paragraph_group", Rect(0, 600, 400, 800), 4, ("d",), "D", "body"),
    )
    plan = ReflowPlanner().plan(DocumentAnalysis((AnalysisPage(0, 1000, 1400, elements),), (role,)))

    document = ReflowDocxRenderer().build(plan)
    grid = document.tables[0].cell(0, 0).tables[0]

    assert all(cell.paragraphs or cell.tables for row in grid.rows for cell in row.cells)


def test_grid_renderer_merges_column_and_row_spans() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    flow = FlowSection(
        "grid",
        FlowKind.GRID,
        ("span", "side"),
        (100, 100, 100),
        grid_cells=(
            GridCell(0, 2, ("side",), row_span=2),
            GridCell(1, 0, ("span",), column_span=2),
        ),
        row_heights_pt=(30, 40),
    )
    elements = {
        "span": PlannedElement("span", "paragraph_group", "body", "image area"),
        "side": PlannedElement("side", "paragraph_group", "body", "side flow"),
    }
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._render_grid(container, flow, elements, {"body": role}, 1.0)

    assert container.tables[0]._tbl.xpath('.//w:gridSpan[@w:val="2"]')
    assert container.tables[0]._tbl.xpath(".//w:vMerge")
    assert [row.height.pt for row in container.tables[0].rows] == pytest.approx((30, 40))


def test_wrapped_flow_uses_floating_editable_media_table() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    flow = FlowSection(
        "wrapped",
        FlowKind.WRAPPED,
        ("body", "image"),
        floating_element_id="image",
        floating_width_pt=40,
    )
    elements = {
        "body": PlannedElement("body", "paragraph_group", "body", "Editable body"),
        "image": PlannedElement(
            "image",
            "figure_group",
            "body",
            payload={"image_base64": _png_base64(), "caption": "Editable caption"},
        ),
    }
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._render_wrapped(container, flow, elements, {"body": role}, 1.0, 100)

    positioning = container.tables[0]._tbl.tblPr.find(qn("w:tblpPr"))
    assert positioning.get(qn("w:tblpXSpec")) == "right"
    assert container.tables[0].cell(0, 0).paragraphs[-1].text == "Editable caption"
    assert container.paragraphs[-1].text == "Editable body"


def test_grid_image_does_not_exceed_its_source_physical_width() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    element = PlannedElement(
        "image",
        "figure_group",
        "body",
        payload={
            "image_base64": _png_base64(),
            "source_bbox": (0, 0, 50, 20),
            "source_scale": 0.5,
            "width_fraction": 1.0,
        },
    )

    ReflowDocxRenderer()._write_image(container, element, 1.0, 100)

    assert container.paragraphs[-1].runs[0]._r.xpath(".//wp:extent")[0].get("cx") == str(int(25 * 12700))


def test_left_wrapped_flow_uses_absolute_position_and_editable_text() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    flow = FlowSection(
        "wrapped",
        FlowKind.WRAPPED,
        ("body", "note"),
        floating_element_id="note",
        floating_width_pt=40,
        floating_side="left",
        floating_offset_x_pt=20,
    )
    elements = {
        "body": PlannedElement("body", "paragraph_group", "body", "Editable body"),
        "note": PlannedElement("note", "paragraph_group", "body", "Editable note"),
    }
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._render_wrapped(container, flow, elements, {"body": role}, 1.0, 100)

    positioning = container.tables[0]._tbl.tblPr.find(qn("w:tblpPr"))
    assert positioning.get(qn("w:tblpX")) == "400"
    assert positioning.get(qn("w:tblpXSpec")) is None
    assert container.tables[0].cell(0, 0).paragraphs[-1].text == "Editable note"


def test_native_table_respects_element_width_and_sets_fixed_table_width() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    element = PlannedElement(
        "table",
        "table_group",
        "body",
        payload={"html": "<table><tr><td>A</td><td>B</td></tr></table>", "width_fraction": 0.75},
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._write_native_table(container, element, {"body": role}, 0.8, 100)

    grid = container.tables[0]._tbl.tblGrid
    assert sum(int(column.get(qn("w:w"))) for column in grid.gridCol_lst) == pytest.approx(75 * 20, abs=2)
    table_width = container.tables[0]._tbl.tblPr.find(qn("w:tblW"))
    assert table_width.get(qn("w:type")) == "dxa"
    assert int(table_width.get(qn("w:w"))) == pytest.approx(75 * 20, abs=2)
    assert container.tables[0]._tbl.tblPr.find(qn("w:tblCaption")).get(qn("w:val")) == "docflow-native-table"


def test_native_table_respects_source_left_anchor() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    element = PlannedElement(
        "table",
        "table_group",
        "body",
        payload={
            "html": "<table><tr><td>A</td></tr></table>",
            "width_fraction": 0.5,
            "left_indent_pt": 24,
        },
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._write_native_table(container, element, {"body": role}, 1.0, 100)

    table = container.tables[0]
    assert table.alignment == WD_TABLE_ALIGNMENT.LEFT
    assert table._tbl.tblPr.find(qn("w:tblInd")).get(qn("w:w")) == "480"


def test_native_table_preserves_sparse_source_columns_and_header_span() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    element = PlannedElement(
        "table",
        "table_group",
        "body",
        payload={
            "html": (
                "<table><tr><td>Letter</td><td colspan='3'>Use</td></tr>"
                "<tr><td>K</td><td>kite</td><td></td><td></td></tr>"
                "<tr><td>L</td><td></td><td></td><td>lion</td></tr></table>"
            )
        },
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._write_native_table(container, element, {"body": role}, 1.0, 100)

    table = container.tables[0]
    assert len(table.columns) == 4
    assert table.cell(1, 1).text == "kite"
    assert table.cell(2, 3).text == "lion"
    assert table._tbl.xpath('.//w:gridSpan[@w:val="3"]')


def test_numbered_equation_clears_word_default_cell_paragraph_spacing() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    element = PlannedElement(
        "equation",
        "equation_group",
        "body",
        payload={"image_base64": _png_base64(), "number": "(1)", "width_fraction": 0.8},
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._write_equation(container, element, {"body": role}, 1.0, 100)

    for cell in container.tables[0].rows[0].cells:
        assert cell.paragraphs[0].paragraph_format.space_before.pt == 0
        assert cell.paragraphs[0].paragraph_format.space_after.pt == 0


def test_native_table_uses_inferred_table_font() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    element = PlannedElement(
        "table",
        "table_group",
        "body",
        payload={"html": "<table><tr><td>内容</td></tr></table>", "font_family": "仿宋"},
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._write_native_table(container, element, {"body": role}, 1.0, 100)

    fonts = container.tables[0].cell(0, 0).paragraphs[0].runs[0]._r.rPr.find(qn("w:rFonts"))
    assert fonts.get(qn("w:eastAsia")) == "仿宋"


def test_native_table_uses_source_height_as_row_minimum() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    element = PlannedElement(
        "table",
        "table_group",
        "body",
        payload={"html": "<table><tr><td>A</td></tr><tr><td>B</td></tr></table>", "table_height_pt": 100},
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._write_native_table(container, element, {"body": role}, 0.8, 100)

    assert all(row.height.pt == pytest.approx(40) for row in container.tables[0].rows)


def test_heading_preserves_visual_rows_but_joins_overlapping_lines() -> None:
    element = PlannedElement(
        "heading",
        "heading",
        "heading",
        text="2 Chapter title Continued",
        payload={
            "lines": ("2", "Chapter title", "Continued"),
            "line_tops_px": (10, 12, 40),
            "line_heights_px": (30, 24, 20),
        },
    )

    assert ReflowDocxRenderer._visual_text(element) == "2 Chapter title\nContinued"


def test_question_block_preserves_option_rows() -> None:
    element = PlannedElement(
        "question",
        "paragraph_group",
        "body",
        text="1 Question A. First B. Second C. Third",
        payload={"lines": ("1 Question", "A. First", "B. Second", "C. Third")},
        text_structure=TextStructure(preserve_source_lines=True, is_list=True),
    )

    assert ReflowDocxRenderer._visual_text(element) == "1 Question\nA. First\nB. Second\nC. Third"

    paragraph = Document().add_paragraph()
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)
    ReflowDocxRenderer()._write_text(paragraph, element, {"body": role}, 1.0, 100)
    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT


def test_vertical_text_remains_editable_in_a_vertical_cell() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    element = PlannedElement(
        "aside",
        "paragraph_group",
        "body",
        text="公司证券研究报告",
        payload={"lines": ("公司证券研究报告",), "source_bbox": (20, 100, 60, 700), "source_scale": 0.5},
        text_structure=TextStructure(orientation="vertical"),
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._render_element(container, element, {"body": role}, 1.0, 100)

    assert container.tables[0].cell(0, 0).text == "公司证券研究报告"
    assert container.tables[0]._tbl.xpath('.//w:textDirection[@w:val="tbRl"]')


def test_multiline_heading_font_fits_each_preserved_source_line() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    element = PlannedElement(
        "heading",
        "heading",
        "heading",
        text="A deliberately long heading Continued",
        payload={
            "lines": ("A deliberately long heading", "Continued"),
            "line_tops_px": (0, 20),
            "line_heights_px": (18, 18),
            "alignment": "left",
        },
    )
    role = TypographicRole("heading", "宋体", "Times New Roman", 20, 1.0)

    ReflowDocxRenderer()._render_element(container, element, {"heading": role}, 1.0, 100)

    assert container.paragraphs[-1].runs[0].font.size.pt < role.font_size_pt


def test_column_layout_collapses_word_trailing_paragraph() -> None:
    container = Document().add_table(rows=1, cols=1).cell(0, 0)
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)
    elements = {
        "left": PlannedElement("left", "paragraph_group", "body", text="left", payload={"column": 0}),
        "right": PlannedElement("right", "paragraph_group", "body", text="right", payload={"column": 1}),
    }
    flow = FlowSection("columns", FlowKind.SEQUENTIAL_COLUMNS, tuple(elements), column_widths_pt=(50, 50))

    ReflowDocxRenderer()._render_columns(container, flow, elements, {"body": role}, 1.0)

    spacing = container.paragraphs[-1]._p.pPr.find(qn("w:spacing"))
    assert spacing.get(qn("w:line")) == "1"


def test_reflow_docx_keeps_text_tables_and_equation_numbers_editable(tmp_path) -> None:
    roles = (
        TypographicRole("heading", "黑体", "Times New Roman", 18, 1.0, bold=True),
        TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0),
    )
    elements = (
        SemanticElement("heading", "heading", Rect(100, 100, 900, 160), 1, ("r1",), text="Editable heading", role_id="heading"),
        SemanticElement(
            "table",
            "table_group",
            Rect(100, 220, 900, 500),
            2,
            ("r2",),
            payload={"html": "<table><tr><th>Name</th><th>Value</th></tr><tr><td>A</td><td>7</td></tr></table>"},
        ),
        SemanticElement(
            "equation",
            "equation_group",
            Rect(200, 560, 800, 640),
            3,
            ("r3",),
            payload={"image_base64": _png_base64(), "number": "(7)"},
        ),
    )
    analysis = DocumentAnalysis((AnalysisPage(0, 1000, 1400, elements),), roles)
    plan = ReflowPlanner().plan(analysis)
    output = tmp_path / "result.docx"

    ReflowDocxRenderer().render(plan, str(output))
    document = Document(output)
    text = " ".join(node.text or "" for node in document.element.body.iter(qn("w:t")))

    assert "Editable heading" in text
    assert "Name" in text and "Value" in text
    assert "(7)" in text
    assert len(document.element.body.xpath(".//w:tbl")) >= 3
    native_table = document.element.body.xpath('.//w:tbl[w:tblPr/w:tblStyle[@w:val="TableGrid"]]')[0]
    assert native_table.xpath('.//w:pPr/w:spacing[@w:lineRule="exact"]')


def test_reflow_docx_creates_one_unlinked_section_per_source_page(tmp_path) -> None:
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)
    pages = tuple(
        AnalysisPage(
            index,
            1000,
            1400,
            (
                SemanticElement(f"header-{index}", "header", Rect(100, 20, 900, 60), 1, (f"h{index}",), text=f"H{index}", role_id="body"),
                SemanticElement(f"body-{index}", "paragraph_group", Rect(100, 100, 900, 1200), 2, (f"b{index}",), text=f"B{index}", role_id="body"),
            ),
        )
        for index in range(2)
    )
    output = tmp_path / "two-pages.docx"

    ReflowDocxRenderer().render(ReflowPlanner().plan(DocumentAnalysis(pages, (role,))), str(output))
    document = Document(output)

    assert len(document.sections) == 2
    first_header = " ".join(paragraph.text for paragraph in document.sections[0].header.paragraphs)
    second_header = " ".join(paragraph.text for paragraph in document.sections[1].header.paragraphs)
    assert "H0" in first_header
    assert "H1" in second_header
    assert not document.sections[1].header.is_linked_to_previous


def test_reflow_docx_contains_each_source_page_in_an_exact_height_frame(tmp_path) -> None:
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)
    element = SemanticElement("body", "paragraph_group", Rect(100, 100, 900, 1200), 1, ("r1",), text="Body", role_id="body")
    plan = ReflowPlanner().plan(DocumentAnalysis((AnalysisPage(0, 1000, 1400, (element,)),), (role,)))
    output = tmp_path / "page-frame.docx"

    ReflowDocxRenderer().render(plan, str(output))
    document = Document(output)
    frame = document.tables[0]
    row = frame.rows[0]
    usable_height = plan.pages[0].geometry.height_pt - plan.pages[0].geometry.margin_top_pt - plan.pages[0].geometry.margin_bottom_pt

    assert row.height_rule == WD_ROW_HEIGHT_RULE.EXACTLY
    assert row.height.pt == pytest.approx(usable_height, abs=0.1)
    assert row._tr.xpath("./w:trPr/w:cantSplit")
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0]._p.xpath('./w:pPr/w:spacing[@w:line="1"]')
    assert document.element.body[-3].tag == qn("w:tbl")
    assert not document.element.body.xpath(".//w:sectPr/w:docGrid")


def test_scaled_page_does_not_add_a_trailing_body_paragraph() -> None:
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)
    element = SemanticElement(
        "body",
        "paragraph_group",
        Rect(100, 100, 900, 1300),
        1,
        ("r1",),
        text="x" * 10000,
        role_id="body",
    )
    plan = ReflowPlanner().plan(
        DocumentAnalysis(
            (AnalysisPage(0, 1000, 1400, (element,)),),
            (role,),
        )
    )
    page = plan.pages[0]

    document = ReflowDocxRenderer().build(plan)

    assert page.fit_scale < 1.0
    assert not document.paragraphs


def test_layout_table_gutter_preserves_planned_content_width() -> None:
    table = Document().add_table(rows=1, cols=3)

    ReflowDocxRenderer._format_layout_table(table, (100, 100, 100), 20)

    assert [cell.width.pt for cell in table.rows[0].cells] == [110, 120, 110]
    assert not table._tbl.xpath(".//w:cantSplit")


def test_nested_table_trailing_paragraph_is_collapsed() -> None:
    cell = Document().add_table(rows=1, cols=1).cell(0, 0)
    cell.add_table(rows=1, cols=1)

    ReflowDocxRenderer._collapse_trailing_paragraph(cell)

    assert cell.paragraphs[-1]._p.xpath('./w:pPr/w:spacing[@w:line="1"]')
    assert cell.paragraphs[-1]._p.xpath("./w:pPr/w:rPr/w:vanish")


def test_reflow_docx_writes_planned_vertical_spacing(tmp_path) -> None:
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)
    elements = (
        SemanticElement("first", "heading", Rect(100, 100, 900, 200), 1, ("r1",), text="First", role_id="body"),
        SemanticElement("second", "paragraph_group", Rect(100, 300, 900, 400), 2, ("r2",), text="Second", role_id="body"),
    )
    plan = ReflowPlanner().plan(DocumentAnalysis((AnalysisPage(0, 1000, 1400, elements),), (role,)))
    output = tmp_path / "spacing.docx"

    ReflowDocxRenderer().render(plan, str(output))
    paragraph = next(item for item in Document(output).tables[0].cell(0, 0).paragraphs if item.text == "Second")

    assert plan.pages[0].elements[1].payload["space_before_pt"] > 0
    assert paragraph.paragraph_format.space_before.pt == pytest.approx(
        plan.pages[0].elements[1].payload["space_before_pt"] * plan.pages[0].fit_scale,
        abs=0.1,
    )


def test_background_paragraph_uses_source_width() -> None:
    paragraph = Document().add_paragraph()
    element = PlannedElement(
        "label",
        "paragraph_group",
        "body",
        "Label",
        payload={"background_color": "#892549", "width_fraction": 0.25},
    )
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)

    ReflowDocxRenderer()._write_text(paragraph, element, {"body": role}, 0.8, 100)

    assert paragraph.paragraph_format.right_indent.pt == pytest.approx(80)


def test_reflow_docx_writes_source_bbox_paragraph_indents(tmp_path) -> None:
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)
    element = SemanticElement(
        "author",
        "paragraph_group",
        Rect(300, 100, 700, 150),
        1,
        ("r1",),
        text="Author Name",
        role_id="body",
        payload={"lines": ("Author Name",)},
    )
    body = SemanticElement("body", "paragraph_group", Rect(100, 200, 900, 300), 2, ("r2",), text="Body", role_id="body")
    plan = ReflowPlanner().plan(DocumentAnalysis((AnalysisPage(0, 1000, 1400, (element, body)),), (role,)))
    output = tmp_path / "indents.docx"

    ReflowDocxRenderer().render(plan, str(output))
    paragraph = next(item for item in Document(output).tables[0].cell(0, 0).paragraphs if item.text == "Author Name")

    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert paragraph.paragraph_format.left_indent.pt == pytest.approx(paragraph.paragraph_format.right_indent.pt, abs=0.1)


def test_reflow_docx_uses_deterministic_line_height(tmp_path) -> None:
    role = TypographicRole("body", "宋体", "Times New Roman", 10.5, 1.0)
    element = SemanticElement(
        "body",
        "paragraph_group",
        Rect(100, 100, 900, 200),
        1,
        ("r1",),
        text="First line Second line",
        role_id="body",
        payload={"lines": ("First line", "Second line"), "line_heights_px": (20, 20)},
    )
    plan = ReflowPlanner().plan(DocumentAnalysis((AnalysisPage(0, 1000, 1400, (element,)),), (role,)))
    output = tmp_path / "source-lines.docx"

    ReflowDocxRenderer().render(plan, str(output))
    paragraph = next(item for item in Document(output).tables[0].cell(0, 0).paragraphs if "First line" in item.text)

    assert plan.pages[0].elements[0].payload["line_height_pt"] == pytest.approx(20 * 841.89 / 1400)
    assert paragraph.text == "First line Second line"
    assert paragraph.paragraph_format.line_spacing_rule == WD_LINE_SPACING.EXACTLY
    assert paragraph.paragraph_format.widow_control is False
    assert paragraph.paragraph_format.keep_together is False
    assert paragraph.paragraph_format.keep_with_next is False
